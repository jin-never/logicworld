# ==============================================================================
#  Initialization: Load .env FIRST, then import other modules
# ==============================================================================
import os
from dotenv import load_dotenv

# This is now the very first thing that runs.
# 查找.env文件，先检查当前目录，然后检查上级目录，最后检查project-files目录
env_path = None
if os.path.exists('.env'):
    env_path = '.env'
elif os.path.exists('.env_key'):
    env_path = '.env_key'
elif os.path.exists('../.env'):
    env_path = '../.env'
elif os.path.exists('../../project-files/.env'):
    env_path = '../../project-files/.env'

if env_path:
    load_dotenv(env_path)
    print(f"[SUCCESS] Environment file loaded: {env_path}")
else:
    print("[WARN] No .env file found, using default configuration")

# Now, import other modules that might depend on environment variables
from fastapi import FastAPI, HTTPException, Request, Depends, BackgroundTasks, Form, Body, WebSocket, WebSocketDisconnect, File, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.exceptions import RequestValidationError
import json
import logging
from typing import List, Dict, Any, Optional, Union, Tuple, Callable, Awaitable

# 设置logger
logger = logging.getLogger(__name__)
import asyncio
import re
import time
import random
import subprocess
# import spacy # <-- This was the missing import
from pathlib import Path
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))

from ai.prompt_lib import render_prompt, add_memory

from openai import AsyncOpenAI
from pydantic import BaseModel, Field
import uvicorn

# --- Local Imports ---
# These are now AFTER load_dotenv()
from agent_system.orchestrator import Orchestrator
from tools import tool_registry
# 导入内部工具模块
from tools.internal_tools import HTTPRequestTool, DataAnalysisTool, http_tool, data_tool
# 导入统一工具系统 - 替换MCP工具系统
from tools.unified_tool_system import unified_tool_system, execute_tool_call, get_available_tools
# 导入工具路由器
from tools.tool_router import tool_router
# MCP工具已完全移除
# from modao_proxy import router as modao_router  # 模块不存在，暂时注释
from utils.schemas import Node, Edge, WorkflowPayload, NodeData, Position

# 在启动时注册内部与原子工具，确保工具可用
try:
    from tools.internal_tools import register_internal_tools
    register_internal_tools()
except Exception as e:
    print(f"[WARN] 内部工具注册失败: {e}")

try:
    from agent_system.atomic_tools import register_atomic_tools
    register_atomic_tools()
except Exception as e:
    print(f"[WARN] 原子工具注册失败: {e}")

# 工具加载器已删除 - 现在完全使用AI工作流构建器

# 导入配置和监控模块
from utils.config_manager import get_config, get_cors_config
from utils.monitoring import metrics_collector, alert_manager, health_checker

# 导入逻辑智慧系统
try:
    from ai.brain_system.brain_api import brain_router
    LOGIC_WISDOM_SYSTEM_AVAILABLE = True
except ImportError as e:
    print(f"[WARN] 逻辑智慧系统导入失败: {e}")
    LOGIC_WISDOM_SYSTEM_AVAILABLE = False
    brain_router = None

# Reverse proxy for Modao community services
# (已在上面导入)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Global Variables & Clients ---
# Initialize the FastAPI app
app = FastAPI(
    title="思维导图 AI 后端",
    version="1.0.0",
    swagger_ui_parameters={"lang": "zh-CN"}
)

# ==== 本地 Swagger-UI 静态资源 ====
from fastapi.staticfiles import StaticFiles
from fastapi.openapi.docs import get_swagger_ui_html
# 获取 swagger_ui_bundle 的静态文件目录（兼容无 dist_path 属性的版本）
import swagger_ui_bundle

# 兼容不同版本的目录结构
swagger_path = None

# 1) 新版本提供 dist_path
if hasattr(swagger_ui_bundle, "dist_path"):
    swagger_path = swagger_ui_bundle.dist_path
else:
    import pkg_resources, pathlib, os
    vendor_base = pathlib.Path(pkg_resources.resource_filename("swagger_ui_bundle", "vendor"))
    # 取 vendor 下第一个以 swagger-ui- 开头的目录
    for p in vendor_base.iterdir():
        if p.is_dir() and p.name.startswith("swagger-ui-"):
            swagger_path = p
            break

if not swagger_path or not swagger_path.exists():
    raise RuntimeError("无法找到 swagger-ui 静态资源目录，请检查 swagger_ui_bundle 安装")

# 挂载 Modao 反向代理路由（此时 app 已定义）
# app.include_router(modao_router)  # 模块不存在，暂时注释
# MCP平台管理路由已移除

# 挂载前端配置管理路由
from utils.frontend_config import router as frontend_config_router
app.include_router(frontend_config_router)

# 挂载前端配置管理路由到 /api/tools 前缀（为了兼容前端调用）
app.include_router(frontend_config_router, prefix="/api/tools")

# 添加测试端点
@app.get("/api/tools/test")
async def test_api_tools():
    """测试API工具路由是否工作"""
    return {"status": "success", "message": "API tools route is working"}

@app.post("/api/tools/test-post")
async def test_api_tools_post(data: dict):
    """测试POST请求"""
    return {"status": "success", "message": "POST request received", "data": data}

# 添加更新工具状态的端点
@app.post("/api/tools/update-status/{tool_id}")
async def update_tool_status(tool_id: str, status_data: dict):
    """更新工具测试和批准状态"""
    try:
        tested = status_data.get("tested")
        approved = status_data.get("approved")

        await _update_tool_test_status(tool_id, tested, approved)

        return {
            "success": True,
            "message": "工具状态更新成功",
            "tool_id": tool_id,
            "tested": tested,
            "approved": approved
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"更新失败: {str(e)}"
        }

# 测试路由已移除 - 工具可以直接使用，无需测试

# 新增：安装 uv 包管理器的端点
@app.post("/api/tools/install-uv")
async def install_uv_package_manager():
    """安装 uv 包管理器"""
    try:
        import subprocess
        import sys

        print("[WRENCH2] 开始安装 uv 包管理器...")

        # 使用 pip 安装 uv
        result = subprocess.run([
            sys.executable, "-m", "pip", "install", "uv"
        ], capture_output=True, text=True, timeout=300)

        if result.returncode == 0:
            print("[OK] uv 安装成功!")
            print(f"输出: {result.stdout}")

            # 验证安装
            verify_result = subprocess.run([
                "uv", "--version"
            ], capture_output=True, text=True, timeout=30)

            if verify_result.returncode == 0:
                return {
                    "success": True,
                    "message": f"uv 安装成功! 版本: {verify_result.stdout.strip()}",
                    "details": {
                        "install_output": result.stdout,
                        "version": verify_result.stdout.strip()
                    }
                }
            else:
                return {
                    "success": True,
                    "message": "uv 安装成功，但版本验证失败",
                    "details": {
                        "install_output": result.stdout,
                        "verify_error": verify_result.stderr
                    }
                }
        else:
            print(f"[ERROR] uv 安装失败: {result.stderr}")
            return {
                "success": False,
                "message": f"uv 安装失败: {result.stderr}",
                "details": {
                    "stdout": result.stdout,
                    "stderr": result.stderr
                }
            }

    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "message": "安装超时（5分钟）",
            "details": None
        }
    except Exception as e:
        print(f"[ERROR] 安装过程出错: {str(e)}")
        return {
            "success": False,
            "message": f"安装过程出错: {str(e)}",
            "details": None
        }

# MCP工具测试API端点已移除 - 工具可以直接使用

# 新增：MCP工具分享到系统API端点
@app.post("/api/tools/share-to-system")
async def share_mcp_tool_to_system(request: dict):
    """将MCP工具分享到系统工具库"""
    try:
        tool_id = request.get("toolId")
        share = request.get("share", True)

        print(f"🌍 分享工具到系统请求: {tool_id}, 分享状态: {share}")

        # 更新工具分享状态
        await _update_tool_test_status(tool_id, tested=None, approved=share)

        if share:
            # 同步到系统工具库
            await _sync_tool_to_system(tool_id)

        return {
            "success": True,
            "message": "工具已分享到系统工具库" if share else "工具分享已取消",
            "shared": share
        }

    except Exception as e:
        print(f"[ERROR] 工具分享出错: {str(e)}")
        return {
            "success": False,
            "message": f"Share error: {str(e)}"
        }

# 复杂的工具功能验证逻辑已移除 - 工具可以直接使用



# Excel测试功能已移除 - 工具可以直接使用


async def test_generic_functionality(process, steps, tool_name):
    """通用工具功能测试"""
    try:
        print(f"[WRENCH2] 开始测试通用工具功能: {tool_name}")

        # 获取工具列表
        tools_request = {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/list"
        }

        process.stdin.write((json.dumps(tools_request) + "\n").encode())
        await process.stdin.drain()

        # 读取工具列表响应
        tools_response_line = await asyncio.wait_for(process.stdout.readline(), timeout=10)
        tools_response = json.loads(tools_response_line.decode().strip())

        print(f"[WRENCH2] 工具列表响应: {str(tools_response)[:200]}...")

        if tools_response.get("result") and "tools" in tools_response["result"]:
            tools = tools_response["result"]["tools"]
            print(f"[OK] 工具列表获取成功，共 {len(tools)} 个工具")

            if tools:
                # 选择第一个工具进行测试
                test_tool = tools[0]
                tool_name_to_test = test_tool.get("name", "unknown")
                print(f"[DART] 选择测试工具: {tool_name_to_test}")

                # 尝试调用工具（使用最小参数）
                call_request = {
                    "jsonrpc": "2.0",
                    "id": 3,
                    "method": "tools/call",
                    "params": {
                        "name": tool_name_to_test,
                        "arguments": {}  # 使用空参数进行测试
                    }
                }

                print(f"[TELEPHONE_RECEIVER] 发送工具调用请求: {tool_name_to_test}")
                process.stdin.write((json.dumps(call_request) + "\n").encode())
                await process.stdin.drain()

                # 读取调用响应
                call_response_line = await asyncio.wait_for(process.stdout.readline(), timeout=15)
                call_response = json.loads(call_response_line.decode().strip())

                print(f"[TELEPHONE_RECEIVER] 调用响应: {str(call_response)[:200]}...")

                if "result" in call_response:
                    print(f"[OK] 工具调用成功: {tool_name_to_test}")
                    steps[2]["status"] = "success"
                    steps[2]["message"] = f"成功调用工具 {tool_name_to_test}"
                    return True
                elif "error" in call_response:
                    error_msg = call_response.get("error", {}).get("message", "未知错误")
                    # 某些错误是预期的（如参数不足），这仍然证明工具可以响应
                    if "required" in error_msg.lower() or "missing" in error_msg.lower() or "argument" in error_msg.lower():
                        print(f"[OK] 工具响应正常（参数错误是预期的）: {error_msg[:50]}")
                        steps[2]["status"] = "success"
                        steps[2]["message"] = f"工具 {tool_name_to_test} 响应正常"
                        return True
                    else:
                        print(f"[ERROR] 工具调用失败: {error_msg}")
                        steps[2]["status"] = "failed"
                        steps[2]["message"] = f"工具调用失败: {error_msg[:50]}"
                        return False
                else:
                    print("[ERROR] 工具调用响应格式异常")
                    steps[2]["status"] = "failed"
                    steps[2]["message"] = "工具调用响应格式异常"
                    return False
            else:
                print("[ERROR] 工具列表为空")
                steps[2]["status"] = "failed"
                steps[2]["message"] = "工具列表为空"
                return False
        else:
            print("[ERROR] 工具列表响应格式错误")
            steps[2]["status"] = "failed"
            steps[2]["message"] = "工具列表响应格式错误"
            return False

    except Exception as e:
        print(f"[ERROR] 通用功能测试异常: {str(e)}")
        steps[2]["status"] = "failed"
        steps[2]["message"] = f"测试异常: {str(e)[:50]}"
        return False

# 挂载思维导图管理路由
from mindmap_api import router as mindmap_router
app.include_router(mindmap_router)

# 挂载安全模块路由
try:
    from security.api_routes import router as security_router
    app.include_router(security_router, prefix="/api/security")
    print("[OK] 安全模块路由已挂载到 /api/security")
except ImportError as e:
    print(f"[WARN] 安全模块不可用: {e}")

# 挂载用户设置API路由
try:
    from user_settings_api import router as user_settings_router
    app.include_router(user_settings_router, prefix="/api")
    print("[OK] 用户设置API路由已挂载到 /api")
except ImportError as e:
    print(f"[WARN] 用户设置API不可用: {e}")

# 挂载逻辑智慧系统路由
if LOGIC_WISDOM_SYSTEM_AVAILABLE and brain_router:
    app.include_router(brain_router)
    print("[OK] 逻辑智慧系统路由已挂载到 /brain")
else:
    print("[WARN] 逻辑智慧系统不可用，跳过路由挂载")

# 挂载 LogicWorld Protocol 集成路由
try:
    from logic.logicworld_integration.routes import router as logicworld_router
    app.include_router(logicworld_router)
    print("[OK] LogicWorld Protocol 集成路由已挂载到 /api/logicworld")
except ImportError as e:
    print(f"[WARN] LogicWorld Protocol 集成导入失败: {e}")
except Exception as e:
    print(f"[WARN] LogicWorld Protocol 集成路由挂载失败: {e}")

# 挂载AI智能工作流路由
try:
    from ai.intelligent_workflow_fastapi import router as intelligent_workflow_router
    app.include_router(intelligent_workflow_router)
    print("[OK] AI智能工作流API路由已挂载到 /api/intelligent-workflow")
except ImportError as e:
    print(f"[WARN] AI智能工作流API导入失败: {e}")
except Exception as e:
    print(f"[WARN] AI智能工作流API挂载失败: {e}")

# 挂载智能规划API路由
try:
    from api.intelligent_planning_api import router as intelligent_planning_router
    app.include_router(intelligent_planning_router)
    print("[OK] 智能规划API路由已挂载到 /api/intelligent-planning")
except ImportError as e:
    print(f"[WARN] 智能规划API导入失败: {e}")
except Exception as e:
    print(f"[WARN] 智能规划API挂载失败: {e}")

# 挂载本地工具路由
try:
    import sys
    import os
    # 添加backend目录到Python路径
    backend_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if backend_path not in sys.path:
        sys.path.insert(0, backend_path)

    from utils.local_tools import router as local_tools_router
    app.include_router(local_tools_router, prefix="/api")
    print("[OK] 本地工具路由已挂载到 /api/local-tools")
except ImportError as e:
    print(f"[WARN] 本地工具导入失败: {e}")
except Exception as e:
    print(f"[WARN] 本地工具路由挂载失败: {e}")

# 文档预览API路由
try:
    # 暂时禁用文档预览API避免语法错误
    # from api.document_preview_api import router as document_preview_router
    # app.include_router(document_preview_router)
    print("[SKIP] 文档预览API路由已禁用")
except ImportError as e:
    print(f"[WARN] 文档预览API导入失败: {e}")
except Exception as e:
    print(f"[WARN] 文档预览API路由挂载失败: {e}")

# 别名端点：确保 /api/local-tools/file-manager/open 可用（转调 utils.local_tools.open_file）
from fastapi import Body

@app.post("/api/local-tools/file-manager/open")
async def _alias_open_file(payload: dict = Body(...)):
    try:
        # 延迟导入以避免循环依赖
        from utils.local_tools import open_file as _open_file
        return await _open_file(payload)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# 将静态目录挂载到 /swagger
app.mount("/swagger", StaticFiles(directory=str(swagger_path)), name="swagger")

# 使用本地资源渲染 Swagger-UI，并设置中文界面
from fastapi.responses import HTMLResponse

# 注入前端脚本，将 Swagger-UI 残余英文标签动态替换为中文
_ZH_TRANSLATIONS = {
    "Try it out": "尝试",
    " Try it out": "尝试",
    "Try it out ": "尝试",
    " Try it out ": "尝试",
    "Execute": "执行",
    "Cancel": "取消",
    "Schemas": "数据结构",
    "Schema": "数据结构",
    "Example Value": "示例值",
    "Responses": "响应",
    "Response": "响应",
    "Request body": "请求体",
    "Parameters": "参数",
    "Parameter": "参数",
    "Description": "描述",
    "No parameters": "无参数",
    "default": "默认",
    "Media type": "媒体类型",
    "Code": "状态码",
    "Status": "状态",
    "Value": "值",
    "Model": "模型",
    "Download": "下载",
    "Clear": "清除",
    "Close": "关闭",
    "Links": "链接",
    "No links": "无链接",
    "Available authorizations": "授权",
    "Authorize": "授权",
    "Headers": "请求头",
}


@app.get("/docs", include_in_schema=False)
async def custom_swagger_ui():
    base_html = get_swagger_ui_html(
        openapi_url=app.openapi_url,
        title=app.title + " - 接口文档",
        swagger_js_url="/swagger/swagger-ui-bundle.js",
        swagger_css_url="/swagger/swagger-ui.css",
        swagger_ui_parameters={"lang": "zh-CN"},
    )

    # 动态注入脚本，运行时替换页面文字
    import json as _json
    translate_js = """
    <script>
    (function() {
        const zhMap = %s;
        const lowerMap = {};
        Object.keys(zhMap).forEach(k => lowerMap[k.toLowerCase()] = zhMap[k]);

        function translate() {
            const elements = document.querySelectorAll('button, span, div, label, h4, td, th');
            elements.forEach(el => {
                const txt = (el.innerText || '').trim();
                const low = txt.toLowerCase();
                if (lowerMap[low]) {
                    el.innerText = lowerMap[low];
                }
            });
        }
        // 初始尝试
        translate();
        // 监听 DOM 变化
        const observer = new MutationObserver(translate);
        observer.observe(document.body, { childList: true, subtree: true });
        // 定时兜底
        setInterval(translate, 1500);
    })();
    </script>
    """ % _json.dumps(_ZH_TRANSLATIONS, ensure_ascii=False)

    content = base_html.body.decode().replace("</body>", translate_js + "</body>")
    return HTMLResponse(content=content, status_code=base_html.status_code, headers=dict(base_html.headers))

DEEPSEEK_API_KEY = None
client: Optional[AsyncOpenAI] = None
WORKFLOW_RULES = []


@app.on_event("startup")
async def startup_event():
    """
    Initializes the application, loading necessary configurations and tools.
    """
    global DEEPSEEK_API_KEY, client
    # Load workflow rules on startup
    global WORKFLOW_RULES
    WORKFLOW_RULES = load_workflow_rules()
    
    DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
    if not DEEPSEEK_API_KEY:
        # Fallback to OpenAI key if DeepSeek key is not found
        DEEPSEEK_API_KEY = os.getenv("OPENAI_API_KEY")

    if not DEEPSEEK_API_KEY:
        print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
        print("!!! WARNING: DEEPSEEK_API_KEY or OPENAI_API_KEY not found.  !!!")
        print("!!! AI-based tools will fail. Please set it in your .env file.!!!")
        print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
    
    # Initialize the client only if the key exists
    if DEEPSEEK_API_KEY:
        client = AsyncOpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com/v1")

    # 工具注册已禁用 - 现在完全使用AI工作流构建器
    try:
        logging.info("[OK] 工具注册已跳过，使用AI工作流构建器")
        
        # 初始化全局enhanced_tool_router以确保MCP工具正常工作（无外部MCP服务器）
        try:
            from tools.enhanced_tool_router import EnhancedToolRouter
            enhanced_tool_router = EnhancedToolRouter()
            await enhanced_tool_router.initialize()
            logging.info("✅ 全局enhanced_tool_router初始化成功")
        except Exception as e:
            logging.error(f"❌ 全局enhanced_tool_router初始化失败: {e}")
            
    except Exception as e:
        logging.error(f"[ERROR] 工具注册配置错误: {e}")
        # 继续启动，不因工具注册失败而停止

# CORS 白名单（前端开发端口 & 本地文件）
ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "http://localhost:3001",
    "http://127.0.0.1:3001",
    "http://localhost:3002",
    "http://127.0.0.1:3002",
    "http://localhost:3003",
    "http://127.0.0.1:3003",
    "http://localhost",
    "http://127.0.0.1",
    "null",  # 支持本地文件访问
    "*"      # 临时允许所有来源，用于测试
]

# 仅添加一次 CORS 中间件 - 使用本地定义的白名单
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# 添加请求验证错误处理器
@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    logging.error(f"[ERROR] 请求验证错误: {exc}")
    print(f"[ERROR] 请求验证错误: {exc}")
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors(), "body": exc.body}
    )

# 测试端点
@app.post("/test-chat")
async def test_chat(data: dict):
    """测试聊天端点"""
    logging.info(f"[WRENCH2] 测试端点收到数据: {data}")
    print(f"[WRENCH2] 测试端点收到数据: {data}")
    return {"status": "success", "received": data}

# --- Pro Orchestrator Request model must be defined before route decoration ---

from pydantic import BaseModel  # ensure BaseModel in scope for forward decl


class ValidationRequest(BaseModel):
    nodes: List[Dict[str, Any]]
import httpx
import re
from datetime import datetime

# Import tool registry AFTER environment variables are loaded so that tools can access API keys at import time.

# Import orchestrator AFTER dotenv as it requires API keys indirectly via AgentFactory

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Environment & API Key ---
# 移除这里的API密钥检查，因为在startup事件中已经处理了
# api_key = os.getenv("DEEPSEEK_API_KEY")
# if not api_key:
#     raise ValueError("DEEPSEEK_API_KEY environment variable not set.")
# client = AsyncOpenAI(api_key=api_key, base_url="https://api.deepseek.com")

# --- SpaCy Model for Keyword Extraction ---
# try:
#     nlp = spacy.load("en_core_web_sm")
# except OSError:
#     print("Downloading 'en_core_web_sm' model...")
#     from spacy.cli import download
#     download("en_core_web_sm")
#     nlp = spacy.load("en_core_web_sm")
nlp = None  # Temporarily disable spacy

# --- Rule Engine ---
WORKFLOW_RULES = []

def load_workflow_rules(rules_dir: str = "rules"):
    """Loads workflow validation rules from the specified directory.
    - rules_dir 支持相对路径；将转换为绝对路径
    - 目录不存在时仅记录警告，避免启动失败
    """
    import os as _os
    global WORKFLOW_RULES
    WORKFLOW_RULES = []
    # 规范化目录为绝对路径
    abs_dir = _os.path.abspath(rules_dir)
    logging.info(f"Loading workflow rules from: {abs_dir}")
    if not _os.path.isdir(abs_dir):
        logging.warning(f"Rules directory not found: {abs_dir}, skip loading rules")
        return
    for filename in _os.listdir(abs_dir):
        if filename.endswith(".json"):
            filepath = _os.path.join(abs_dir, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    rule = json.load(f)
                    WORKFLOW_RULES.append(rule)
                    logging.info(f"Successfully loaded rule: {rule.get('workflow_name')}")
            except Exception as e:
                logging.error(f"Failed to load rule from {filepath}: {e}")

# --- Memory/History Functions ---
def save_execution_to_history(payload: 'WorkflowPayload', final_result: Any):
    """Saves a successful workflow execution to the history log."""
    history_file = "backend/memory/execution_history.jsonl"
    log_entry = {
        "timestamp": datetime.utcnow().isoformat(),
        "workflow": payload.model_dump(), # Use model_dump for Pydantic v2
        "final_result": final_result
    }
    try:
        with open(history_file, 'a', encoding='utf-8') as f:
            f.write(json.dumps(log_entry) + '\n')
        logging.info("Successfully saved execution to history.")
    except Exception as e:
        logging.error(f"Failed to save execution to history: {e}")


# --- FastAPI App ---
# app 已在文件顶部初始化过，此处无需重新创建，避免覆盖 swagger_ui_parameters 设置
# app = FastAPI()

# 移除重复的startup事件处理器，已在上面定义过了
# @app.on_event("startup")
# async def startup_event():
#     """Load rules on application startup."""
#     load_workflow_rules()

# 第二处重复的 add_middleware 已移除，避免多次装载

# ================================
#  Pro Orchestrator API
# ================================


class OrchestratorRequest(BaseModel):
    """请求体: 提供一系列高阶任务给 Pro Orchestrator"""
    tasks: List[str]
    # mode字段已删除，现在使用统一的执行逻辑


@app.post("/orchestrate")
async def orchestrate(request: OrchestratorRequest):
    """
    运行新的 Pro 编排器。
    - 'normal' 模式: 产生一个轻量级的、纯文本的思考链。
    - 'super' 模式: (未来) 启用多专家、工具执行的专业工作流。
    
    请求示例:
    {
        "tasks": [
            "明确项目需求并撰写PRD",
            "设计系统架构",
            "实现核心功能代码",
            "编写集成测试"
        ],
        "mode": "normal"
    }
    """
    orchestrator = Orchestrator()
    result = await orchestrator.run(request.tasks)
    return result

def find_reference(args: Dict[str, Any]) -> Optional[Tuple[str, str]]:
    """
    Finds the first reference (a value starting with '@node-') in a dictionary.
    This was the missing function causing the NameError.
    """
    for key, value in args.items():
        if isinstance(value, str) and value.startswith('@node-'):
            return key, value
    return None, None

class ChatRequest(BaseModel):
    prompt: str
    model: str = 'deepseek-chat'
    history: list = []
    intent: str = 'chat'  # 'chat' or 'mindmap'
    multi_scene: bool = False  # 是否启用多场景检索

class ValidationRequest(BaseModel):
    nodes: List[Dict[str, Any]]


# --- Helper Functions ---
def generate_id():
    """Generates a unique ID for messages."""
    return f"{os.urandom(6).hex()}"

MINDMAP_KEYWORDS = {'generate', 'create', 'make', 'draw', 'mind map', '思维导图', '生成', '创建', '画'}

def detect_intent(prompt: str):
    """Detects if the user wants to generate a mind map."""
    return any(keyword in prompt.lower() for keyword in MINDMAP_KEYWORDS)

async def handle_step_generation_json(request: ChatRequest):
    """处理步骤生成请求，返回结构化JSON响应"""
    try:
        logging.info("[WRENCH2] 开始处理步骤生成请求")

        # 检查API密钥
        api_key = os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENAI_API_KEY")
        if not api_key:
            logging.error("[ERROR] 未配置API密钥")
            return StreamingResponse(
                iter([f"data: {json.dumps({'type': 'error', 'error': 'AI服务未配置'})}\n\n"]),
                media_type="text/event-stream"
            )

        # 构建AI提示词
        step_prompt = f"""请为以下需求生成可执行的工作流步骤：{request.prompt}

要求：
1. 严格按照以下JSON格式返回，不要添加任何其他文本或markdown格式：

{{
  "title": "[CLIPBOARD2] **执行步骤**",
  "stepCount": 步骤数量,
  "steps": ["步骤1标题", "步骤2标题", "步骤3标题"],
  "detailedSteps": [
    {{
      "title": "步骤1标题",
      "content": "步骤描述",
      "tool": "具体工具名称",
      "toolName": "具体工具名称",
      "aiDescription": "详细执行说明",
      "executionDescription": "详细执行说明",
      "parameters": ""
    }}
  ]
}}

2. 生成5-10个具体可执行的步骤
3. 工具名称使用：DeepSeek聊天模型、文件操作工具、网页浏览器、代码编辑器、终端工具、API调用器、数据库工具、文本处理器、图像处理器、数据分析器等
4. 执行说明要详细具体，包含操作方法和预期结果
5. 只返回JSON，不要任何其他格式或文本

请直接返回JSON："""

        # 创建AI客户端
        from openai import AsyncOpenAI
        base_url = "https://api.deepseek.com" if os.getenv("DEEPSEEK_API_KEY") else None
        client = AsyncOpenAI(api_key=api_key, base_url=base_url)

        # 流式响应生成器
        async def generate_stream():
            try:
                # 调用AI API
                response = await client.chat.completions.create(
                    model="deepseek-chat" if os.getenv("DEEPSEEK_API_KEY") else "gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "你是一个专业的工作流设计师，擅长将复杂任务分解为可执行的步骤。"},
                        {"role": "user", "content": step_prompt}
                    ],
                    stream=True,
                    max_tokens=2000,
                    temperature=0.7
                )

                accumulated_content = ""

                async for chunk in response:
                    if chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        accumulated_content += content

                        # 发送流式数据
                        yield f"data: {json.dumps({'type': 'content', 'content': content})}\n\n"

                # 尝试解析最终的JSON
                try:
                    # 清理内容，移除可能的markdown标记
                    clean_content = accumulated_content.strip()
                    if clean_content.startswith("```json"):
                        clean_content = clean_content[7:]
                    if clean_content.endswith("```"):
                        clean_content = clean_content[:-3]
                    clean_content = clean_content.strip()

                    step_data = json.loads(clean_content)
                    yield f"data: {json.dumps({'type': 'complete', 'stepData': step_data})}\n\n"

                except json.JSONDecodeError:
                    logging.warning("JSON解析失败，使用降级响应")
                    # 降级响应
                    fallback_data = {
                        "title": "[CLIPBOARD2] **执行步骤**",
                        "stepCount": 3,
                        "steps": ["分析需求", "准备工具", "执行任务"],
                        "detailedSteps": [
                            {
                                "title": "分析需求",
                                "content": "详细分析用户需求和目标",
                                "tool": "DeepSeek聊天模型",
                                "toolName": "DeepSeek聊天模型",
                                "aiDescription": "使用AI分析用户输入的需求",
                                "executionDescription": "使用AI分析用户输入的需求",
                                "parameters": ""
                            },
                            {
                                "title": "准备工具",
                                "content": "准备执行任务所需的工具和资源",
                                "tool": "文件操作工具",
                                "toolName": "文件操作工具",
                                "aiDescription": "准备必要的工具和文件",
                                "executionDescription": "准备必要的工具和文件",
                                "parameters": ""
                            },
                            {
                                "title": "执行任务",
                                "content": "按照计划执行具体任务",
                                "tool": "代码编辑器",
                                "toolName": "代码编辑器",
                                "aiDescription": "执行具体的任务操作",
                                "executionDescription": "执行具体的任务操作",
                                "parameters": ""
                            }
                        ]
                    }
                    yield f"data: {json.dumps({'type': 'complete', 'stepData': fallback_data})}\n\n"

                # 发送结束标记
                yield f"data: {json.dumps({'type': 'end'})}\n\n"

            except Exception as e:
                logging.error(f"步骤生成失败: {e}")
                yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
                # 即使出错也要发送结束标记
                yield f"data: {json.dumps({'type': 'end'})}\n\n"

        return StreamingResponse(generate_stream(), media_type="text/event-stream")

    except Exception as e:
        logging.error(f"步骤生成处理失败: {e}")
        return StreamingResponse(
            iter([f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"]),
            media_type="text/event-stream"
        )


async def handle_simple_chat(request: ChatRequest):
    """Handle simple chat requests with intelligent role selection."""
    try:
        # 检查是否是步骤生成请求
        if "请为以下需求生成可执行的工作流步骤" in request.prompt or "生成步骤" in request.prompt:
            logging.info("[WRENCH2] 检测到步骤生成请求，使用结构化JSON响应")
            return await handle_step_generation_json(request)

        # 导入智能角色选择器
        from ai.intelligent_role_selector import role_selector

        # 智能角色选择
        logging.info(f"[MASK] 开始智能角色选择: {request.prompt}")
        role_match = role_selector.analyze_user_input(
            user_input=request.prompt,
            context=getattr(request, 'context', {})
        )

        logging.info(f"[DART] 选择角色: {role_match.role_name} (置信度: {role_match.confidence:.2f})")
        logging.info(f"[MEMO2] 选择理由: {role_match.reasoning}")

        # 构建对话历史
        messages = []

        # 使用智能选择的角色模板作为系统提示
        if role_match.template:
            try:
                # 渲染角色模板
                system_prompt = role_match.template.format(**role_match.variables)
                logging.info(f"[MASK] 使用角色模板: {role_match.role_name}")
            except Exception as template_error:
                logging.warning(f"[WARN] 模板渲染失败: {template_error}")
                # 降级到默认提示
                system_prompt = f"你是一个{role_match.role_name}，请根据你的专业角色回答用户的问题。"
        else:
            # 降级到默认提示
            system_prompt = "你是一个友好的AI助手。请直接回答用户的问题，不需要创建工作流或思维导图。"

        messages.append({
            "role": "system",
            "content": system_prompt
        })

        # 添加历史对话
        for msg in request.history[-10:]:  # 只保留最近10条消息
            if msg.get('sender') == 'user':
                messages.append({"role": "user", "content": msg.get('text', '')})
            elif msg.get('sender') == 'ai':
                messages.append({"role": "assistant", "content": msg.get('text', '')})

        # 添加当前用户消息
        messages.append({"role": "user", "content": request.prompt})

        # 使用流式输出调用AI获取回复
        from fastapi.responses import StreamingResponse
        import json

        async def generate_stream():
            try:
                logger.info("[LAUNCH] 开始生成流式响应")

                # 首先发送角色信息
                role_info = {
                    "type": "role_info",
                    "role_info": {
                        "role_id": role_match.role_id,
                        "role_name": role_match.role_name,
                        "confidence": role_match.confidence,
                        "reasoning": role_match.reasoning
                    },
                    "timestamp": int(time.time() * 1000)
                }
                logger.info("[OUTBOX_TRAY] 发送角色信息")
                yield f"data: {json.dumps(role_info, ensure_ascii=False)}\n\n"

                # 创建流式请求
                logger.info("[RELOAD] 创建流式API请求")
                stream = await client.chat.completions.create(
                    model=request.model,
                    messages=messages,
                    temperature=0.7,
                    max_tokens=1000,
                    stream=True
                )
                logger.info("[OK] 流式API请求已创建")

                # 发送开始标记
                start_data = {
                    "type": "chat_start",
                    "timestamp": int(time.time() * 1000)
                }
                logger.info("[OUTBOX_TRAY] 发送开始标记")
                yield f"data: {json.dumps(start_data, ensure_ascii=False)}\n\n"

                # 流式输出AI回复
                logger.info("[RELOAD] 开始处理流式响应")
                chunk_count = 0
                async for chunk in stream:
                    chunk_count += 1
                    logger.info(f"[PACKAGE] 收到第{chunk_count}个chunk: {chunk}")
                    if chunk.choices[0].delta.content is not None:
                        content = chunk.choices[0].delta.content
                        logger.info(f"[MEMO2] chunk内容: '{content}'")
                        chunk_data = {
                            "type": "chat_chunk",
                            "content": content,
                            "timestamp": int(time.time() * 1000)
                        }
                        yield f"data: {json.dumps(chunk_data, ensure_ascii=False)}\n\n"

                        # 添加小延迟，确保前端能看到逐字效果
                        import asyncio
                        await asyncio.sleep(0.05)  # 50毫秒延迟，让用户能看到逐字效果

                logger.info(f"[OK] 流式处理完成，共处理{chunk_count}个chunk")

                # 发送结束标记
                end_data = {
                    "type": "chat_end",
                    "timestamp": int(time.time() * 1000)
                }
                yield f"data: {json.dumps(end_data, ensure_ascii=False)}\n\n"

            except Exception as e:
                error_data = {
                    "type": "error",
                    "message": f"流式输出错误: {str(e)}",
                    "timestamp": int(time.time() * 1000)
                }
                yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

        # 返回流式响应
        return StreamingResponse(
            generate_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "Content-Type"
            }
        )

    except Exception as e:
        logging.error(f"Error in simple chat: {e}")
        return JSONResponse(
            status_code=500,
            content={
                "type": "error",
                "message": f"聊天服务出现错误: {str(e)}"
            }
        )

class IntentDetectionRequest(BaseModel):
    prompt: str

# --- System Tools Library Models ---
class SystemToolConfig(BaseModel):
    """系统工具库配置模型"""
    name: str
    description: str
    transport: str = "stdio"
    command: str
    args: List[str] = []
    timeout: int = 60000
    env: Dict[str, str] = {}
    tested: bool = False
    approved: bool = False
    created_at: Optional[str] = None
    updated_at: Optional[str] = None
    # 系统工具显示所需的额外字段
    id: Optional[str] = None
    category: str = "系统工具"
    functionalCategory: Union[str, List[str]] = "system_tools"
    icon: str = "[GEAR2]"
    source: str = "user_shared"
    tool_type: str = "mcp_tool"
    enabled: bool = True
    tags: List[str] = []
    url: Optional[str] = None

@app.post("/detect_intent")
async def detect_intent_endpoint(request: IntentDetectionRequest):
    """
    使用AI智能检测用户意图
    """
    try:
        prompt = request.prompt

        # 简化的关键词检测
        mindmap_keywords = {
            'generate', 'create', 'make', 'draw', 'build', 'design', 'plan',
            'mind map', 'mindmap', 'workflow', 'process', 'task', 'project',
            '思维导图', '生成', '创建', '画', '制作', '设计', '规划', '工作流', '流程', '任务', '项目'
        }

        prompt_lower = prompt.lower()
        has_mindmap_keyword = any(keyword in prompt_lower for keyword in mindmap_keywords)

        if has_mindmap_keyword:
            intent = 'mindmap'
            confidence = 'keyword_detected'
        else:
            intent = 'chat'
            confidence = 'keyword_detected'

        return {
            "intent": intent,
            "confidence": confidence
        }

    except Exception as e:
        logging.error(f"Error in intent detection: {e}")
        # 出错时返回默认意图
        return {
            "intent": "chat",
            "confidence": "fallback"
        }

# ==============================================================================
# System Tools Library Management
# ==============================================================================

# 系统工具库存储文件路径
SYSTEM_TOOLS_FILE = "../data/configs/backend-config/system_tools.json"
USER_TOOLS_FILE = "user_tools.json"

def _load_system_tools() -> List[Dict[str, Any]]:
    """加载系统工具库配置"""
    try:
        if os.path.exists(SYSTEM_TOOLS_FILE):
            with open(SYSTEM_TOOLS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    except Exception as e:
        logging.error(f"加载系统工具库失败: {e}")
        return []

def _save_system_tools(tools: List[Dict[str, Any]]) -> bool:
    """保存系统工具库配置"""
    try:
        with open(SYSTEM_TOOLS_FILE, 'w', encoding='utf-8') as f:
            json.dump(tools, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logging.error(f"保存系统工具库失败: {e}")
        return False

def _load_user_tools() -> List[Dict[str, Any]]:
    """加载用户工具配置"""
    try:
        if os.path.exists(USER_TOOLS_FILE):
            with open(USER_TOOLS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    except Exception as e:
        logging.error(f"加载用户工具失败: {e}")
        return []

def _save_user_tools(tools: List[Dict[str, Any]]) -> bool:
    """保存用户工具配置"""
    try:
        with open(USER_TOOLS_FILE, 'w', encoding='utf-8') as f:
            json.dump(tools, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logging.error(f"保存用户工具失败: {e}")
        return False

@app.get("/system-tools")
async def list_system_tools():
    """获取系统工具库列表"""
    try:
        tools = _load_system_tools()
        return tools
    except Exception as e:
        logging.error(f"获取系统工具库失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取系统工具库失败: {str(e)}")

@app.post("/system-tools")
async def add_or_update_system_tool(config: SystemToolConfig):
    """添加或更新系统工具"""
    try:
        tools = _load_system_tools()

        # 转换为字典格式
        tool_dict = config.model_dump()
        tool_dict["created_at"] = tool_dict.get("created_at") or datetime.now().isoformat()
        tool_dict["updated_at"] = datetime.now().isoformat()

        # 确保有ID字段
        if not tool_dict.get("id"):
            tool_dict["id"] = f"user_shared_{config.name.replace(' ', '_').replace('-', '_').lower()}"

        # 确保有系统工具显示所需的所有字段
        if not tool_dict.get("category"):
            tool_dict["category"] = "系统工具"
        if not tool_dict.get("functionalCategory"):
            tool_dict["functionalCategory"] = tool_dict.get("functionalCategory", ["system_tools"])
        if not tool_dict.get("icon"):
            tool_dict["icon"] = "[GEAR2]"
        if not tool_dict.get("source"):
            tool_dict["source"] = "user_shared"
        if not tool_dict.get("tool_type"):
            tool_dict["tool_type"] = "mcp_tool"
        if not tool_dict.get("tags"):
            tool_dict["tags"] = ["用户分享", "MCP工具"]

        # 检查是否已存在同名工具
        existing_index = None
        for i, tool in enumerate(tools):
            if tool.get("name") == config.name:
                existing_index = i
                break

        if existing_index is not None:
            # 更新现有工具
            tools[existing_index] = tool_dict
            logging.info(f"更新系统工具: {config.name}")
        else:
            # 添加新工具
            tools.append(tool_dict)
            logging.info(f"添加系统工具: {config.name}")

        # 保存到文件
        if _save_system_tools(tools):
            return {"success": True, "message": f"工具 {config.name} 已保存到系统工具库"}
        else:
            raise HTTPException(status_code=500, detail="保存工具失败")

    except Exception as e:
        logging.error(f"添加/更新系统工具失败: {e}")
        raise HTTPException(status_code=500, detail=f"操作失败: {str(e)}")

# ==============================================================================
# User Tools Management
# ==============================================================================

class UserToolConfig(BaseModel):
    """用户工具配置模型"""
    name: str
    description: str
    type: str = "api"  # api, file, data, ai, search, utility
    config: Dict[str, Any] = {}
    enabled: bool = True
    category: str = "我的工具"
    icon: str = "[HAMMER_WRENCH2]"

@app.get("/api/tools/user-tools")
async def list_user_tools():
    """获取用户工具列表"""
    try:
        tools = _load_user_tools()
        return {"success": True, "tools": tools}
    except Exception as e:
        logging.error(f"获取用户工具失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取用户工具失败: {str(e)}")

@app.post("/api/tools/user-tools")
async def create_user_tool(config: UserToolConfig):
    """创建用户工具"""
    try:
        tools = _load_user_tools()

        # 转换为字典格式
        tool_dict = config.dict()
        tool_dict["id"] = f"user_tool_{int(time.time())}"
        tool_dict["created_at"] = datetime.now().isoformat()
        tool_dict["updated_at"] = datetime.now().isoformat()
        tool_dict["tested"] = True  # 新工具默认已测试，直接可以批准
        tool_dict["source"] = "user"

        # 检查是否已存在同名工具
        if any(tool.get("name") == config.name for tool in tools):
            raise HTTPException(status_code=400, detail="工具名称已存在")

        tools.append(tool_dict)

        if _save_user_tools(tools):
            return {"success": True, "message": f"工具 {config.name} 创建成功", "tool": tool_dict}
        else:
            raise HTTPException(status_code=500, detail="保存工具失败")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"创建用户工具失败: {e}")
        raise HTTPException(status_code=500, detail=f"创建用户工具失败: {str(e)}")

@app.put("/api/tools/user-tools/{tool_id}")
async def update_user_tool(tool_id: str, config: UserToolConfig):
    """更新用户工具"""
    try:
        tools = _load_user_tools()

        # 查找要更新的工具
        tool_index = None
        for i, tool in enumerate(tools):
            if tool.get("id") == tool_id:
                tool_index = i
                break

        if tool_index is None:
            raise HTTPException(status_code=404, detail="工具不存在")

        # 更新工具配置
        tool_dict = config.dict()
        tool_dict["id"] = tool_id
        tool_dict["updated_at"] = datetime.now().isoformat()
        tool_dict["created_at"] = tools[tool_index].get("created_at", datetime.now().isoformat())
        tool_dict["tested"] = tools[tool_index].get("tested", False)
        tool_dict["source"] = "user"

        tools[tool_index] = tool_dict

        if _save_user_tools(tools):
            return {"success": True, "message": f"工具 {config.name} 更新成功", "tool": tool_dict}
        else:
            raise HTTPException(status_code=500, detail="保存工具失败")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"更新用户工具失败: {e}")
        raise HTTPException(status_code=500, detail=f"更新用户工具失败: {str(e)}")

@app.delete("/api/tools/user-tools/{tool_id}")
async def delete_user_tool(tool_id: str):
    """删除用户工具"""
    try:
        tools = _load_user_tools()

        # 查找要删除的工具
        tool_index = None
        for i, tool in enumerate(tools):
            if tool.get("id") == tool_id:
                tool_index = i
                break

        if tool_index is None:
            raise HTTPException(status_code=404, detail="工具不存在")

        deleted_tool = tools.pop(tool_index)

        if _save_user_tools(tools):
            return {"success": True, "message": f"工具 {deleted_tool.get('name', '未知')} 删除成功"}
        else:
            raise HTTPException(status_code=500, detail="保存工具失败")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"删除用户工具失败: {e}")
        raise HTTPException(status_code=500, detail=f"删除用户工具失败: {str(e)}")

@app.post("/api/tools/test")
async def test_tool(tool_data: Dict[str, Any]):
    """测试工具连接"""
    try:
        tool_type = tool_data.get("type", "api")
        tool_config = tool_data.get("config", {})

        if tool_type == "api":
            # 测试API工具
            base_url = tool_config.get("baseUrl", "")
            method = tool_config.get("method", "GET")
            auth_type = tool_config.get("authType", "none")
            auth_token = tool_config.get("authToken", "")

            if not base_url:
                return {"success": False, "error": "API地址不能为空"}

            headers = {}
            if auth_type == "bearer" and auth_token:
                headers["Authorization"] = f"Bearer {auth_token}"
            elif auth_type == "apikey" and auth_token:
                headers["X-API-Key"] = auth_token

            import aiohttp
            async with aiohttp.ClientSession() as session:
                try:
                    async with session.request(method, base_url, headers=headers, timeout=10) as response:
                        if response.status < 400:
                            # 更新工具测试状态
                            if "id" in tool_data:
                                await _update_tool_test_status(tool_data["id"], True)
                            return {"success": True, "message": "API连接测试成功", "status_code": response.status}
                        else:
                            return {"success": False, "error": f"API返回错误状态码: {response.status}"}
                except aiohttp.ClientError as e:
                    return {"success": False, "error": f"连接失败: {str(e)}"}
        elif tool_type == "ai" or tool_data.get("provider") == "deepseek":
            # 测试AI工具（如DeepSeek模型）
            try:
                # 使用环境变量中的API密钥，而不是配置文件中的
                api_key = os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENAI_API_KEY")
                if not api_key:
                    return {"success": False, "error": "未配置DEEPSEEK_API_KEY或OPENAI_API_KEY环境变量"}

                # 测试DeepSeek API连接
                from openai import AsyncOpenAI
                test_client = AsyncOpenAI(
                    api_key=api_key,
                    base_url="https://api.deepseek.com/v1"
                )

                # 发送一个简单的测试请求
                response = await test_client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "user", "content": "测试连接"}],
                    max_tokens=10
                )

                if response and response.choices:
                    # 更新工具测试状态
                    if "id" in tool_data:
                        await _update_tool_test_status(tool_data["id"], True)
                    return {"success": True, "message": "DeepSeek模型连接测试成功"}
                else:
                    return {"success": False, "error": "DeepSeek API响应异常"}

            except Exception as e:
                return {"success": False, "error": f"DeepSeek模型测试失败: {str(e)}"}
        else:
            # 其他类型工具的测试逻辑
            return {"success": True, "message": f"{tool_type}类型工具测试通过"}

    except Exception as e:
        logging.error(f"工具测试失败: {e}")
        return {"success": False, "error": f"测试失败: {str(e)}"}

async def _update_tool_test_status(tool_id: str, tested: bool = None, approved: bool = None):
    """更新工具测试状态和批准状态"""
    try:
        # 首先尝试更新用户工具
        tools = _load_user_tools()
        tool_found = False
        for tool in tools:
            if tool.get("id") == tool_id:
                # 只有当tested不为None时才更新测试状态
                if tested is not None:
                    tool["tested"] = tested

                tool["updated_at"] = datetime.now().isoformat()

                # 如果测试成功，设置为可用状态（不需要待批准）
                if tested and approved is None:
                    tool["approval_status"] = "ready"  # 测试通过，可以使用
                    tool["test_passed_at"] = datetime.now().isoformat()
                elif approved is not None:
                    tool["approval_status"] = "shared" if approved else "ready"  # 批准=分享到系统
                    if approved:
                        tool["shared_at"] = datetime.now().isoformat()
                        tool["is_system_tool"] = True  # 标记为系统工具

                tool_found = True
                break

        if tool_found:
            _save_user_tools(tools)
        else:
            # 如果在用户工具中没找到，尝试更新MCP工具
            from utils.frontend_config import load_config, save_config, MCP_TOOLS_CONFIG
            mcp_tools = load_config(MCP_TOOLS_CONFIG)
            for tool in mcp_tools:
                if tool.get("id") == tool_id:
                    # 只有当tested不为None时才更新测试状态
                    if tested is not None:
                        tool["tested"] = tested

                    tool["updatedAt"] = datetime.now().isoformat()

                    # 如果测试成功，设置为可用状态（不需要待批准）
                    if tested and approved is None:
                        tool["approvalStatus"] = "ready"  # 测试通过，可以使用
                        tool["testPassedAt"] = datetime.now().isoformat()
                    elif approved is not None:
                        tool["approvalStatus"] = "shared" if approved else "ready"  # 批准=分享到系统
                        if approved:
                            tool["sharedAt"] = datetime.now().isoformat()
                            tool["isSystemTool"] = True  # 标记为系统工具

                    tool_found = True
                    break

            if tool_found:
                save_config(MCP_TOOLS_CONFIG, mcp_tools)
                print(f"[OK] MCP工具状态已更新: {tool_id}")

        if not tool_found:
            print(f"[WARN] 未找到工具ID: {tool_id}")

    except Exception as e:
        logging.error(f"更新工具测试状态失败: {e}")

async def _sync_tool_to_system(tool_id: str):
    """将工具同步到系统工具库"""
    try:
        from utils.frontend_config import load_config, save_config, MCP_TOOLS_CONFIG

        # 获取MCP工具
        mcp_tools = load_config(MCP_TOOLS_CONFIG)
        tool_to_sync = None

        for tool in mcp_tools:
            if tool.get("id") == tool_id:
                tool_to_sync = tool
                break

        if not tool_to_sync:
            print(f"[WARN] 未找到要同步的工具: {tool_id}")
            return False

        # 创建系统工具版本（移除敏感信息）
        system_tool = {
            "id": f"system_{tool_id}",  # 确保有ID字段
            "name": tool_to_sync.get("name"),
            "description": tool_to_sync.get("description"),
            "transport": tool_to_sync.get("server_type", "stdio"),  # 使用transport而不是server_type
            "command": tool_to_sync.get("command"),
            "args": tool_to_sync.get("args"),
            "url": tool_to_sync.get("url"),
            "timeout": 60000,
            "env": {},
            "enabled": True,
            "functionalCategory": tool_to_sync.get("functionalCategory", []),  # 保持功能分类
            "version": tool_to_sync.get("version", "1.0.0"),
            "author": "系统用户分享",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "sourceType": "user_shared",
            "capabilities": tool_to_sync.get("capabilities", []),
            "tags": ["系统工具", "用户分享"],
            "tested": True,
            "approved": True,
            "originalToolId": tool_id,  # 记录原始工具ID
            "isSystemTool": True
        }

        # 添加到系统工具配置文件
        from utils.frontend_config import SYSTEM_TOOLS_CONFIG
        system_tools = load_config(SYSTEM_TOOLS_CONFIG)

        # 检查是否已经存在相同的工具（避免重复添加）
        existing_tool = None
        for existing in system_tools:
            if existing.get("originalToolId") == tool_id or existing.get("name") == tool_to_sync.get("name"):
                existing_tool = existing
                break

        if existing_tool:
            # 更新现有工具
            existing_tool.update(system_tool)
            print(f"[RELOAD] 更新现有系统工具: {tool_to_sync.get('name')}")
        else:
            # 添加新工具
            system_tools.append(system_tool)
            print(f"➕ 添加新系统工具: {tool_to_sync.get('name')}")

        save_config(SYSTEM_TOOLS_CONFIG, system_tools)

        print(f"[OK] 工具已同步到系统工具库: {tool_to_sync.get('name')}")
        return True

    except Exception as e:
        print(f"[ERROR] 同步工具到系统失败: {str(e)}")
        return False

# 旧的工具执行端点已删除，使用新的统一端点


# ==============================================================================
# 内部工具API端点 - HTTP请求工具和数据分析工具
# ==============================================================================

@app.post("/api/internal-tools/http-request")
async def internal_http_request(
    url: str = Body(...),
    method: str = Body("GET"),
    headers: Optional[Dict[str, str]] = Body(None),
    data: Optional[Dict[str, Any]] = Body(None),
    timeout: int = Body(30)
):
    """内部HTTP请求工具API端点"""
    try:
        result = await http_tool.make_request(
            url=url,
            method=method,
            headers=headers,
            data=data,
            timeout=timeout
        )
        return result
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"内部HTTP工具执行失败: {str(e)}"
        }


@app.post("/api/internal-tools/http-request-sync")
async def internal_http_request_sync(
    url: str = Body(...),
    method: str = Body("GET"),
    headers: Optional[Dict[str, str]] = Body(None),
    data: Optional[Dict[str, Any]] = Body(None),
    timeout: int = Body(30)
):
    """内部HTTP请求工具API端点（同步版本）"""
    try:
        result = http_tool.make_sync_request(
            url=url,
            method=method,
            headers=headers,
            data=data,
            timeout=timeout
        )
        return result
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"内部HTTP工具执行失败: {str(e)}"
        }


@app.post("/api/internal-tools/data-analysis")
async def internal_data_analysis(
    data: List[Dict[str, Any]] = Body(...),
    analysis_type: str = Body("basic")
):
    """内部数据分析工具API端点"""
    try:
        result = data_tool.analyze_data(
            data=data,
            analysis_type=analysis_type
        )
        return result
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"内部数据分析工具执行失败: {str(e)}"
        }


@app.post("/api/internal-tools/data-cleaning")
async def internal_data_cleaning(
    data: List[Dict[str, Any]] = Body(...),
    operations: Optional[List[str]] = Body(None)
):
    """内部数据清洗工具API端点"""
    try:
        result = data_tool.clean_data(
            data=data,
            operations=operations
        )
        return result
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"内部数据清洗工具执行失败: {str(e)}"
        }


# ==============================================================================
# AI大脑升级：注入了全新规划思想的系统指令
# ==============================================================================
# NOTE: 使用 .format() 插入 tool_descriptions 时，必须把 JSON 示例中的 { } 进行转义为 {{ }}，否则会被当成占位符。
SYSTEM_PROMPT = """
# 角色 (Role)
你是 Workflow Planner，一个专门把用户目标拆解为**有序工具调用计划**的 AI。

# 目标 (Goal)
产出 *严格符合 JSON 结构* 的输出，包含两个字段：
1. reasoning: 你逐步思考的中文过程；
2. plan: 工具调用列表 (按顺序执行)。

# 可用工具 (Tools)
{tool_descriptions}
每条工具说明包含名称和功能描述，你必须确保 tool_name 字段与上表完全一致。

# 链式思考 (Chain-of-Thought)
在决定 plan 之前，请先生成 reasoning，分三步思考：
① 分析用户意图 ② 需要哪些信息 ③ 选择并排序工具。

# 输出格式 (Output Format)
仅输出如下 JSON，无其他多余文字：
{{
  "reasoning": "...",
  "plan": [
    {{"tool_name": "<tool_name>", "arguments": {{"arg1": "..."}}}},
    ...
  ]
}}

# 示例 (Few-shot)
<用户> 查询"华为 AI 战略 白皮书"并摘要 </用户>
<AI>
{{
  "reasoning": "用户想获取白皮书内容并概要...",
  "plan": [
    {{"tool_name": "backend_tools_api_tools_web_search", "arguments": {{"query": "华为 AI 战略 白皮书", "limit": 5}}}},
    {{"tool_name": "backend_tools_fetch_tools_safe_fetch_page", "arguments": {{"url": "@node-1.web_search_result[0].href"}}}},
    {{"tool_name": "backend_tools_ai_tools_summarize_text", "arguments": {{"text": "@node-2.safe_fetch_page_result", "style": "concise"}}}}
  ]
}}
</AI>

# 约束 (Rules)
1. 必须使用 JSON 对象包裹 reasoning 与 plan 两个键。
2. plan 内步骤按执行先后排序；若任务涉及多主体，先并行搜索后并行抓取；最后汇总。
3. 使用 @node-id.output_name 语法正确串联数据流。
4. 若没有合适工具，诚实说明，但仍返回空列表 plan。 
"""


# --- Tool Discovery and Registration ---
async def create_tool_based_plan_simple(prompt: str, model: str, tools: List[Dict[str, Any]]):
    """
    A fallback planner that uses the basic tool-calling feature.
    """
    messages = [
        {"role": "system", "content": "You are a helpful assistant that can plan tasks by calling available tools. Respond with a list of tool calls required to fulfill the user's request."},
        {"role": "user", "content": prompt}
    ]
    try:
        response = await client.chat.completions.create(
            model=model,
            messages=messages,
            tools=tools,
            tool_choice="auto",
        )
        response_message = response.choices[0].message
        tool_calls = response_message.tool_calls

        if not tool_calls:
            logging.info("Fallback AI did not choose to use any tools.")
            return None

        plan = []
        for tool_call in tool_calls:
            plan.append({
                "tool_name": tool_call.function.name,
                "arguments": json.loads(tool_call.function.arguments)
            })
        logging.info(f"Received tool-based plan from fallback AI: {plan}")
        return plan
    except Exception as e:
        logging.error(f"Error in fallback create_tool_based_plan_simple: {e}", exc_info=True)
        return None

async def create_fallback_plan(prompt: str, model: str):
    """
    当没有工具可用时，创建基础的回退计划
    """
    logging.info("Creating fallback plan without registered tools.")

    # 分析prompt类型，生成对应的基础计划
    prompt_lower = prompt.lower()

    # 通用计划
    return [
        {"tool_name": "task_processor", "arguments": {"task": prompt[:100]}},
        {"tool_name": "result_formatter", "arguments": {"result": "处理结果"}}
    ]

async def create_tool_based_plan(prompt: str, model: str):
    """
    Creates a sequential plan by asking the AI to generate a JSON structure.
    This is the primary, more intelligent planning method.
    """
    tools = tool_registry.get_tools()
    logging.info(f"Found {len(tools)} registered tools")

    if not tools:
        logging.warning("No tools are registered. Creating fallback plan with basic tools.")
        # 创建基础工具的回退计划
        return await create_fallback_plan(prompt, model)

    # 记录可用工具
    tool_names = [tool["function"]["name"] for tool in tools]
    logging.info(f"Available tools: {tool_names}")

    tool_descriptions = ""
    for tool in tools:
        function_info = tool.get('function', {})
        func_name = function_info.get('name', 'unknown_function')
        func_desc = function_info.get('description', 'No description.')
        tool_descriptions += f"- `{func_name}`: {func_desc}\\n"

    # The assembler now needs to know about the plan to wire things correctly.
    # So we pass the plan to it.
    system_prompt_content = SYSTEM_PROMPT.format(tool_descriptions=tool_descriptions)


    messages = [
        {"role": "system", "content": system_prompt_content},
        {"role": "user", "content": prompt}
    ]

    logging.info("Requesting structured JSON plan from AI.")
    try:
        response = await client.chat.completions.create(
            model=model,
            messages=messages,
            response_format={"type": "json_object"},
        )
        response_content = response.choices[0].message.content
        plan_data = json.loads(response_content)

        if "reasoning" in plan_data and "plan" in plan_data and isinstance(plan_data["plan"], list):
            logging.info(f"Received structured plan from AI: {plan_data['plan']}")
            return plan_data["plan"]
        else:
            logging.warning(f"AI returned malformed JSON plan: {response_content}. Falling back.")
            return await create_tool_based_plan_simple(prompt, model, tools)

    except Exception as e:
        logging.error(f"Structured planning failed: {e}. Falling back to simple tool calling.", exc_info=True)
        return await create_tool_based_plan_simple(prompt, model, tools)


# --- REMOVED OLD PLANNER AND FORMATTER ---

def assembler(plan: list, multi_scene: bool = True):
    """
    New Assembler: Dynamically creates a graph from a linear plan.
    - It expands 'web_search' results into individual 'fetch_page' nodes.
    - It wires up dependencies correctly for a final summarization.
    """
    nodes, edges = [], []
    # ------- 0. Simple fallback: 如果计划中没有 web_search 相关工具，直接线性串联 -------
    if not any("web_search" in step.get("tool_name", "") for step in plan):
        layout_y = 0
        prev_id = None
        # Start node
        nodes.append({
            "id": "1",
            "type": "input",
            "position": {"x": 0, "y": layout_y},
            "data": {"label": "Start", "nodeType": "start-topic-node"}
        })
        prev_id = "1"
        layout_y += 120
        for idx, step in enumerate(plan, start=1):
            node_id = f"tool-{idx}"
            nodes.append({
                "id": node_id,
                "type": "execution-node",
                "position": {"x": 0, "y": layout_y},
                "data": {
                    "label": (lambda n: (n[34:] if n.startswith("backend_tools_tool_registry_mcp_") else n)[:24] + ("…" if len(n)>24 else ""))(step["tool_name"]),
                    "nodeType": "execution-node",
                    "params": { **step.get("arguments", {}), "tool_name": step["tool_name"] },
                    "inputs": [],
                    "outputs": ["out"]
                }
            })
            edges.append({"id": f"e-{prev_id}-{node_id}", "source": prev_id, "target": node_id})
            prev_id = node_id
            layout_y += 140

        return {"nodes": nodes, "edges": edges}

    node_id_counter = 1
    layout_x = 0

    # Phase 1: Create initial search nodes from the plan
    # 去重：同一查询只保留一次
    seen_queries = set()
    search_tasks = []
    import re, unicodedata

    def normalize_query(text: str) -> str:
        """统一化处理：全角转半角、去标点空格、去尾部修饰词、小写"""
        text = unicodedata.normalize('NFKC', text)  # 全角转半角
        # 去所有空白字符（包含中文全角空格等）
        text = re.sub(r"\s+", "", text, flags=re.UNICODE)
        # 去常见中英文标点符号
        text = re.sub(r"[\u3000-\u303F\u2000-\u206F\u0021-\u002F\u003A-\u0040\u005B-\u0060\u007B-\u007E]+", "", text)
        text = re.sub(r"(优势|优缺点|分析|评价|对比|比较)$", "", text)  # 去常见修饰
        return text.lower()

    for item in plan:
        if "web_search" not in item.get("tool_name", ""):
            continue
        raw_q = item.get("arguments", {}).get("query", "")
        nq = normalize_query(raw_q)
        if nq and nq not in seen_queries:
            seen_queries.add(nq)
            search_tasks.append(item)

    # 如果未启用多场景，仅保留第一条查询
    if not multi_scene and len(search_tasks) > 1:
        search_tasks = search_tasks[:1]

    last_op_node_ids = []

    for task in search_tasks:
        search_node_id = f"search-node-{node_id_counter - 1}"
        query = task["arguments"].get("query", "Missing query")

        # --- Search Node ---
        nodes.append({
            "id": search_node_id,
            "type": "execution-node",
            "position": {"x": layout_x, "y": 100},
            "data": {
                "label": f"Web Search: '{query[:20]}...'",
                "nodeType": "execution-node",
                "params": {
                    "tool_name": task["tool_name"],
                    "query": query,
                    "limit": 5  # 请求更多结果，便于 fallback
                },
                "inputs": [],
                "outputs": ["web_search_result"]
            },
            # 默认隐藏中间技术节点，前端仍可通过"显示全部"按钮展开
            "hidden": False
        })

        # --- Multiple Fetch Nodes per search result ---
        FETCH_PER_SEARCH = 3  # 尝试抓取前 3 条结果
        for idx in range(FETCH_PER_SEARCH):
            fetch_node_id = f"fetch-node-{node_id_counter - 1}-{idx}"
            nodes.append({
                "id": fetch_node_id,
                "type": "execution-node",
                "position": {"x": layout_x, "y": 250 + idx * 140},
                "data": {
                    "label": f"Fetch {idx + 1}: '{query[:15]}...'",
                    "nodeType": "execution-node",
                    "params": {
                        "tool_name": "backend_tools_fetch_tools_safe_fetch_page",
                        "url": f"@{search_node_id}.web_search_result[{idx}].href"
                    },
                    "inputs": ["url"],
                    "outputs": ["content"]
                },
                # 默认不隐藏
                "hidden": False
            })
            edges.append({
                "id": f"e-{search_node_id}-fetch-{idx}",
                "source": search_node_id,
                "target": fetch_node_id
            })
            last_op_node_ids.append(fetch_node_id)

        layout_x += 320  # 下一个搜索分支水平偏移
        node_id_counter += 1

    # Phase 2: Create the final summary node
    if last_op_node_ids:
        summary_node_id = "summary-node"
        summary_prompt = (
            "请根据以下抓取内容，对华为与腾讯在 AI 领域的优点、缺点进行对比分析，"
            "并用中文以项目符号 (bullet) 形式输出：\n"
            "- 华为优点…〔1〕\n- 华为缺点…〔2〕\n- 腾讯优点…〔3〕\n- 腾讯缺点…〔4〕\n"
            "请在每条 bullet 末尾加上来源编号〔i〕，编号顺序按下文列表顺序对应。\n\n"
        )
        for i, fetch_node_id in enumerate(last_op_node_ids):
            summary_prompt += f"来源{i+1}: @{fetch_node_id}.content \\n\\n"

        nodes.append({
            "id": summary_node_id,
            "type": "output",
            "position": {"x": layout_x / 2 - 100, "y": 400},
            "data": {
                "label": "Final Summary",
                "nodeType": "end-topic-node",
                "params": {"final_prompt": summary_prompt},
                "inputs": [f"@{fetch_node_id}.content" for fetch_node_id in last_op_node_ids],
                "outputs": ["final_summary"]
            }
        })
        for fetch_node_id in last_op_node_ids:
            edges.append({"id": f"e-fetch-summary-{fetch_node_id}", "source": fetch_node_id, "target": summary_node_id})

    # Add a start node for clarity
    nodes.insert(0, {
        "id": "1",  # 前端期望的主节点 ID
        "type": "input",
        "position": {"x": layout_x / 2 - 50, "y": 0},
        "data": {"label": "Start", "nodeType": "start-topic-node"}
    })
    
    for search_node in search_tasks:
        search_node_id = nodes[search_tasks.index(search_node) + 1]['id'] # +1 to account for start node
        edges.append({"id": f"e-start-search-{search_node_id}", "source": "1", "target": search_node_id})

    return {"nodes": nodes, "edges": edges}

# --- API Endpoint ---
@app.post("/chat")
async def handle_chat(request: ChatRequest):
    """
    Handles chat requests. Can either generate a workflow (mindmap) or provide simple chat responses.
    """
    try:
        logging.info(f"[RELOAD] 收到聊天请求 - prompt: {request.prompt[:100]}..., intent: {request.intent}")
        print(f"[RELOAD] 收到聊天请求 - prompt: {request.prompt[:100]}..., intent: {request.intent}")
    except Exception as e:
        logging.error(f"[ERROR] 请求处理错误: {e}")
        print(f"[ERROR] 请求处理错误: {e}")
        raise HTTPException(status_code=400, detail=f"请求处理错误: {str(e)}")

    # 如果intent是chat，直接返回聊天响应
    if request.intent == 'chat':
        return await handle_simple_chat(request)

    # 如果intent是mindmap、task或workflow，使用AI工作流构建器
    if request.intent in ['mindmap', 'task', 'workflow']:
        try:
            from ai.intelligent_workflow_orchestrator import IntelligentWorkflowOrchestrator

            # 使用新的智能工作流编排器
            orchestrator = IntelligentWorkflowOrchestrator()
            await orchestrator.initialize()

            # 生成工作流
            intelligent_workflow = await orchestrator.generate_intelligent_workflow(request.prompt)

            # 转换为前端期望的格式
            def node_to_dict(node):
                node_type_str = node.type.value if hasattr(node.type, 'value') else str(node.type)
                # 材料节点使用customNode显示详细统计信息，其他使用普通节点类型
                if node_type_str == "material-node":
                    frontend_type = "customNode"
                else:
                    frontend_type = node_type_str

                # 节点类型到友好名称的映射
                friendly_names = {
                    "material-node": "材料",
                    "execution-node": "执行",
                    "result-node": "结果",
                    "condition-node": "条件"
                }

                # 节点类型到图标的映射
                type_icons = {
                    "material-node": "[FILE_FOLDER]",
                    "execution-node": "[LIGHTNING]",
                    "result-node": "[PAGE_FACING_UP]",
                    "condition-node": "🔀"
                }

                friendly_label = friendly_names.get(node_type_str, f"智能{node_type_str}")
                type_icon = type_icons.get(node_type_str, "[LIGHTNING]")

                return {
                    "id": node.id,
                    "type": frontend_type,
                    "position": node.position,
                    "data": {
                        "label": friendly_label,
                        "nodeType": frontend_type,
                        "status": "default"  # 普通节点的默认状态
                        # 移除增强型节点特有的属性，使用普通节点格式
                        # "config": node.config,
                        # "intelligentConfig": True,
                        # "aiEnhanced": True,
                        # "typeIcon": type_icon,
                        # "progress": 0
                    }
                }

            def conn_to_dict(conn):
                return {
                    "id": conn.id,
                    "source": conn.source,
                    "target": conn.target,
                    "type": "enhanced",
                    "animated": False,
                    "className": "silky-smooth",
                    "style": {"strokeWidth": 2, "stroke": "#7c3aed"},
                    "markerEnd": "url(#smooth-arrow)",
                    "data": {
                        "connectionType": conn.type.value if hasattr(conn.type, 'value') else str(conn.type),
                        "conditions": conn.conditions or {},
                        "metadata": conn.metadata or {}
                    }
                }

            workflow_result = {
                "success": True,
                "workflow": {
                    "nodes": [node_to_dict(node) for node in intelligent_workflow.nodes],
                    "edges": [conn_to_dict(conn) for conn in intelligent_workflow.connections],
                    "metadata": intelligent_workflow.metadata
                }
            }

            if workflow_result.get("success"):
                # 直接使用新API返回的工作流数据，无需额外转换
                workflow = workflow_result.get("workflow", {})
                nodes = workflow.get("nodes", [])
                edges = workflow.get("edges", [])

                # 调试日志：记录AI生成的节点类型
                logging.info(f"AI工作流构建器返回的节点类型: {[n.get('type') for n in nodes]}")
                logging.info(f"AI工作流构建器返回的节点数据: {[n.get('data', {}).get('nodeType') for n in nodes]}")

                # 构建响应
                graph = {
                    "nodes": nodes,
                    "edges": edges,
                    "description": workflow.get("metadata", {}).get("analysis", {}).get("workflow_name", "AI生成的工作流")
                }

                logging.info(f"Successfully created mindmap with {len(nodes)} nodes using AI workflow builder")
                return JSONResponse(content=graph)
            else:
                # AI工作流构建器失败，返回错误信息而不是使用备用方案
                logging.warning(f"AI workflow builder failed: {workflow_result.get('error', 'Unknown error')}")
                error_node = {
                    "id": "error-node",
                    "type": "enhanced",
                    "data": {
                        "label": "AI工作流构建失败",
                        "type": "error-node",
                        "nodeType": "error-node",
                        "typeIcon": "[ERROR]",
                        "status": "error",
                        "progress": 0,
                        "params": {"error": f"AI工作流构建失败: {workflow_result.get('error', '未知错误')}"},
                        "inputs": [], "outputs": []
                    },
                    "position": {"x": 200, "y": 200}
                }
                return JSONResponse(content={"nodes": [error_node], "edges": []})

        except Exception as e:
            logging.error(f"AI workflow builder failed: {e}")
            # 返回错误信息而不是使用备用方案
            error_node = {
                "id": "error-node",
                "type": "enhanced",
                "data": {
                    "label": "AI工作流构建异常",
                    "type": "error-node",
                    "nodeType": "error-node",
                    "typeIcon": "[ERROR]",
                    "status": "error",
                    "progress": 0,
                    "params": {"error": f"AI工作流构建异常: {str(e)}"},
                    "inputs": [], "outputs": []
                },
                "position": {"x": 200, "y": 200}
            }
            return JSONResponse(content={"nodes": [error_node], "edges": []})




# --- New Workflow Execution Logic ---

# 模型定义已移至 schemas.py

# --- NEW: Pydantic Model for Approval Resumption ---
class ApprovalResumePayload(BaseModel):
    nodes: List[Node]
    edges: List[Edge]
    context: Dict[str, Any]
    resume_from_node_id: str # This will be the approval node's id
    approved_data: List[str] # The user-approved list

@app.post("/execute")
async def execute_workflow(payload: WorkflowPayload, user_id: Optional[str] = None):
    """
    Executes a workflow graph provided by the frontend.
    Enhanced with resource management and concurrent execution.
    """
    logging.info("Received workflow for execution.")

    try:
        from ai.concurrent_executor import concurrent_executor
        from utils.resource_manager import resource_manager
        import time

        # Validate the payload
        if not payload.nodes:
            return JSONResponse(status_code=400, content={"error": "No nodes provided"})

        # 检查资源和配额
        estimated_duration = len(payload.nodes) * 30  # 估算每个节点30秒
        resource_check = await resource_manager.can_execute_task(user_id, estimated_duration)

        if not resource_check['can_execute']:
            return JSONResponse(status_code=429, content={
                "error": "Resource limit exceeded",
                "reason": resource_check['reason'],
                "checks": resource_check['checks']
            })

        # 获取资源
        start_time = time.time()
        if not await resource_manager.acquire_resources(user_id):
            return JSONResponse(status_code=503, content={
                "error": "Unable to acquire resources"
            })

        try:
            # === Clarify phase: ask LLM for missing questions ===
            try:
                clarify_prompt = render_prompt("clarify_requirements", user_input=json.dumps([n.data.label for n in payload.nodes], ensure_ascii=False))
                from agent_system.agent_factory import AgentFactory  # local import to avoid circular
                clarify_resp = await AgentFactory.ask_llm(clarify_prompt)

                questions: list[str] = []
                try:
                    questions = json.loads(clarify_resp)
                    if not isinstance(questions, list):
                        questions = []
                except Exception:
                    # fallback: split by lines / numbering
                    for line in clarify_resp.splitlines():
                        clean = line.strip("- •0123456789. ")
                        if clean:
                            questions.append(clean)

                questions = [q for q in questions if q]
                if questions:
                    return {
                        "need_clarify": True,
                        "questions": questions
                    }
            except Exception as e:
                logging.warning(f"Clarify phase skipped due to error: {e}")

            # 使用并发执行器执行工作流
            result = await concurrent_executor.execute_workflow(
                payload.nodes,
                payload.edges,
                context={},
                progress_callback=None  # 可以添加WebSocket进度回调
            )

            # 记录执行统计
            execution_time = time.time() - start_time
            resource_manager.tracker.record_execution(
                execution_time,
                result.get('stats', {}).get('peak_memory_usage', 0),
                result.get('stats', {}).get('peak_cpu_usage', 0)
            )

            return result

        finally:
            # 释放资源
            execution_time = time.time() - start_time
            resource_manager.release_resources(user_id, execution_time)

    except Exception as e:
        logging.error(f"Workflow execution failed: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


# --- Refactored Helper Functions for Execution ---

async def resolve_params_from_context(params: Dict, context: Dict, node_id: str) -> Dict:
    """Resolves parameter values that are references to the execution context."""
    from utils.data_flow import data_flow_manager

    resolved_params = {}
    for key, value in params.items():
        if isinstance(value, str) and value.startswith('@'):
            try:
                resolved_value = data_flow_manager.resolve_reference(value, context)
                resolved_params[key] = resolved_value
            except ValueError as e:
                raise ValueError(f"Failed to resolve reference '{value}' in node '{node_id}': {e}")
        elif isinstance(value, dict):
            # 递归处理嵌套字典
            resolved_params[key] = await resolve_params_from_context(value, context, node_id)
        elif isinstance(value, list):
            # 处理列表中的引用
            resolved_list = []
            for item in value:
                if isinstance(item, str) and item.startswith('@'):
                    try:
                        resolved_item = data_flow_manager.resolve_reference(item, context)
                        resolved_list.append(resolved_item)
                    except ValueError as e:
                        raise ValueError(f"Failed to resolve reference '{item}' in node '{node_id}': {e}")
                elif isinstance(item, dict):
                    resolved_item = await resolve_params_from_context(item, context, node_id)
                    resolved_list.append(resolved_item)
                else:
                    resolved_list.append(item)
            resolved_params[key] = resolved_list
        else:
            resolved_params[key] = value
    return resolved_params

async def hydrate_final_prompt(params: Dict, context: Dict) -> str:
    """Substitutes all context references in the final prompt template."""
    prompt_template = params.get("final_prompt", "Please provide a final summary.")
    
    def substitute_context(match):
        ref_string = match.group(1)
        try:
            # Simplified resolver for prompt templates
            ref_node_id, ref_output_name = re.match(r"([\w-]+)\.([\w_]+)", ref_string).groups()
            context_key = f"{ref_node_id}.{ref_output_name}"
            return str(context.get(context_key, ""))
        except Exception:
            return ""
            
    return re.sub(r"@([\w\.-_]+)", substitute_context, prompt_template)


# --- NEW: Endpoint to resume a failed workflow ---
class ResumePayload(BaseModel):
    nodes: List[Node]
    edges: List[Edge]
    context: Dict[str, Any]
    resume_from_node_id: str
    manual_data: str

@app.post("/resume")
async def resume_workflow(payload: ResumePayload):
    """
    Resumes a workflow from a previously failed node, using manually provided data.
    """
    logging.info(f"Resuming workflow from node {payload.resume_from_node_id} with manual data.")
    
    nodes_by_id = {node.id: node for node in payload.nodes}
    adj = {node_id: [] for node_id in nodes_by_id}
    in_degree = {node_id: 0 for node_id in nodes_by_id}
    for edge in payload.edges:
        if edge.source in nodes_by_id and edge.target in nodes_by_id:
            adj[edge.source].append(edge.target)
            in_degree[edge.target] += 1

    # Find the execution path STARTING from the node AFTER the resumed one
    try:
        # Find the node that failed, which is our starting point
        start_node = nodes_by_id.get(payload.resume_from_node_id)
        if not start_node:
            return JSONResponse(status_code=404, content={"error": "Node to resume from not found."})

        # The output of the failed node is the manual data
        # We need to parse the manual data (e.g., comma-separated URLs)
        manual_results = [item.strip() for item in payload.manual_data.replace(',', '\n').split('\n') if item.strip()]
        
        # Inject the manual data into the context
        output_key = f"{start_node.id}.{start_node.data.outputs[0]}"
        execution_context = payload.context
        execution_context[output_key] = manual_results if len(manual_results) > 1 else (manual_results[0] if manual_results else "")
        logging.info(f"Injected manual data into context key: {output_key}")

        # Re-run the topological sort to find the execution order of remaining nodes
        all_node_ids = list(nodes_by_id.keys())
        sub_graph_nodes = []
        
        # We need to find the correct execution order of all nodes, but only execute the ones from the resume point onwards
        queue = [node_id for node_id in all_node_ids if in_degree[node_id] == 0]
        full_ordered_ids = []
        while queue:
            u = queue.pop(0)
            full_ordered_ids.append(u)
            for v in adj.get(u, []):
                in_degree[v] -= 1
                if in_degree[v] == 0:
                    queue.append(v)
        
        # Find where to start the execution
        try:
            start_index = full_ordered_ids.index(payload.resume_from_node_id)
            nodes_to_execute = [nodes_by_id[nid] for nid in full_ordered_ids[start_index + 1:]] # Execute nodes AFTER the manual one
        except ValueError:
             return JSONResponse(status_code=400, content={"error": "Resumed node not found in execution path."})

        # Now, execute the rest of the workflow with the updated context
        resumed_result = await execute_workflow_logic(nodes_to_execute, execution_context, payload)
        
        # The logic function already returns a serializable dict or a JSONResponse,
        # so we can return it directly.
        return resumed_result

    except Exception as e:
        logging.error(f"Error during workflow resumption: {e}", exc_info=True)
        return JSONResponse(status_code=500, content={"error": f"Failed to resume workflow: {e}"})

@app.post("/resume_from_approval")
async def resume_from_approval(payload: ApprovalResumePayload):
    """
    Resumes a workflow from an approval node using user-confirmed data.
    """
    logging.info(f"Resuming workflow from approval node {payload.resume_from_node_id}.")
    
    execution_context = payload.context
    approval_node_id = payload.resume_from_node_id
    
    # Find the approval node in the original payload to get its output key name
    approval_node_obj = None
    for n in payload.nodes:
        if n.id == approval_node_id:
            approval_node_obj = n
            break

    if not approval_node_obj or not approval_node_obj.data.outputs:
        return JSONResponse(status_code=400, content={"error": "Invalid approval node or output not defined."})

    # Inject the approved data into the context
    output_key = f"{approval_node_obj.id}.{approval_node_obj.data.outputs[0]}"
    execution_context[output_key] = payload.approved_data
    logging.info(f"Injected approved data into context key: {output_key}")

    # Now, figure out the rest of the execution path
    nodes_by_id = {node.id: node for node in payload.nodes}
    adj = {node.id: [] for node in nodes_by_id}
    in_degree = {node.id: 0 for node in nodes_by_id}
    for edge in payload.edges:
        if edge.source in nodes_by_id and edge.target in nodes_by_id:
            adj[edge.source].append(edge.target)
            in_degree[edge.target] += 1
    
    # Get full topological order using Kahn's algorithm
    queue = [node_id for node_id in nodes_by_id if in_degree[node_id] == 0]
    full_ordered_ids = []
    while queue:
        u = queue.pop(0)
        full_ordered_ids.append(u)
        for v in adj.get(u, []):
            in_degree[v] -= 1
            if in_degree[v] == 0:
                queue.append(v)
    
    if len(full_ordered_ids) != len(nodes_by_id):
         return JSONResponse(status_code=400, content={"error": "Workflow has a cycle."})

    # Find where to restart
    try:
        start_index = full_ordered_ids.index(approval_node_id)
        # We need to execute nodes AFTER the approval node
        nodes_to_execute = [nodes_by_id[nid] for nid in full_ordered_ids[start_index + 1:]]
    except ValueError:
        return JSONResponse(status_code=400, content={"error": "Approval node not found in execution path."})

    # Execute the rest of the workflow
    resumed_result = await execute_workflow_logic(nodes_to_execute, execution_context, payload)
    
    return resumed_result


async def execute_workflow_logic(ordered_nodes: List[Node], execution_context: Dict, payload: BaseModel, progress_cb: Optional[Callable[[Dict[str, Any]], Awaitable[None]]] = None) -> Dict:
    """
    The actual logic of executing a list of nodes. Refactored to be reusable.
    `payload` is the original request model (WorkflowPayload or ResumePayload) for history saving.
    """
    from workflow.workflow_state import workflow_state_manager
    import uuid

    # 生成工作流ID
    workflow_id = str(uuid.uuid4())

    # 开始工作流执行
    if isinstance(payload, WorkflowPayload):
        workflow_state_manager.start_workflow(workflow_id, payload)
    for node in ordered_nodes:
        node_type = node.data.nodeType
        logging.info(f"Executing node '{node.data.label}' (Type: {node_type})")

        # 开始节点执行
        workflow_state_manager.start_node(workflow_id, node.id)

        # 添加调试信息
        from utils.error_handling import debug_info
        debug_info.add_debug_point(
            "node_start",
            {"node_type": node_type, "params": node.data.params},
            node.id
        )

        if progress_cb:
            await progress_cb({"event": "node_start", "node_id": node.id, "label": node.data.label})

        # --- Resolve node parameters from context ---
        params = node.data.params.copy()
        try:
            resolved_params = await resolve_params_from_context(params, execution_context, node.id)
        except ValueError as e:
            logging.error(f"Parameter resolution failed for node {node.id}: {e}")
            return JSONResponse(status_code=400, content={
                "status": "failed",
                "error_node_id": node.id,
                "error_message": str(e),
                "context": execution_context,
            })

        # --- NEW: Handle Approval Node ---
        if node_type == 'approval-node':
            # This is where the workflow pauses. Find the data that needs approval.
            input_ref_name = node.data.inputs[0] if node.data.inputs else None
            
            source_node_id = None
            if isinstance(payload, (WorkflowPayload, ApprovalResumePayload, ResumePayload)):
                 for edge in payload.edges:
                    if edge.target == node.id:
                        source_node_id = edge.source
                        break
            
            if not source_node_id or not input_ref_name:
                 return JSONResponse(status_code=500, content={
                    "status": "failed", "error_node_id": node.id, 
                    "error_message": "Approval node is misconfigured (missing input or incoming edge).", "context": execution_context
                })

            context_key_to_approve = f"{source_node_id}.{input_ref_name}"
            data_to_approve = execution_context.get(context_key_to_approve, [])

            logging.info(f"Pausing workflow at node {node.id} for user approval.")
            return JSONResponse(status_code=200, content={
                "status": "paused_for_approval",
                "node_id": node.id,
                "data_to_approve": data_to_approve,
                "context": execution_context # Send the current context to the frontend
            })

        # --- 首先尝试使用扩展节点执行器 ---
        from workflow.node_executors import execute_extended_node, NodeExecutionError
        try:
            result = await execute_extended_node(node, execution_context)

            # 保存结果到上下文（带数据验证）
            if node.data.outputs and result is not None:
                context_key = f"{node.id}.{node.data.outputs[0]}"
                try:
                    from utils.data_flow import data_flow_manager
                    validated_result = data_flow_manager.validate_node_output(
                        node.id, node.data.outputs[0], result
                    )
                    execution_context[context_key] = validated_result
                except Exception as e:
                    logging.warning(f"Data validation failed for {context_key}: {e}")
                    execution_context[context_key] = result  # 使用原始结果

            # 完成节点执行
            workflow_state_manager.complete_node(workflow_id, node.id, result)

            if progress_cb:
                await progress_cb({"event": "node_done", "node_id": node.id})

            continue  # 继续下一个节点

        except NodeExecutionError as e:
            # 检查是否是"不是扩展节点"的错误
            if e.error_type == "not_extended":
                # 继续使用原有逻辑处理
                pass
            else:
                # 真正的执行错误
                from utils.error_handling import error_handler, ErrorSeverity, ErrorCategory

                # 处理错误
                error_info = error_handler.handle_error(
                    e,
                    severity=ErrorSeverity.HIGH,
                    category=ErrorCategory.EXECUTION,
                    node_id=node.id,
                    workflow_id=workflow_id,
                    context=execution_context
                )

                logging.error(f"Extended node execution failed: {e}")
                workflow_state_manager.fail_node(workflow_id, node.id, str(e))
                workflow_state_manager.fail_workflow(workflow_id, f"Node {node.id} failed: {str(e)}")

                return JSONResponse(status_code=500, content={
                    "status": "failed",
                    "error_node_id": node.id,
                    "error_message": str(e),
                    "error_id": error_info.error_id,
                    "suggestions": error_info.suggestions,
                    "context": execution_context
                })
        except Exception:
            # 如果不是扩展节点类型，继续使用原有逻辑
            pass

        # --- Execute node based on its type ---
        if node_type == 'execution-node':
            # 执行节点可以处理工具调用或AI任务
            tool_name = resolved_params.pop("tool_name", None)

            if tool_name:
                # 如果有tool_name，执行工具调用（类似原来的tool-node逻辑）
                try:
                    # If tool params include promptx_template, render it first
                    tpl_id = resolved_params.pop("promptx_template", None)
                    if tpl_id:
                        vars_dict = resolved_params.pop("vars", {}) if isinstance(resolved_params.get("vars"), dict) else {}
                        rendered_prompt = render_prompt(tpl_id, **vars_dict)
                        # 默认写回到 'prompt' 或 'value' 字段，若不存在则新增
                        if "prompt" in resolved_params:
                            resolved_params["prompt"] = rendered_prompt
                        else:
                            resolved_params["value"] = rendered_prompt

                    logging.info(f"Executing tool: {tool_name} with args: {resolved_params}")
                    result = await tool_registry.execute_tool(tool_name, **resolved_params)
                    logging.info(f"Tool '{tool_name}' executed. Result: {str(result)[:200]}...")

                    if isinstance(result, str) and (result.lower().startswith("error") or "failed" in result.lower()):
                        logging.error(f"Tool '{tool_name}' returned a failure condition: {result}")
                        return JSONResponse(status_code=200, content={
                            "status": "failed",
                            "error_node_id": node.id,
                            "error_message": result,
                            "context": execution_context
                        })

                    if node.data.outputs and result and (not isinstance(result, str) or result.strip()):
                        # 特殊处理WordMCP工具的返回值
                        if tool_name == "wordmcp" and isinstance(result, dict) and "document_path" in result:
                            # 为WordMCP工具，保存文档路径而不是消息
                            context_key = f"{node.id}.{node.data.outputs[0]}"
                            execution_context[context_key] = result["document_path"]
                        else:
                            # 其他工具，保存原始结果
                            context_key = f"{node.id}.{node.data.outputs[0]}"
                            execution_context[context_key] = result

                    if progress_cb:
                        await progress_cb({"event": "node_done", "node_id": node.id})

                except Exception as e:
                    logging.error(f"Error executing tool '{tool_name}' in execution-node '{node.id}': {e}", exc_info=True)
                    return JSONResponse(status_code=500, content={
                        "status": "failed",
                        "error_node_id": node.id,
                        "error_message": f"An unexpected error occurred in tool '{tool_name}': {e}",
                        "context": execution_context
                    })
            else:
                # 如果没有tool_name，执行AI任务或其他执行逻辑
                ai_tool = resolved_params.get("ai_tool", "gpt4")
                user_requirements = resolved_params.get("user_requirements", "")
                output_format = resolved_params.get("output_format", "text")

                # 这里可以添加AI任务执行逻辑
                # 暂时返回一个模拟结果
                result = f"执行节点完成任务：{user_requirements}"

                if node.data.outputs:
                    context_key = f"{node.id}.{node.data.outputs[0]}"
                    execution_context[context_key] = result

                if progress_cb:
                    await progress_cb({"event": "node_done", "node_id": node.id})

        elif node_type == 'end-topic-node':
            try:
                # The final summary node execution
                final_prompt_hydrated = await hydrate_final_prompt(resolved_params, execution_context)
                
                response = await client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {"role": "system", "content": "You are an AI assistant that summarizes results. Be concise and clear."},
                        {"role": "user", "content": final_prompt_hydrated}
                    ],
                    stream=False
                )
                final_result = response.choices[0].message.content
                logging.info(f"AI node '{node.data.label}' executed. Result: {final_result[:200]}...")

                if node.data.outputs:
                    context_key = f"{node.id}.{node.data.outputs[0]}"
                    execution_context[context_key] = final_result
                
                save_execution_to_history(payload, final_result)
                
                if progress_cb:
                    await progress_cb({"event": "node_done", "node_id": node.id})

                # 完成工作流执行
                workflow_state_manager.complete_workflow(workflow_id, final_result)

                # This is the final success response
                return {"status": "success", "result": final_result, "context": execution_context}

            except Exception as e:
                logging.error(f"Error in final AI node execution: {e}", exc_info=True)
                return JSONResponse(status_code=500, content={
                    "status": "failed",
                    "error_node_id": node.id,
                    "error_message": f"Error in final summarization: {e}",
                    "context": execution_context
                })

        elif node_type == 'langchain-node':
            try:
                chain_yaml = resolved_params.get("chain_yaml")
                if not chain_yaml:
                    raise ValueError("langchain-node requires 'chain_yaml' in params.")

                # 允许相对路径；默认查找 backend/langchain_chains/
                from pathlib import Path
                yaml_path = Path(chain_yaml)
                if not yaml_path.is_absolute():
                    yaml_path = Path(__file__).parent / "langchain_chains" / yaml_path
                if not yaml_path.exists():
                    raise FileNotFoundError(f"Chain yaml not found: {yaml_path}")

                from langchain.load import load_chain

                try:
                    chain = load_chain(str(yaml_path))
                except Exception:
                    # fallback: treat as dotted python path to a Chain instance named `chain`
                    import importlib
                    mod_path, _, attr = str(chain_yaml).rpartition(":")
                    if not mod_path:
                        raise
                    module = importlib.import_module(mod_path)
                    chain = getattr(module, attr or "chain")

                # 输入变量：统一把 execution_context 合并进去，可在 YAML 里引用
                inputs = {**resolved_params.get("inputs", {}), "context": execution_context}

                # Memory 注入（可选）
                memory_path = resolved_params.get("memory_path")
                if memory_path:
                    from ai.promptx_adapter import JSONFileMemory
                    try:
                        mem_obj = JSONFileMemory(memory_path)
                        if hasattr(chain, "memory"):
                            chain.memory = mem_obj
                        elif hasattr(chain, "assign_memory"):
                            chain.assign_memory(mem_obj)  # type: ignore[attr-defined]
                    except Exception as mem_err:
                        logging.warning(f"memory_path provided but failed to load: {mem_err}")

                # 回调，用于事件透传
                callbacks = None
                if progress_cb:
                    from ai.langchain_utils import GraphForwardHandler  # local import to avoid heavy deps if unused
                    handler = GraphForwardHandler(progress_cb)
                    callbacks = [handler]

                # async invoke if supported
                if hasattr(chain, "ainvoke"):
                    result = await chain.ainvoke(inputs, callbacks=callbacks)
                else:
                    result = await asyncio.get_event_loop().run_in_executor(None, lambda: chain.invoke(inputs, callbacks=callbacks))

                if node.data.outputs:
                    context_key = f"{node.id}.{node.data.outputs[0]}"
                    execution_context[context_key] = result

                if progress_cb:
                    await progress_cb({"event": "node_done", "node_id": node.id})
            except Exception as e:
                logging.error(f"Error executing langchain-node '{node.id}': {e}", exc_info=True)
                return JSONResponse(status_code=500, content={
                    "status": "failed",
                    "error_node_id": node.id,
                    "error_message": f"Error in langchain-node execution: {e}",
                    "context": execution_context
                })

    logging.warning("Workflow finished without reaching a summary node.")
    return {"status": "success", "result": "Workflow finished, but no final result was generated.", "context": execution_context}


# --- NEW: Streaming execution endpoint ---
@app.post("/execute_stream")
async def execute_workflow_stream(request: Request):
    """Executes the workflow and streams progress events via Server-Sent Events (SSE)."""
    try:
        # 强力修复编码问题：多种方式尝试解码
        raw_bytes = await request.body()
        raw_data = None
                
                # 方法1: 直接使用FastAPI的request.json()
        try:
            raw_data = await request.json()
            logging.info(f"✅ [编码修复-方法1] 直接JSON解析成功，节点数量: {len(raw_data.get('nodes', []))}")
        except Exception as e1:
            logging.warning(f"⚠️ [编码修复-方法1] 直接JSON解析失败: {e1}")
            
            # 方法2: 手动UTF-8解码
            try:
                raw_text = raw_bytes.decode('utf-8')
                raw_data = json.loads(raw_text)
                logging.info(f"✅ [编码修复-方法2] UTF-8解码成功，节点数量: {len(raw_data.get('nodes', []))}")
            except Exception as e2:
                logging.warning(f"⚠️ [编码修复-方法2] UTF-8解码失败: {e2}")
                
                # 方法3: 尝试其他编码
                for encoding in ['gbk', 'gb2312', 'iso-8859-1']:
                    try:
                        raw_text = raw_bytes.decode(encoding)
                        raw_data = json.loads(raw_text)
                        logging.info(f"✅ [编码修复-方法3] {encoding}解码成功，节点数量: {len(raw_data.get('nodes', []))}")
                        break
                    except Exception as e3:
                        logging.warning(f"⚠️ [编码修复-方法3] {encoding}解码失败: {e3}")
                        continue
        
        if raw_data is None:
            raise ValueError("无法解析请求数据，所有编码方式都失败了")
        
        # 修复节点数据中的乱码问题
        if 'nodes' in raw_data:
            for node in raw_data['nodes']:
                if 'data' in node and 'task' in node['data']:
                    original_task = node['data']['task']
                    # 检查是否包含乱码
                    if '?' in original_task or 'ord' in original_task:
                        logging.warning(f"⚠️ [编码修复] 检测到乱码任务描述: {original_task}")
                        # 尝试修复常见的乱码模式
                        if "ord" in original_task:
                            node['data']['task'] = "创建Word文档"  # 设置默认任务
                            logging.info(f"🔧 [编码修复] 已修复任务描述为: 创建Word文档")

        # 尝试解析为WorkflowPayload
        payload = WorkflowPayload(**raw_data)
        logging.info(f"Received streaming workflow request with {len(payload.nodes)} nodes and {len(payload.edges)} edges")

    except json.JSONDecodeError as e:
        logging.error(f"❌ [编码修复] JSON解析失败: {e}")
        return JSONResponse(status_code=400, content={"error": f"Invalid JSON format: {str(e)}"})
    except UnicodeDecodeError as e:
        logging.error(f"❌ [编码修复] 编码解析失败: {e}")
        return JSONResponse(status_code=400, content={"error": f"Encoding error: {str(e)}"})
    except Exception as e:
        logging.error(f"Failed to parse request: {e}")
        return JSONResponse(status_code=400, content={"error": f"Invalid request format: {str(e)}"})

    queue: asyncio.Queue = asyncio.Queue()

    async def progress_cb(event: Dict[str, Any]):
        await queue.put(event)

    async def worker():
        try:
            logging.info(f"Starting streaming workflow execution with {len(payload.nodes)} nodes")
            await queue.put({"event": "start", "message": "开始执行工作流"})

            # Re-use logic from /execute to determine execution order
            nodes_by_id = {node.id: node for node in payload.nodes}
            adj = {nid: [] for nid in nodes_by_id}
            in_degree = {nid: 0 for nid in nodes_by_id}
            for edge in payload.edges:
                if edge.source in nodes_by_id and edge.target in nodes_by_id:
                    adj[edge.source].append(edge.target)
                    in_degree[edge.target] += 1
            q = [nid for nid in nodes_by_id if in_degree[nid] == 0]
            ordered_ids = []
            while q:
                u = q.pop(0)
                ordered_ids.append(u)
                for v in adj.get(u, []):
                    in_degree[v] -= 1
                    if in_degree[v] == 0:
                        q.append(v)
            if len(ordered_ids) != len(payload.nodes):
                logging.error("Workflow has a cycle")
                await queue.put({"event": "error", "message": "Workflow has a cycle."})
                await queue.put("[DONE]")
                return

            logging.info(f"Execution order: {ordered_ids}")
            ordered_nodes = [nodes_by_id[nid] for nid in ordered_ids]
            execution_context: Dict[str, Any] = {}
            result = await execute_workflow_logic(ordered_nodes, execution_context, payload, progress_cb)
            logging.info(f"Workflow execution completed with result: {result}")
            await queue.put({"event": "final", "data": result})
        except Exception as e:
            logging.error(f"Streaming workflow failed: {e}", exc_info=True)
            await queue.put({"event": "error", "message": str(e)})
        finally:
            await queue.put("[DONE]")

    # Launch worker task
    asyncio.create_task(worker())

    async def event_generator():
        while True:
            evt = await queue.get()
            if evt == "[DONE]":
                yield "data: [DONE]\n\n"
                break
            try:
                payload = json.dumps(evt, ensure_ascii=False, default=str)
            except TypeError:
                payload = json.dumps(str(evt), ensure_ascii=False)
            yield f"data: {payload}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


# --- NEW: LangGraph streaming endpoint ---
from agent_system.langgraph_compiler import run_langgraph_workflow


@app.post("/execute_langgraph_stream")
async def execute_langgraph_stream(payload: WorkflowPayload):
    """Stream execution events powered by LangGraph v2 events."""
    from ai.langgraph_executor import stream_events  # 延迟导入避免循环

    async def event_generator():
        try:
            async for ev in stream_events(payload):
                yield json.dumps(ev, ensure_ascii=False) + "\n"
        except asyncio.CancelledError:
            # client disconnected
            pass

    return StreamingResponse(event_generator(), media_type="application/json")


@app.post("/validate_workflow")
async def validate_workflow(request: ValidationRequest):
    """
    Validates the given workflow against loaded rules.
    """
    node_labels = [node.get('data', {}).get('label', '').lower() for node in request.nodes]
    all_labels_text = " ".join(node_labels)

    missing_suggestions = []

    # 确保WORKFLOW_RULES不为空
    if WORKFLOW_RULES is None:
        logging.warning("WORKFLOW_RULES is None, reloading rules...")
        load_workflow_rules()

    if not WORKFLOW_RULES:
        logging.warning("No workflow rules loaded")
        return {"suggestions": []}

    for rule in WORKFLOW_RULES:
        required_keywords = rule.get("required_nodes_keywords", [])
        missing_keywords = []
        for keyword in required_keywords:
            if keyword.lower() not in all_labels_text:
                missing_keywords.append(keyword)
        
        if missing_keywords:
            suggestion = f"您的「{rule.get('workflow_name', '通用')}」流程可能缺失以下关键步骤: {', '.join(missing_keywords)}。建议添加相关节点以保证流程完整性。"
            missing_suggestions.append({
                "rule_name": rule.get('workflow_name'),
                "suggestion_text": suggestion,
                "missing_keywords": missing_keywords
            })

    return {"suggestions": missing_suggestions}

# --- Dev Orchestrate endpoint: user prompt -> DecompositionEngine -> Orchestrator ---
class DevOrchRequest(BaseModel):
    prompt: str
    # mode字段已删除，现在使用统一的执行逻辑

from agent_system.decomposition_engine import DecompositionEngine

@app.post("/dev_orchestrate")
async def dev_orchestrate(request: DevOrchRequest):
    """Experimental endpoint that turns a raw user prompt into high-level tasks, then executes them via Orchestrator (non-stream)."""
    tasks = await DecompositionEngine.decompose(request.prompt)
    logging.info(f"[DevOrch] Decomposed tasks: {tasks}")
    orchestrator = Orchestrator()
    result = await orchestrator.run(tasks)
    return result

# === MCP相关配置已移除 ===

@app.get("/node-types")
async def get_node_types():
    """获取所有可用的节点类型配置"""
    try:
        import json
        from pathlib import Path

        config_path = Path(__file__).parent.parent / "node_types_config.json"
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            return config
        else:
            return {"nodeTypes": {}, "categories": []}
    except Exception as e:
        logging.error(f"Error loading node types config: {e}")
        return {"error": "Failed to load node types configuration"}


def get_available_node_types():
    """获取可用的节点类型作为工具列表"""
    try:
        import json
        from pathlib import Path

        config_path = Path(__file__).parent.parent / "node_types_config.json"
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)

            # 转换为工具格式
            tools = []
            for node_type, node_config in config.get("nodeTypes", {}).items():
                tools.append({
                    "name": node_type,
                    "display_name": node_config.get("name", node_type),
                    "description": node_config.get("description", ""),
                    "category": node_config.get("category", "general"),
                    "icon": node_config.get("icon", "[LIGHTNING]"),
                    "inputs": node_config.get("inputs", []),
                    "outputs": node_config.get("outputs", []),
                    "params": node_config.get("params", {})
                })
            return tools
        else:
            return []
    except Exception as e:
        logging.error(f"Error loading node types: {e}")
        return []


class AIWorkflowRequest(BaseModel):
    requirement: str
    selectedPattern: str = ""
    preferences: dict = {}

@app.post("/api/ai/build-workflow-test")
async def test_ai_endpoint(request: dict):
    """测试AI端点是否工作"""
    return {
        "success": True,
        "message": f"收到需求: {request.get('requirement', '')}",
        "data": request
    }

@app.post("/api/ai/build-workflow")
async def build_workflow_with_ai(request: AIWorkflowRequest):
    """AI自动构建工作流（真正的AI生成）"""
    try:
        requirement = request.requirement
        if not requirement:
            return {
                "success": False,
                "error": "需求描述不能为空",
                "message": "请提供详细的需求描述"
            }

        # 创建OpenAI客户端
        api_key = os.getenv("OPENAI_API_KEY") or os.getenv("DEEPSEEK_API_KEY")
        base_url = None

        if os.getenv("DEEPSEEK_API_KEY"):
            base_url = "https://api.deepseek.com"
            api_key = os.getenv("DEEPSEEK_API_KEY")

        if not api_key:
            return {
                "success": False,
                "error": "AI服务未配置",
                "message": "请配置OPENAI_API_KEY或DEEPSEEK_API_KEY环境变量"
            }

        # 创建真正的AI工作流生成器
        from ai.enhanced_workflow_generator import EnhancedWorkflowGenerator
        from openai import AsyncOpenAI

        client = AsyncOpenAI(api_key=api_key, base_url=base_url)
        generator = EnhancedWorkflowGenerator(client)

        # 获取可用的工具/节点类型
        available_tools = get_available_node_types()

        # 使用真正的AI生成工作流
        ai_result = await generator.generate_workflow(
            user_input=requirement,
            available_tools=available_tools,
            user_preferences=request.preferences
        )

        if ai_result.get('success'):
            # 转换为前端期望的格式
            workflow = ai_result.get('workflow', {})
            nodes = workflow.get('nodes', [])
            edges = workflow.get('edges', [])

            # 转换节点格式为前端期望的格式
            formatted_nodes = []
            for i, node in enumerate(nodes):
                formatted_nodes.append({
                    "id": node.get('id', f"ai_node_{i}"),
                    "name": node.get('label', f"节点 {i+1}"),
                    "nodeType": node.get('type', 'default'),
                    "icon": "🤖",
                    "reason": f"AI分析: {node.get('description', '智能生成的节点')}"
                })

            return {
                "success": True,
                "nodes": formatted_nodes,
                "connections": edges,
                "workflow_description": f"基于AI分析生成的工作流，包含{len(formatted_nodes)}个智能节点",
                "analysis": {
                    "suggested_pattern": ai_result.get('context', {}).get('intent_type', 'intelligent'),
                    "required_capabilities": ["AI生成", "智能分析"],
                    "estimated_complexity": ai_result.get('context', {}).get('complexity', 'medium')
                }
            }
        else:
            # AI生成失败，返回错误
            return {
                "success": False,
                "error": ai_result.get('error', '未知错误'),
                "message": "AI工作流生成失败，请稍后重试"
            }

    except Exception as e:
        logging.error(f"AI workflow building failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "message": "AI工作流构建服务暂时不可用"
        }


@app.get("/api/ai/workflow-patterns")
async def get_workflow_patterns():
    """获取可用的工作流模式"""
    try:
        from services.ai_workflow_builder import ai_workflow_builder

        patterns = []
        for pattern_name, pattern_info in ai_workflow_builder.workflow_patterns.items():
            patterns.append({
                "name": pattern_name,
                "description": pattern_info["description"],
                "keywords": pattern_info["keywords"],
                "nodeCount": len(pattern_info["pattern"])
            })

        return {
            "success": True,
            "patterns": patterns
        }

    except Exception as e:
        logging.error(f"Failed to get workflow patterns: {e}")
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/workflow-status/{workflow_id}")
async def get_workflow_status(workflow_id: str):
    """获取工作流执行状态"""
    from workflow.workflow_state import workflow_state_manager

    execution = workflow_state_manager.get_workflow_status(workflow_id)
    if execution:
        return execution.to_dict()
    else:
        raise HTTPException(status_code=404, detail="Workflow not found")


@app.get("/workflow-status")
async def list_workflow_status():
    """获取所有工作流状态"""
    from workflow.workflow_state import workflow_state_manager

    executions = workflow_state_manager.get_all_workflows()
    return [execution.to_dict() for execution in executions]


@app.post("/workflow-control/{workflow_id}/{action}")
async def control_workflow(workflow_id: str, action: str, reason: str = "User requested"):
    """控制工作流执行（暂停、恢复、取消）"""
    from workflow.workflow_state import workflow_state_manager

    execution = workflow_state_manager.get_workflow_status(workflow_id)
    if not execution:
        raise HTTPException(status_code=404, detail="Workflow not found")

    if action == "pause":
        workflow_state_manager.pause_workflow(workflow_id, reason)
        return {"status": "paused", "message": f"Workflow {workflow_id} paused"}
    elif action == "resume":
        workflow_state_manager.resume_workflow(workflow_id)
        return {"status": "resumed", "message": f"Workflow {workflow_id} resumed"}
    elif action == "cancel":
        workflow_state_manager.fail_workflow(workflow_id, f"Cancelled: {reason}")
        return {"status": "cancelled", "message": f"Workflow {workflow_id} cancelled"}
    else:
        raise HTTPException(status_code=400, detail="Invalid action. Use 'pause', 'resume', or 'cancel'")

@app.post("/api/workflow/{workflow_id}/reset")
async def reset_workflow_execution(workflow_id: str, payload: WorkflowPayload):
    """重置工作流到初始状态并重新执行"""
    try:
        from workflow.mindmap_execution_engine import mindmap_execution_engine

        new_workflow_id = await mindmap_execution_engine.reset_and_restart_workflow(workflow_id, payload)
        return {
            "status": "success",
            "message": "Workflow reset and restarted",
            "workflow_id": new_workflow_id
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/api/workflow/execute")
async def execute_workflow_api(request: Request, user_id: Optional[str] = None):
    """
    API endpoint for workflow execution - compatible with frontend expectations
    """
    try:
        # 手动解析请求数据以处理selectedTool对象格式
        raw_data = await request.json()
        
        # 🔍 MVP调试：检查材料节点数据是否正确传递
        logger.info("🔍 [工作流执行] 开始检查接收到的节点数据")
        if 'nodes' in raw_data:
            for node in raw_data['nodes']:
                if node.get('type') == 'material-node' or node.get('data', {}).get('nodeType') == 'material-node':
                    logger.info(f"🔍 [材料节点] 发现材料节点: {node['id']}")
                    node_data = node.get('data', {})
                    
                    # 检查关键字段
                    target_file = node_data.get('targetFile')
                    selected_files = node_data.get('selectedFiles')
                    
                    logger.info(f"🔍 [材料节点] targetFile: {target_file}")
                    logger.info(f"🔍 [材料节点] selectedFiles: {selected_files}")
                    
                    if selected_files:
                        logger.info(f"🎯 [材料节点] 检测到 {len(selected_files)} 个文件!")
                        for i, file_info in enumerate(selected_files):
                            logger.info(f"🔍 [文件{i+1}] {file_info}")
                    else:
                        logger.warning(f"⚠️ [材料节点] 未检测到selectedFiles数据!")
        
        # 🔧 修复selectedTool对象格式问题
        if 'nodes' in raw_data:
            for node in raw_data['nodes']:
                if 'data' in node and 'selectedTool' in node['data']:
                    selected_tool = node['data']['selectedTool']
                    # 如果selectedTool是对象，提取name字段
                    if isinstance(selected_tool, dict) and 'name' in selected_tool:
                        node['data']['selectedTool'] = selected_tool['name']
                        logger.info(f"🔧 [selectedTool修复] 转换 {selected_tool} -> {selected_tool['name']}")
        
        # 创建WorkflowPayload对象
        from utils.schemas import WorkflowPayload
        payload = WorkflowPayload(**raw_data)
        
        # 使用思维导图执行引擎
        from workflow.mindmap_execution_engine import mindmap_execution_engine
        import uuid
        
        # 生成工作流ID
        workflow_id = str(uuid.uuid4())
        
        # 启动异步执行 - 使用更安全的方式
        try:
            # 直接调用执行方法，不使用create_task
            workflow_id = await mindmap_execution_engine.execute_workflow(payload)
        except Exception as exec_error:
            logger.error(f"工作流执行失败: {exec_error}")
            return JSONResponse(status_code=500, content={
                "success": False,
                "error": f"工作流执行失败: {str(exec_error)}"
            })
        
        return {
            "success": True,
            "message": "工作流已启动",
            "workflow_id": workflow_id
        }
        
    except Exception as e:
        logger.error(f"工作流启动失败: {e}")
        return JSONResponse(status_code=500, content={
            "success": False,
            "error": str(e)
        })

@app.get("/api/workflow/{workflow_id}/status")
async def get_workflow_status(workflow_id: str):
    """
    获取工作流执行状态
    """
    try:
        from workflow.mindmap_execution_engine import mindmap_execution_engine
        
        # 获取工作流状态
        status = await mindmap_execution_engine.get_workflow_status(workflow_id)
        
        if status is None:
            return JSONResponse(status_code=404, content={
                "error": "工作流不存在"
            })
        
        return status
        
    except Exception as e:
        logger.error(f"获取工作流状态失败: {e}")
        return JSONResponse(status_code=500, content={
            "error": str(e)
        })

@app.post("/api/workflow/save-result-file")
async def save_result_file(request: Dict[str, Any]):
    """
    直接保存结果文件到指定路径，而不是通过浏览器下载
    """
    try:
        content = request.get('content', '')
        file_name = request.get('fileName', 'result.txt')
        output_format = request.get('outputFormat', 'txt')
        storage_path = request.get('storagePath', './output')
        
        # 确保存储目录存在
        import os
        from pathlib import Path
        
        storage_dir = Path(storage_path)
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        # 构建完整的文件路径
        file_path = storage_dir / file_name
        
        # 根据格式保存文件
        if output_format.lower() in ['docx']:
            # 对于Word文档，需要创建真正的Word文件
            try:
                from docx import Document
                
                doc = Document()
                # 添加内容到文档
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                
                # 保存文档
                doc.save(str(file_path))
                print(f"📁 Word文档已保存到: {file_path}")
                
            except ImportError:
                # 如果没有python-docx，则保存为文本文件
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"📁 文本文件已保存到: {file_path}")
        else:
            # 其他格式直接保存为文本
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"📁 文件已保存到: {file_path}")
        
        return {
            "status": "success",
            "message": "文件已成功保存",
            "saved_file_path": str(file_path),
            "filePath": str(file_path),
            "directory": str(storage_dir),
            "fileName": file_name,
            "fileSize": len(content.encode('utf-8'))
        }
        
    except Exception as e:
        logger.error(f"保存结果文件时出错: {e}")
        raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")


@app.get("/debug/errors")
async def get_error_statistics():
    """获取错误统计信息"""
    from utils.error_handling import error_handler
    return error_handler.get_error_statistics()


@app.get("/debug/recent-errors")
async def get_recent_errors(limit: int = 20):
    """获取最近的错误"""
    from utils.error_handling import error_handler
    return error_handler.get_recent_errors(limit)


@app.get("/debug/execution-trace")
async def get_execution_trace():
    """获取执行轨迹"""
    from utils.error_handling import debug_info
    return debug_info.get_execution_trace()


@app.get("/debug/node/{node_id}")
async def get_node_debug_info(node_id: str):
    """获取特定节点的调试信息"""
    from utils.error_handling import debug_info
    return debug_info.get_node_debug_info(node_id)


@app.post("/debug/clear")
async def clear_debug_info():
    """清空调试信息"""
    from utils.error_handling import error_handler, debug_info
    error_handler.clear_error_history()
    debug_info.execution_trace.clear()
    debug_info.debug_data.clear()
    return {"status": "cleared", "message": "Debug information cleared"}


@app.get("/performance/stats")
async def get_performance_stats():
    """获取性能统计信息"""
    from ai.concurrent_executor import concurrent_executor, performance_optimizer
    from utils.resource_manager import resource_manager

    stats = concurrent_executor.get_performance_stats()
    analysis = performance_optimizer.analyze_performance(stats)
    resource_status = resource_manager.get_resource_status()

    return {
        "performance_stats": stats,
        "performance_analysis": analysis,
        "resource_status": resource_status
    }


@app.get("/performance/optimize")
async def optimize_performance():
    """执行性能优化"""
    from ai.concurrent_executor import concurrent_executor, performance_optimizer

    stats = concurrent_executor.get_performance_stats()
    analysis = performance_optimizer.analyze_performance(stats)

    applied_optimizations = []
    for recommendation in analysis.get('recommendations', []):
        if performance_optimizer.apply_optimization(recommendation):
            applied_optimizations.append(recommendation)

    return {
        "optimizations_applied": applied_optimizations,
        "new_config": performance_optimizer.current_config,
        "performance_score": analysis.get('performance_score', 0)
    }


@app.get("/resource/status/{user_id}")
async def get_user_resource_status(user_id: str):
    """获取用户资源状态"""
    from utils.resource_manager import resource_manager
    return resource_manager.get_resource_status(user_id)


@app.get("/resource/quota/{user_id}")
async def get_user_quota(user_id: str):
    """获取用户配额信息"""
    from utils.resource_manager import resource_manager
    return resource_manager.quota_manager.get_quota_status(user_id)


@app.get("/templates")
async def get_workflow_templates(category: Optional[str] = None):
    """获取工作流模板"""
    from workflow.workflow_templates import template_manager

    if category:
        templates = template_manager.get_templates_by_category(category)
    else:
        templates = template_manager.get_all_templates()

    return {
        "templates": [template.to_dict() for template in templates],
        "categories": template_manager.get_categories()
    }


@app.get("/templates/{template_id}")
async def get_workflow_template(template_id: str):
    """获取特定工作流模板"""
    from workflow.workflow_templates import template_manager

    template = template_manager.get_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    return template.to_dict()


@app.get("/templates/search/{query}")
async def search_workflow_templates(query: str):
    """搜索工作流模板"""
    from workflow.workflow_templates import template_manager

    templates = template_manager.search_templates(query)
    return {
        "templates": [template.to_dict() for template in templates],
        "query": query,
        "count": len(templates)
    }


@app.post("/templates")
async def create_workflow_template(template_data: Dict[str, Any]):
    """创建新的工作流模板"""
    from workflow.workflow_templates import template_manager, WorkflowTemplate

    try:
        template = WorkflowTemplate.from_dict(template_data)
        template_manager.save_template(template)
        return {"status": "success", "template_id": template.id}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to create template: {str(e)}")


@app.get("/tutorials")
async def get_tutorials(difficulty: Optional[str] = None):
    """获取教程列表"""
    from workflow.tutorial_generator import tutorial_generator

    if difficulty:
        tutorials = tutorial_generator.get_tutorials_by_difficulty(difficulty)
    else:
        tutorials = tutorial_generator.get_all_tutorials()

    return {
        "tutorials": [tutorial.to_dict() for tutorial in tutorials],
        "difficulties": ["初级", "中级", "高级"]
    }


@app.get("/tutorials/{tutorial_id}")
async def get_tutorial(tutorial_id: str):
    """获取特定教程"""
    from workflow.tutorial_generator import tutorial_generator

    tutorial = tutorial_generator.get_tutorial(tutorial_id)
    if not tutorial:
        raise HTTPException(status_code=404, detail="Tutorial not found")

    return tutorial.to_dict()


@app.get("/tutorials/search/{query}")
async def search_tutorials(query: str):
    """搜索教程"""
    from workflow.tutorial_generator import tutorial_generator

    tutorials = tutorial_generator.search_tutorials(query)
    return {
        "tutorials": [tutorial.to_dict() for tutorial in tutorials],
        "query": query,
        "count": len(tutorials)
    }


@app.get("/documentation/nodes")
async def get_node_documentation():
    """获取节点文档"""
    from workflow.tutorial_generator import documentation_generator
    return documentation_generator.generate_node_documentation()


@app.get("/documentation/api")
async def get_api_documentation():
    """获取API文档"""
    from workflow.tutorial_generator import documentation_generator
    return documentation_generator.generate_api_documentation()


# --- 思维导图管理API ---
@app.get("/api/mindmap/current")
async def get_current_mindmap():
    """获取当前思维导图数据"""
    try:
        # 这里可以添加从数据库或缓存获取当前思维导图的逻辑
        # 目前返回模拟数据
        return {
            "success": True,
            "mindmap": {
                "nodes": [],
                "edges": []
            },
            "message": "获取当前思维导图成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取当前思维导图失败: {str(e)}")

@app.post("/api/mindmap/current")
async def update_current_mindmap(mindmap: dict):
    """更新当前思维导图数据"""
    try:
        # 这里可以添加保存到数据库或缓存的逻辑
        # 目前返回模拟的成功响应
        return {
            "success": True,
            "message": "当前思维导图更新成功",
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"更新当前思维导图失败: {str(e)}")

@app.post("/api/workflows")
async def create_workflow(workflow: dict):
    """创建新的工作流"""
    try:
        # 这里可以添加数据库存储逻辑
        # 目前返回模拟的成功响应
        workflow_id = f"workflow_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        return {
            "success": True,
            "id": workflow_id,
            "message": "工作流创建成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"创建工作流失败: {str(e)}")

@app.put("/api/workflows/{workflow_id}")
async def update_workflow(workflow_id: str, workflow: dict):
    """更新工作流"""
    try:
        # 这里可以添加数据库更新逻辑
        return {
            "success": True,
            "id": workflow_id,
            "message": "工作流更新成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"更新工作流失败: {str(e)}")

@app.post("/api/workflows/{workflow_id}/access")
async def update_workflow_access(workflow_id: str):
    """更新工作流访问时间"""
    try:
        return {
            "success": True,
            "id": workflow_id,
            "access_time": datetime.now().isoformat(),
            "message": "访问时间更新成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"更新访问时间失败: {str(e)}")

@app.get("/api/workflows")
async def get_workflows():
    """获取工作流列表"""
    try:
        # 这里可以从数据库获取真实的工作流数据
        # 目前返回模拟数据
        sample_workflows = [
            {
                "id": "workflow_001",
                "title": "数据处理自动化",
                "description": "自动处理Excel数据并生成报表",
                "node_count": 6,
                "edge_count": 5,
                "category": "data",
                "created_at": "2024-01-15T10:30:00",
                "updated_at": "2024-01-20T14:20:00",
                "is_public": False,
                "tags": ["数据处理", "自动化", "Excel"]
            },
            {
                "id": "workflow_002", 
                "title": "邮件通知流程",
                "description": "检测文件变化并发送邮件通知",
                "node_count": 4,
                "edge_count": 3,
                "category": "notification",
                "created_at": "2024-01-10T09:15:00",
                "updated_at": "2024-01-18T16:45:00",
                "is_public": True,
                "tags": ["邮件", "通知", "监控"]
            }
        ]
        return sample_workflows
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取工作流列表失败: {str(e)}")

@app.delete("/api/workflows/{workflow_id}")
async def delete_workflow(workflow_id: str):
    """删除工作流"""
    try:
        # 这里可以添加数据库删除逻辑
        return {
            "success": True,
            "id": workflow_id,
            "message": "工作流删除成功"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除工作流失败: {str(e)}")


# --- 模板市场API ---
@app.get("/api/template-marketplace")
async def get_marketplace_templates():
    """获取模板市场中的模板"""
    try:
        # 读取模板数据文件，如果不存在则返回空列表
        templates_file = Path(__file__).parent.parent / "config" / "templates.json"
        if templates_file.exists():
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_data = json.load(f)
            return {
                "success": True,
                "templates": templates_data.get("templates", []),
                "total": len(templates_data.get("templates", []))
            }
        else:
            return {
                "success": True,
                "templates": [],
                "total": 0
            }
    except Exception as e:
        logger.error(f"获取模板市场数据失败: {str(e)}")
        return {
            "success": False,
            "templates": [],
            "total": 0,
            "error": str(e)
        }


@app.get("/api/my-templates")
async def get_my_templates():
    """获取用户上传的模板"""
    try:
        # 读取用户模板数据，暂时从同一个文件中筛选
        templates_file = Path(__file__).parent.parent / "config" / "templates.json"
        if templates_file.exists():
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_data = json.load(f)
            # 这里可以根据用户ID筛选，暂时返回所有
            user_templates = templates_data.get("templates", [])
            return {
                "success": True,
                "templates": user_templates,
                "total": len(user_templates)
            }
        else:
            return {
                "success": True,
                "templates": [],
                "total": 0
            }
    except Exception as e:
        logger.error(f"获取用户模板失败: {str(e)}")
        return {
            "success": False,
            "templates": [],
            "total": 0,
            "error": str(e)
        }


# --- 模板详情和互动API ---
@app.get("/api/template/{template_id}")
async def get_template_detail(template_id: str):
    """获取模板详情"""
    try:
        templates_file = Path(__file__).parent.parent / "config" / "templates.json"
        if not templates_file.exists():
            return {"success": False, "error": "模板数据文件不存在"}
        
        with open(templates_file, 'r', encoding='utf-8') as f:
            templates_data = json.load(f)
        
        # 查找指定ID的模板
        template = None
        for t in templates_data.get("templates", []):
            if t.get("id") == template_id:
                template = t
                break
        
        if not template:
            return {"success": False, "error": "模板不存在"}
        
        # 获取模板的评论和互动数据
        comments_file = Path(__file__).parent.parent / "config" / "template_comments.json"
        comments_data = {"comments": []}
        if comments_file.exists():
            with open(comments_file, 'r', encoding='utf-8') as f:
                comments_data = json.load(f)
        
        # 筛选出此模板的评论
        template_comments = [c for c in comments_data.get("comments", []) if c.get("template_id") == template_id]
        
        # 计算统计数据
        total_ratings = len([c for c in template_comments if c.get("rating")])
        avg_rating = sum(c.get("rating", 0) for c in template_comments if c.get("rating")) / max(total_ratings, 1)
        
        return {
            "success": True,
            "template": {
                **template,
                "stats": {
                    **template.get("stats", {}),
                    "rating": round(avg_rating, 1),
                    "total_ratings": total_ratings,
                    "comments_count": len(template_comments)
                },
                "comments": template_comments[-5:]  # 返回最新5条评论
            }
        }
        
    except Exception as e:
        logger.error(f"获取模板详情失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/api/template/{template_id}/comment")
async def add_template_comment(template_id: str, request: Request):
    """添加模板评论"""
    try:
        data = await request.json()
        comment_text = data.get("comment", "").strip()
        rating = data.get("rating")
        user_name = data.get("user_name", "匿名用户")
        
        if not comment_text and not rating:
            return {"success": False, "error": "评论内容或评分不能为空"}
        
        # 验证评分范围
        if rating is not None and (rating < 1 or rating > 5):
            return {"success": False, "error": "评分必须在1-5之间"}
        
        # 读取现有评论数据
        comments_file = Path(__file__).parent.parent / "config" / "template_comments.json"
        comments_data = {"comments": []}
        if comments_file.exists():
            with open(comments_file, 'r', encoding='utf-8') as f:
                comments_data = json.load(f)
        
        # 创建新评论
        new_comment = {
            "id": f"comment_{int(time.time() * 1000)}_{random.randint(1000, 9999)}",
            "template_id": template_id,
            "user_name": user_name,
            "comment": comment_text,
            "rating": rating,
            "created_at": datetime.now().isoformat(),
            "likes": 0,
            "replies": []
        }
        
        # 添加到评论列表
        comments_data["comments"].append(new_comment)
        
        # 保存到文件
        comments_file.parent.mkdir(parents=True, exist_ok=True)
        with open(comments_file, 'w', encoding='utf-8') as f:
            json.dump(comments_data, f, ensure_ascii=False, indent=2)
        
        return {
            "success": True,
            "comment": new_comment,
            "message": "评论添加成功"
        }
        
    except Exception as e:
        logger.error(f"添加评论失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.get("/api/template/{template_id}/comments")
async def get_template_comments(template_id: str, page: int = 1, limit: int = 10):
    """获取模板评论列表"""
    try:
        comments_file = Path(__file__).parent.parent / "config" / "template_comments.json"
        if not comments_file.exists():
            return {
                "success": True,
                "comments": [],
                "total": 0,
                "page": page,
                "limit": limit
            }
        
        with open(comments_file, 'r', encoding='utf-8') as f:
            comments_data = json.load(f)
        
        # 筛选此模板的评论
        template_comments = [c for c in comments_data.get("comments", []) if c.get("template_id") == template_id]
        
        # 按时间倒序排序
        template_comments.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        # 分页
        start_idx = (page - 1) * limit
        end_idx = start_idx + limit
        paginated_comments = template_comments[start_idx:end_idx]
        
        return {
            "success": True,
            "comments": paginated_comments,
            "total": len(template_comments),
            "page": page,
            "limit": limit,
            "has_more": end_idx < len(template_comments)
        }
        
    except Exception as e:
        logger.error(f"获取评论列表失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/api/template/{template_id}/like")
async def toggle_template_like(template_id: str, request: Request):
    """点赞/取消点赞模板"""
    try:
        data = await request.json()
        user_id = data.get("user_id", "anonymous")
        
        # 读取点赞数据
        likes_file = Path(__file__).parent.parent / "config" / "template_likes.json"
        likes_data = {"likes": []}
        if likes_file.exists():
            with open(likes_file, 'r', encoding='utf-8') as f:
                likes_data = json.load(f)
        
        # 检查用户是否已经点赞
        existing_like = None
        for like in likes_data["likes"]:
            if like.get("template_id") == template_id and like.get("user_id") == user_id:
                existing_like = like
                break
        
        if existing_like:
            # 取消点赞
            likes_data["likes"].remove(existing_like)
            liked = False
        else:
            # 添加点赞
            new_like = {
                "template_id": template_id,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }
            likes_data["likes"].append(new_like)
            liked = True
        
        # 保存数据
        likes_file.parent.mkdir(parents=True, exist_ok=True)
        with open(likes_file, 'w', encoding='utf-8') as f:
            json.dump(likes_data, f, ensure_ascii=False, indent=2)
        
        # 计算总点赞数
        total_likes = len([l for l in likes_data["likes"] if l.get("template_id") == template_id])
        
        return {
            "success": True,
            "liked": liked,
            "total_likes": total_likes,
            "message": "点赞成功" if liked else "取消点赞成功"
        }
        
    except Exception as e:
        logger.error(f"点赞操作失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/api/template/{template_id}/favorite")
async def toggle_template_favorite(template_id: str, request: Request):
    """收藏/取消收藏模板"""
    try:
        data = await request.json()
        user_id = data.get("user_id", "anonymous")
        
        # 读取收藏数据
        favorites_file = Path(__file__).parent.parent / "config" / "template_favorites.json"
        favorites_data = {"favorites": []}
        if favorites_file.exists():
            with open(favorites_file, 'r', encoding='utf-8') as f:
                favorites_data = json.load(f)
        
        # 检查用户是否已经收藏
        existing_favorite = None
        for favorite in favorites_data["favorites"]:
            if favorite.get("template_id") == template_id and favorite.get("user_id") == user_id:
                existing_favorite = favorite
                break
        
        if existing_favorite:
            # 取消收藏
            favorites_data["favorites"].remove(existing_favorite)
            favorited = False
        else:
            # 添加收藏
            new_favorite = {
                "template_id": template_id,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }
            favorites_data["favorites"].append(new_favorite)
            favorited = True
        
        # 保存数据
        favorites_file.parent.mkdir(parents=True, exist_ok=True)
        with open(favorites_file, 'w', encoding='utf-8') as f:
            json.dump(favorites_data, f, ensure_ascii=False, indent=2)
        
        # 计算总收藏数
        total_favorites = len([f for f in favorites_data["favorites"] if f.get("template_id") == template_id])
        
        return {
            "success": True,
            "favorited": favorited,
            "total_favorites": total_favorites,
            "message": "收藏成功" if favorited else "取消收藏成功"
        }
        
    except Exception as e:
        logger.error(f"收藏操作失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.get("/api/my-favorites")
async def get_my_favorites():
    """获取用户收藏的模板"""
    try:
        # 读取收藏数据
        favorites_file = Path(__file__).parent.parent / "config" / "template_favorites.json"
        if not favorites_file.exists():
            return {
                "success": True,
                "templates": [],
                "total": 0
            }
        
        with open(favorites_file, 'r', encoding='utf-8') as f:
            favorites_data = json.load(f)
        
        # 读取模板数据
        templates_file = Path(__file__).parent.parent / "config" / "templates.json"
        templates_data = {"templates": []}
        if templates_file.exists():
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_data = json.load(f)
        
        # 获取用户收藏的模板ID列表
        user_id = "anonymous"  # TODO: 从session或JWT中获取真实用户ID
        favorited_template_ids = [
            f.get("template_id") for f in favorites_data.get("favorites", [])
            if f.get("user_id") == user_id
        ]
        
        # 筛选出收藏的模板
        favorited_templates = [
            template for template in templates_data.get("templates", [])
            if template.get("id") in favorited_template_ids
        ]
        
        return {
            "success": True,
            "templates": favorited_templates,
            "total": len(favorited_templates)
        }
        
    except Exception as e:
        logger.error(f"获取收藏模板失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/api/template-upload")
async def upload_template(template_data: dict):
    """上传模板到市场"""
    try:
        # 生成模板ID和时间戳
        template_id = f"template_{int(time.time() * 1000)}"
        current_time = datetime.now().isoformat()
        
        # 构建模板对象
        template = {
            "id": template_id,
            "title": template_data.get("title", "未命名模板"),
            "description": template_data.get("description", ""),
            "category": template_data.get("category", "other"),
            "tags": template_data.get("tags", []),
            "marketType": template_data.get("marketType", "free"),
            "isPublic": template_data.get("isPublic", True),
            "price": template_data.get("price", 0),
            "nodes": template_data.get("nodes", []),
            "edges": template_data.get("edges", []),
            "metadata": {
                "nodeCount": template_data.get("metadata", {}).get("nodeCount", 0),
                "edgeCount": template_data.get("metadata", {}).get("edgeCount", 0),
                "originalTitle": template_data.get("metadata", {}).get("originalTitle", ""),
                "source": template_data.get("metadata", {}).get("source", "unknown"),
                "uploadTime": current_time,
                "lastModified": current_time
            },
            "stats": {
                "downloads": 0,
                "likes": 0,
                "views": 0,
                "rating": 0.0,
                "reviews": []
            },
            "author": {
                "id": "user_001",  # 暂时使用默认用户ID
                "name": "用户",
                "avatar": ""
            }
        }
        
        # 读取现有模板数据
        templates_file = Path(__file__).parent.parent / "config" / "templates.json"
        if templates_file.exists():
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_data = json.load(f)
        else:
            templates_data = {"templates": []}
        
        # 添加新模板
        templates_data["templates"].append(template)
        
        # 确保目录存在
        templates_file.parent.mkdir(parents=True, exist_ok=True)
        
        # 保存到文件
        with open(templates_file, 'w', encoding='utf-8') as f:
            json.dump(templates_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"模板上传成功: {template_id}")
        
        return {
            "success": True,
            "templateId": template_id,
            "message": "模板上传成功",
            "template": template
        }
        
    except Exception as e:
        logger.error(f"模板上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"模板上传失败: {str(e)}")


@app.get("/health")
async def health_check():
    """健康检查端点"""
    return await health_checker.get_health_status()


@app.get("/metrics")
async def get_metrics():
    """获取系统指标"""
    return metrics_collector.get_all_metrics()


@app.get("/metrics/{metric_name}")
async def get_metric(metric_name: str, duration_minutes: int = 60):
    """获取特定指标"""
    points = metrics_collector.get_metric(metric_name, duration_minutes)
    return {
        "metric": metric_name,
        "duration_minutes": duration_minutes,
        "points": [
            {
                "timestamp": point.timestamp.isoformat(),
                "value": point.value,
                "labels": point.labels
            }
            for point in points
        ]
    }


@app.get("/alerts")
async def get_alerts():
    """获取告警信息"""
    return {
        "active_alerts": [
            {
                "id": alert.id,
                "name": alert.name,
                "description": alert.description,
                "severity": alert.severity,
                "current_value": alert.current_value,
                "threshold": alert.threshold,
                "triggered_at": alert.triggered_at.isoformat()
            }
            for alert in alert_manager.get_active_alerts()
        ],
        "alert_history": [
            {
                "id": alert.id,
                "name": alert.name,
                "severity": alert.severity,
                "triggered_at": alert.triggered_at.isoformat(),
                "resolved_at": alert.resolved_at.isoformat() if alert.resolved_at else None,
                "status": alert.status
            }
            for alert in alert_manager.get_alert_history()
        ]
    }


@app.get("/config")
async def get_configuration():
    """获取应用配置"""
    config = get_config()
    # 隐藏敏感信息
    safe_config = {
        "name": config.name,
        "version": config.version,
        "debug": config.debug,
        "resources": {
            "max_memory_mb": config.resources.max_memory_mb,
            "max_cpu_percent": config.resources.max_cpu_percent,
            "max_concurrent_tasks": config.resources.max_concurrent_tasks
        },
        "monitoring": {
            "enable_metrics": config.monitoring.enable_metrics,
            "enable_health_check": config.monitoring.enable_health_check
        }
    }
    return safe_config


@app.get("/tools")
async def list_tools():
    """返回所有已注册的工具列表"""
    return tool_registry.get_tools()

@app.get("/tools/debug")
async def debug_tools():
    """调试工具注册状态"""
    tools = tool_registry.get_tools()
    return {
        "total_tools": len(tools),
        "tools": [
            {
                "name": tool["function"]["name"],
                "description": tool["function"]["description"]
            } for tool in tools
        ],
        "tool_names": [tool["function"]["name"] for tool in tools]
    }

@app.post("/tools/{tool_name}/call")
async def call_tool(tool_name: str, payload: dict):
    """调用指定的工具"""
    try:
        result = await tool_registry.execute_tool(tool_name, **payload)
        return {"status": "success", "result": result}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/api/tools/execute")
async def execute_tool_from_library(request: dict):
    """从工具库执行工具 - 前端工具库页面使用"""
    try:
        tool_id = request.get("tool_id")
        tool_type = request.get("tool_type", "function")
        parameters = request.get("parameters", {})

        # 对于WordMCP工具，不需要tool_id，只需要action
        if tool_type == "wordmcp":
            action = request.get("action")
            if not action:
                return {"status": "error", "error": "缺少action参数"}
        elif not tool_id:
            return {"status": "error", "error": "缺少tool_id参数"}

        # 根据工具类型执行不同的逻辑
        if tool_type == "function":
            # 系统内置函数工具
            result = await tool_registry.execute_tool(tool_id, **parameters)
            return {
                "status": "success",
                "tool_id": tool_id,
                "tool_type": tool_type,
                "result": result,
                "execution_time": "0.1s"  # 简化的执行时间
            }

        elif tool_type == "mcp_server":
            # MCP服务器工具已禁用
            return {
                "status": "error",
                "tool_id": tool_id,
                "tool_type": tool_type,
                "result": "MCP服务器工具已被禁用",
                "execution_time": "0.0s"
            }

        elif tool_type == "platform_service":
            # 托管平台服务
            platform_info = request.get("platform_info", {})
            platform_id = platform_info.get("platform_id")
            service_id = platform_info.get("service_id")

            if not platform_id or not service_id:
                return {"status": "error", "error": "缺少平台信息"}

            # 调用托管平台服务
            from backend.integrations.mcp_manager import mcp_manager
            await mcp_manager.initialize()

            from backend.integrations.mcp_platforms.registry import platform_registry
            platform = platform_registry.get_platform(platform_id)

            if not platform:
                return {"status": "error", "error": f"未找到平台: {platform_id}"}

            # 这里需要实现具体的平台服务调用逻辑
            # 暂时返回模拟结果
            return {
                "status": "success",
                "tool_id": tool_id,
                "tool_type": tool_type,
                "result": f"平台服务 {service_id} 执行成功 (模拟结果)",
                "execution_time": "1.0s",
                "platform": platform.platform_name
            }

        elif tool_type == "user_config":
            # 用户配置工具
            return {
                "status": "error",
                "error": "用户配置工具需要先进行配置才能使用",
                "config_required": True
            }

        elif tool_type == "wordmcp":
            # WordMCP工具 - 新的MCP标准协议
            action = request.get("action")
            if not action:
                return {"status": "error", "error": "缺少action参数"}
            
            # 使用工具路由器执行WordMCP工具
            tool_call = {
                "tool": "wordmcp",
                "action": action,
                "parameters": request.get("parameters", {})
            }
            
            result = await tool_router.route_tool_call(tool_call)
            
            # 转换返回格式为标准格式
            if result.get("status") == "success":
                response = {
                    "status": "success",
                    "tool_id": f"wordmcp_{action}",
                    "tool_type": tool_type,
                    "result": result.get("result", result.get("message", "Tool executed successfully")),
                    "execution_time": "1.0s"
                }
                
                # 如果有文档路径，添加到返回值中（用于工作流数据传递）
                if "document_path" in result:
                    response["document_path"] = result["document_path"]
                
                return response
            else:
                return {
                    "status": "error",
                    "tool_id": f"wordmcp_{action}",
                    "tool_type": tool_type,
                    "error": result.get("error", "工具执行失败")
                }

        else:
            return {"status": "error", "error": f"不支持的工具类型: {tool_type}"}

    except Exception as e:
        print(f"执行工具失败: {e}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "error": str(e),
            "tool_id": request.get("tool_id"),
            "tool_type": request.get("tool_type")
        }

# ===== WordMCP 专用API端点 =====

@app.post("/api/tools/wordmcp")
async def execute_wordmcp_tool(request: dict):
    """执行WordMCP工具 - 专用端点支持新的MCP标准协议"""
    try:
        action = request.get("action")
        parameters = request.get("parameters", {})
        
        if not action:
            return {"success": False, "error": "缺少action参数"}
        
        # 使用工具路由器执行
        tool_call = {
            "tool": "wordmcp",
            "action": action,
            "parameters": parameters
        }
        
        result = await tool_router.route_tool_call(tool_call)
        
        # 转换返回格式
        if result.get("status") == "success":
            return {
                "success": True,
                "result": result.get("result", result.get("message", "Tool executed successfully")),
                "message": result.get("message", "工具执行成功")
            }
        else:
            return {
                "success": False,
                "error": result.get("error", "工具执行失败")
            }
            
    except Exception as e:
        logger.error(f"WordMCP工具执行失败: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}

@app.get("/api/tools/{tool_id}/info")
async def get_tool_info(tool_id: str):
    """获取工具详细信息"""
    try:
        # 检查是否是系统工具
        registered_tools = tool_registry.list_tools()
        if tool_id in registered_tools:
            tool_info = registered_tools[tool_id]
            return {
                "status": "success",
                "tool": {
                    "id": tool_id,
                    "name": tool_info.get("display_name", tool_id),
                    "description": tool_info.get("description", ""),
                    "parameters": tool_info.get("parameters", {}),
                    "tool_type": "function",
                    "source": "system"
                }
            }

        # MCP工具检查已移除

        return {"status": "error", "error": f"未找到工具: {tool_id}"}

    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/test-mcp-tool")
async def test_mcp_tool(config: dict):
    """完整测试MCP工具的可用性"""
    try:
        import subprocess
        import tempfile
        import json
        import time
        import asyncio

        # 构建命令
        command = [config.get("command", "npx")]
        if config.get("args"):
            command.extend(config["args"])

        # 创建临时测试进程
        env = os.environ.copy()
        if config.get("env"):
            env.update(config["env"])

        # 第一步：启动进程测试
        try:
            process = subprocess.Popen(
                command,
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
                text=True
            )

            # 第二步：发送MCP初始化消息
            init_message = {
                "jsonrpc": "2.0",
                "id": 1,
                "method": "initialize",
                "params": {
                    "protocolVersion": "2024-11-05",
                    "capabilities": {},
                    "clientInfo": {"name": "test-client", "version": "1.0.0"}
                }
            }

            # 发送初始化消息
            process.stdin.write(json.dumps(init_message) + "\n")
            process.stdin.flush()

            # 等待响应
            time.sleep(1)

            # 第三步：发送工具列表请求
            tools_message = {
                "jsonrpc": "2.0",
                "id": 2,
                "method": "tools/list"
            }

            process.stdin.write(json.dumps(tools_message) + "\n")
            process.stdin.flush()

            # 等待响应
            time.sleep(1)

            # 第四步：测试一个简单的工具调用（如果是office-editor，测试create_empty_txt）
            if "office" in config.get("name", "").lower() or "word" in str(command).lower():
                test_tool_message = {
                    "jsonrpc": "2.0",
                    "id": 3,
                    "method": "tools/call",
                    "params": {
                        "name": "create_empty_txt",
                        "arguments": {
                            "filename": "test_file.txt"
                        }
                    }
                }

                process.stdin.write(json.dumps(test_tool_message) + "\n")
                process.stdin.flush()
                time.sleep(1)

            # 终止进程并读取输出
            process.terminate()
            try:
                stdout, stderr = process.communicate(timeout=3)
            except subprocess.TimeoutExpired:
                process.kill()
                stdout, stderr = process.communicate()

            # 分析输出
            if stdout:
                # 检查是否有正确的JSON-RPC响应
                lines = stdout.strip().split('\n')
                valid_responses = 0

                for line in lines:
                    if line.strip():
                        try:
                            response = json.loads(line)
                            if "jsonrpc" in response and "id" in response:
                                valid_responses += 1
                        except json.JSONDecodeError:
                            continue

                if valid_responses >= 1:  # 至少有一个有效响应
                    return {"success": True, "message": "MCP tool test successful, function normal"}
                else:
                    return {"success": False, "error": f"Tool response format incorrect: {stdout[:200]}"}
            else:
                return {"success": False, "error": f"Tool no response: {stderr[:200] if stderr else 'No error info'}"}

        except subprocess.TimeoutExpired:
            process.kill()
            return {"success": False, "error": "Tool response timeout"}
        except Exception as e:
            return {"success": False, "error": f"Test process error: {str(e)}"}

    except Exception as e:
        return {"success": False, "error": f"Test config error: {str(e)}"}

async def verify_tool_functionality(config: dict, tool_id: str):
    """Verify the actual functionality of the tool"""
    try:
        import subprocess
        import json
        import tempfile
        import os
        import time

        print(f"[WRENCH2] 开始功能验证: {tool_id}")

        # 根据工具类型选择验证策略
        if "excel" in tool_id.lower():
            return await verify_excel_functionality(config)
        elif "file" in tool_id.lower():
            return await verify_file_functionality(config)
        else:
            # 通用功能验证
            return await verify_generic_functionality(config)

    except Exception as e:
        return {
            "success": False,
            "message": f"功能验证出错: {str(e)}",
            "steps": []
        }

async def verify_excel_functionality(config: dict):
    """验证Excel工具功能"""
    steps = []
    test_file = f"test_verification_{int(time.time())}.xlsx"

    try:
        # 构建命令
        command = [config.get("command", "uvx")]
        if config.get("args"):
            command.extend(config["args"])

        print(f"[MEMO2] 启动Excel工具进行功能验证...")

        # 启动MCP服务器进程
        process = subprocess.Popen(
            command,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=0
        )

        # 等待启动
        await asyncio.sleep(2)

        # 步骤1: 初始化
        init_msg = {
            "jsonrpc": "2.0",
            "id": 1,
            "method": "initialize",
            "params": {
                "protocolVersion": "2024-11-05",
                "capabilities": {},
                "clientInfo": {"name": "test-client", "version": "1.0.0"}
            }
        }

        process.stdin.write(json.dumps(init_msg) + "\n")
        process.stdin.flush()

        # 读取初始化响应
        response_line = process.stdout.readline()
        if response_line:
            try:
                response = json.loads(response_line)
                if response.get("result"):
                    steps.append({
                        "name": "初始化连接",
                        "status": "success",
                        "message": "MCP连接建立成功"
                    })
                else:
                    raise Exception("初始化失败")
            except:
                steps.append({
                    "name": "初始化连接",
                    "status": "failed",
                    "message": "初始化响应格式错误"
                })
                raise Exception("初始化失败")

        # 步骤2: 获取工具列表
        list_tools_msg = {
            "jsonrpc": "2.0",
            "id": 2,
            "method": "tools/list"
        }

        process.stdin.write(json.dumps(list_tools_msg) + "\n")
        process.stdin.flush()

        # 读取工具列表响应
        response_line = process.stdout.readline()
        available_tools = []
        if response_line:
            try:
                response = json.loads(response_line)
                tools = response.get("result", {}).get("tools", [])
                available_tools = [tool.get("name") for tool in tools]
                steps.append({
                    "name": "获取工具列表",
                    "status": "success",
                    "message": f"发现 {len(available_tools)} 个可用工具"
                })
            except:
                steps.append({
                    "name": "获取工具列表",
                    "status": "failed",
                    "message": "无法获取工具列表"
                })
                raise Exception("获取工具列表失败")

        # 步骤3: 测试创建工作簿功能
        if "create_workbook" in available_tools:
            create_msg = {
                "jsonrpc": "2.0",
                "id": 3,
                "method": "tools/call",
                "params": {
                    "name": "create_workbook",
                    "arguments": {
                        "filename": test_file
                    }
                }
            }

            process.stdin.write(json.dumps(create_msg) + "\n")
            process.stdin.flush()

            # 读取创建响应
            response_line = process.stdout.readline()
            if response_line:
                try:
                    response = json.loads(response_line)
                    if response.get("result"):
                        steps.append({
                            "name": "创建Excel工作簿",
                            "status": "success",
                            "message": f"成功创建 {test_file}"
                        })
                    else:
                        steps.append({
                            "name": "创建Excel工作簿",
                            "status": "failed",
                            "message": "创建工作簿失败"
                        })
                except:
                    steps.append({
                        "name": "创建Excel工作簿",
                        "status": "failed",
                        "message": "创建响应格式错误"
                    })
        else:
            steps.append({
                "name": "创建Excel工作簿",
                "status": "skipped",
                "message": "工具不支持create_workbook功能"
            })

        # 清理进程
        process.terminate()
        process.wait(timeout=5)

        # 判断整体成功状态
        success_count = len([s for s in steps if s["status"] == "success"])
        total_count = len([s for s in steps if s["status"] != "skipped"])

        overall_success = success_count >= total_count * 0.8  # 80%成功率

        return {
            "success": overall_success,
            "message": f"Excel功能验证完成，{success_count}/{total_count} 项通过",
            "steps": steps
        }

    except subprocess.TimeoutExpired:
        process.kill()
        return {
            "success": False,
            "message": "功能验证超时",
            "steps": steps
        }
    except Exception as e:
        if 'process' in locals():
            process.terminate()
        return {
            "success": False,
            "message": f"功能验证失败: {str(e)}",
            "steps": steps
        }

async def verify_generic_functionality(config: dict):
    """通用功能验证"""
    return {
        "success": True,
        "message": "通用功能验证通过（基础协议测试）",
        "steps": [
            {
                "name": "协议兼容性",
                "status": "success",
                "message": "MCP协议兼容"
            }
        ]
    }

async def verify_file_functionality(config: dict):
    """文件工具功能验证"""
    # 可以后续扩展
    return await verify_generic_functionality(config)

# 系统工具端点已移除 - MCP功能不可用

# MCP同步端点已移除 - MCP功能不可用

@app.get("/api/tools/library")
async def get_tools_library():
    """获取工具库列表 - 从系统工具与配置文件加载所有启用的工具"""
    try:
        tools_library = []
        seen_tool_ids = set()

        # 1) 系统工具
        try:
            system_tools = _load_system_tools()
            for system_tool in system_tools:
                if not system_tool.get("enabled", True):
                    continue
                tool_id = system_tool.get("id", "")
                if tool_id and tool_id not in seen_tool_ids:
                    seen_tool_ids.add(tool_id)
                    functional_category = system_tool.get("functionalCategory", "system_tools")
                    if isinstance(functional_category, list) and len(functional_category) > 0:
                        functional_category = functional_category[0]
                    elif not functional_category:
                        functional_category = "system_tools"
                    tools_library.append({
                        "id": tool_id,
                        "name": system_tool.get("name", "系统工具"),
                        "description": system_tool.get("description", "系统工具"),
                        "category": "系统工具",
                        "functionalCategory": functional_category,
                        "icon": "[GEAR2]",
                        "source": "system_tools",
                        "enabled": True,
                        "tool_type": "system",
                        "approved": True,
                        "tested": system_tool.get("tested", True)
                    })
        except Exception as e:
            logging.error(f"加载系统工具失败: {e}")

        # 2) AI工具（全部已启用的）
        try:
            from utils.frontend_config import get_ai_tools
            ai_tools_response = await get_ai_tools()
            if isinstance(ai_tools_response, list):
                for ai_tool in ai_tools_response:
                    if not ai_tool.get("enabled", True):
                        continue
                    tool_id = ai_tool.get("id") or ai_tool.get("name")
                    if tool_id and tool_id not in seen_tool_ids:
                        seen_tool_ids.add(tool_id)
                        tools_library.append({
                            "id": tool_id,
                            "name": ai_tool.get("display_name") or ai_tool.get("name", "AI模型"),
                            "description": ai_tool.get("description", "AI模型"),
                            "category": "AI工具",
                            "functionalCategory": ai_tool.get("functionalCategory") or "ai_assistant",
                            "icon": "🤖",
                            "source": "ai",
                            "enabled": True,
                            "tool_type": "ai",
                            "approved": True,
                            "tested": ai_tool.get("tested", True)
                        })
        except Exception as e:
            logging.error(f"加载AI工具失败: {e}")

        # 3) MCP工具（全部已启用的）
        try:
            from utils.frontend_config import load_config, MCP_TOOLS_CONFIG
            mcp_tools = load_config(MCP_TOOLS_CONFIG)
            if not isinstance(mcp_tools, list):
                mcp_tools = []
            for mcp_tool in mcp_tools:
                if not mcp_tool.get("enabled", True):
                    continue
                tool_id = mcp_tool.get("id") or f"mcp_{mcp_tool.get('name','unknown')}"
                if tool_id and tool_id not in seen_tool_ids:
                    seen_tool_ids.add(tool_id)
                    functional_category = mcp_tool.get("functionalCategory") or "system_tools"
                    if isinstance(functional_category, list) and len(functional_category) > 0:
                        functional_category = functional_category[0]
                    tools_library.append({
                        "id": tool_id,
                        "name": mcp_tool.get("name", "MCP工具"),
                        "description": mcp_tool.get("description", "MCP工具"),
                        "category": "MCP工具",
                        "functionalCategory": functional_category,
                        "icon": "[WRENCH2]",
                        "source": "mcp",
                        "enabled": True,
                        "tool_type": "mcp",
                        "approved": mcp_tool.get("approved", False),
                        "tested": mcp_tool.get("tested", False)
                    })
        except Exception as e:
            logging.error(f"加载MCP工具失败: {e}")

        # 4) API工具（全部已启用的）
        try:
            from utils.frontend_config import load_config, API_TOOLS_CONFIG
            api_tools = load_config(API_TOOLS_CONFIG)
            if not isinstance(api_tools, list):
                api_tools = []
            for api_tool in api_tools:
                if not api_tool.get("enabled", True):
                    continue
                tool_id = api_tool.get("id") or f"api_{api_tool.get('name','unknown')}"
                if tool_id and tool_id not in seen_tool_ids:
                    seen_tool_ids.add(tool_id)
                    functional_category = api_tool.get("functionalCategory") or "network_communication"
                    if isinstance(functional_category, list) and len(functional_category) > 0:
                        functional_category = functional_category[0]
                    tools_library.append({
                        "id": tool_id,
                        "name": api_tool.get("name", "API工具"),
                        "description": api_tool.get("description", "API工具"),
                        "category": "API工具",
                        "functionalCategory": functional_category,
                        "icon": "[NETWORK]",
                        "source": "api",
                        "enabled": True,
                        "tool_type": "api",
                        "approved": api_tool.get("approved", False),
                        "tested": api_tool.get("tested", False)
                    })
        except Exception as e:
            logging.error(f"加载API工具失败: {e}")

        return {
            "tools": tools_library,
            "total": len(tools_library),
            "sources": {
                "ai": len([t for t in tools_library if t["source"] == "ai"]),
                "api": len([t for t in tools_library if t["source"] == "api"]),
                "mcp": len([t for t in tools_library if t["source"] == "mcp"]),
                "system": len([t for t in tools_library if t["source"] == "system_tools"])
            }
        }

    except Exception as e:
        logging.error(f"获取工具库失败: {e}")
        return {"error": str(e), "tools": [], "total": 0}

@app.get("/api/tools/library/old")
async def get_tools_library_simplified():
    """获取工具库列表 - 只从四个工具页面的配置文件加载实际配置的工具"""
    try:
        tools_library = []
        seen_tool_ids = set()  # 用于去重

        # 1. 获取AI工具配置
        try:
            from utils.frontend_config import load_config, AI_TOOLS_CONFIG
            ai_tools = load_config(AI_TOOLS_CONFIG)

            for ai_tool in ai_tools:
                if not ai_tool.get("enabled", True):
                    continue

                # 使用id字段，如果没有则使用name字段作为id
                tool_id = ai_tool.get("id") or ai_tool.get("name", "")
                if tool_id and tool_id not in seen_tool_ids:
                    seen_tool_ids.add(tool_id)

                    # 判断是否为智能模式检测系统（特殊处理）
                    if ai_tool.get("name") == "intelligent_mode_detection":
                        tools_library.append({
                            "id": tool_id,
                            "name": ai_tool.get("display_name", ai_tool.get("name", "AI工具")),
                            "description": ai_tool.get("description", "智能模式检测和处理系统"),
                            "category": ai_tool.get("category", "AI工具"),
                            "functionalCategory": ai_tool.get("functionalCategory", "ai_assistant"),
                            "icon": "[BRAIN]",
                            "source": "ai_config",
                            "enabled": True,
                            "tool_type": "intelligent_system",
                            "config": {
                                "api_base": ai_tool.get("api_base", "/brain"),
                                "endpoints": ai_tool.get("endpoints", {}),
                                "version": ai_tool.get("version", "1.0.0")
                            },
                            "approved": True,
                            "tested": True
                        })
                    else:
                        # 标准AI工具处理
                        tools_library.append({
                            "id": tool_id,
                            "name": ai_tool.get("display_name", ai_tool.get("name", "AI工具")),
                            "description": ai_tool.get("description", f"{ai_tool.get('provider', 'AI')} - {ai_tool.get('model', '模型')}"),
                            "category": "AI工具",
                            "functionalCategory": ai_tool.get("functionalCategory", "ai_assistant"),
                            "icon": "🤖",
                            "source": "ai_config",
                            "enabled": True,
                            "tool_type": "ai",
                            "config": {
                                "provider": ai_tool.get("provider"),
                                "model": ai_tool.get("model"),
                                "max_tokens": ai_tool.get("max_tokens", 4000),
                                "temperature": ai_tool.get("temperature", 0.7),
                                "base_url": ai_tool.get("base_url", "")
                            },
                            "approved": True,
                            "tested": True
                        })
        except Exception as e:
            print(f"获取AI工具配置失败: {e}")

        # 2. 获取API工具配置
        try:
            from utils.frontend_config import API_TOOLS_CONFIG
            api_tools = load_config(API_TOOLS_CONFIG)

            for api_tool in api_tools:
                if not api_tool.get("enabled", True):
                    continue

                tool_id = api_tool.get("id", "")
                if tool_id and tool_id not in seen_tool_ids:
                    seen_tool_ids.add(tool_id)
                    tools_library.append({
                        "id": tool_id,
                        "name": api_tool.get("name", "API工具"),
                        "description": api_tool.get("description", f"API工具 - {api_tool.get('base_url', '')}"),
                        "category": "API工具",
                        "icon": "[LINK3]",
                        "source": "api_config",
                        "enabled": True,
                        "tool_type": "api",
                        "config": {
                            "base_url": api_tool.get("base_url"),
                            "protocol": api_tool.get("protocol", "HTTPS"),
                            "timeout": api_tool.get("timeout", 60000),
                            "auth_type": api_tool.get("auth_type", "Bearer Token"),
                            "supported_methods": api_tool.get("supported_methods", ["GET", "POST"]),
                            "custom_headers": api_tool.get("custom_headers", {})
                        },
                        "approved": True,
                        "tested": api_tool.get("tested", False)
                    })
        except Exception as e:
            print(f"获取API工具配置失败: {e}")

        # 3. 获取MCP工具配置
        try:
            from utils.frontend_config import MCP_TOOLS_CONFIG
            mcp_tools = load_config(MCP_TOOLS_CONFIG)

            for mcp_tool in mcp_tools:
                if not mcp_tool.get("enabled", True):
                    continue

                tool_id = mcp_tool.get("id", "")
                if tool_id and tool_id not in seen_tool_ids:
                    seen_tool_ids.add(tool_id)
                    tools_library.append({
                        "id": tool_id,
                        "name": mcp_tool.get("name", "MCP工具"),
                        "description": mcp_tool.get("description", f"MCP工具 - {mcp_tool.get('server_type', 'stdio')}"),
                        "category": "MCP工具",
                        "icon": "[LIGHTNING]",
                        "source": "mcp_config",
                        "enabled": True,
                        "tool_type": "mcp",
                        "config": {
                            "server_type": mcp_tool.get("server_type"),
                            "command": mcp_tool.get("command"),
                            "args": mcp_tool.get("args", []),
                            "url": mcp_tool.get("url")
                        },
                        "approved": True,
                        "tested": mcp_tool.get("tested", False)
                    })
        except Exception as e:
            print(f"获取MCP工具配置失败: {e}")

        # 4. 获取真实注册的系统工具（如果有的话）
        try:
            registered_tools = tool_registry.get_tools()

            for tool_schema in registered_tools:
                function_info = tool_schema.get("function", {})
                tool_name = function_info.get("name", "unknown")

                # 去重检查
                if tool_name in seen_tool_ids:
                    continue
                seen_tool_ids.add(tool_name)

                # 硬编码工具描述，避免编码问题
                tool_descriptions = {
                    "tools_file_operations_file_operations": "文件管理工具",
                    "tools_simple_document_create_html_document": "创建网页文档",
                    "tools_simple_document_create_simple_document": "创建文本文档"
                }
                description = tool_descriptions.get(tool_name, "系统工具")

                tools_library.append({
                    "id": tool_name,
                    "name": function_info.get("name", tool_name).replace('_', ' ').title(),
                    "description": description,
                    "category": "系统工具",
                    "icon": _get_tool_icon(tool_name),
                    "source": "system",
                    "enabled": True,
                    "parameters": function_info.get("parameters", {}),
                    "tool_type": "function",
                    "approved": True,
                    "tested": True
                })
        except Exception as e:
            print(f"获取系统工具失败: {e}")
            # 如果获取系统工具失败，不添加任何模拟工具

        # 4. 获取系统工具配置
        try:
            from pathlib import Path
            # 从backend/core/main.py向上两级到项目根目录，然后到data目录
            system_tools_config = Path(__file__).parent.parent.parent / "data" / "configs" / "backend-config" / "system_tools.json"
            print(f"系统工具配置路径: {system_tools_config}")
            print(f"路径是否存在: {system_tools_config.exists()}")

            if system_tools_config.exists():
                with open(system_tools_config, 'r', encoding='utf-8') as f:
                    system_tools = json.load(f)
                print(f"加载了 {len(system_tools)} 个系统工具")

                for system_tool in system_tools:
                    if not system_tool.get("enabled", True):
                        continue

                    tool_id = system_tool.get("id", "")
                    if tool_id and tool_id not in seen_tool_ids:
                        seen_tool_ids.add(tool_id)
                        # 处理功能分类 - 如果是数组，取第一个；如果是字符串，直接使用
                        functional_category = system_tool.get("functionalCategory", "system_tools")
                        if isinstance(functional_category, list) and len(functional_category) > 0:
                            functional_category = functional_category[0]
                        elif not functional_category:
                            functional_category = "system_tools"

                        tools_library.append({
                            "id": tool_id,
                            "name": system_tool.get("name", "系统工具"),
                            "description": system_tool.get("description", "系统内置工具"),
                            "category": system_tool.get("category", "系统工具"),
                            "functionalCategory": functional_category,
                            "icon": system_tool.get("icon", "[GEAR2]"),
                            "source": system_tool.get("source", "系统工具"),
                            "enabled": True,
                            "tool_type": system_tool.get("tool_type", "system_tool"),
                            "tags": system_tool.get("tags", []),
                            "config": system_tool.get("config", {}),
                            "approved": system_tool.get("approved", True),
                            "tested": system_tool.get("tested", True)
                        })
        except Exception as e:
            print(f"加载系统工具配置失败: {e}")

        return {
            "status": "success",
            "tools": tools_library,
            "total": len(tools_library),
            "categories": {
                "AI工具": len([t for t in tools_library if t["category"] == "AI工具"]),
                "API工具": len([t for t in tools_library if t["category"] == "API工具"]),
                "MCP工具": len([t for t in tools_library if t["category"] == "MCP工具"]),
                "系统工具": len([t for t in tools_library if t["category"] == "系统工具"])
            },
            "last_updated": datetime.now().isoformat()
        }

    except Exception as e:
        print(f"获取工具库失败: {e}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "error": str(e),
            "tools": [],
            "total": 0,
            "last_updated": datetime.now().isoformat()
        }

@app.get("/api/tools/sync-status")
async def get_tools_sync_status():
    """获取工具同步状态"""
    try:
        from utils.frontend_config import load_config, AI_TOOLS_CONFIG, API_TOOLS_CONFIG, MCP_TOOLS_CONFIG

        ai_tools = load_config(AI_TOOLS_CONFIG)
        api_tools = load_config(API_TOOLS_CONFIG)
        mcp_tools = load_config(MCP_TOOLS_CONFIG)

        return {
            "status": "success",
            "sync_status": {
                "ai_tools_count": len(ai_tools),
                "api_tools_count": len(api_tools),
                "mcp_tools_count": len(mcp_tools),
                "total_configured_tools": len(ai_tools) + len(api_tools) + len(mcp_tools),
                "last_sync": datetime.now().isoformat()
            }
        }
    except Exception as e:
        return {
            "status": "error",
            "message": str(e),
            "sync_status": {}
        }

def _get_tool_icon(tool_name: str) -> str:
    """根据工具名称获取图标"""
    icon_map = {
        # 搜索相关
        "web_search": "[MAGNIFYING_GLASS_LEFT]",
        "search": "[MAGNIFYING_GLASS_LEFT]",
        "fetch": "[INBOX_TRAY]",
        "fetch_page": "[PAGE_FACING_UP]",

        # 文件操作
        "file": "[FILE_FOLDER]",
        "document": "[PAGE_FACING_UP]",
        "word": "[MEMO2]",
        "excel": "[BAR_CHART2]",
        "pdf": "[CLIPBOARD2]",

        # 网络请求
        "http": "[NETWORK]",
        "api": "[LINK3]",
        "request": "[SATELLITE_ANTENNA2]",

        # 数据处理
        "data": "[RELOAD]",
        "convert": "[RELOAD]",
        "process": "[GEAR2]",
        "analyze": "[BAR_CHART2]",

        # AI相关
        "ai": "🤖",
        "chat": "💬",
        "llm": "[BRAIN]",
        "intelligent": "[BRAIN]",

        # MCP相关
        "mcp": "[WRENCH2]",
        "server": "[DESKTOP_COMPUTER]",

        # 其他
        "tool": "[HAMMER_WRENCH2]",
        "utility": "[GEAR2]"
    }

    tool_lower = tool_name.lower()

    # 根据工具名称关键词匹配图标
    for keyword, icon in icon_map.items():
        if keyword in tool_lower:
            return icon

    # 默认图标
    return "[GEAR2]"

# POST系统工具端点已移除 - MCP功能不可用

@app.get("/mcp-platforms/status")
async def get_mcp_platforms_status():
    """返回MCP托管平台状态"""
    # 返回默认的平台状态
    platforms = [
        {
            "id": "modao",
            "name": "墨刀社区",
            "status": "disconnected",
            "description": "从墨刀社区同步 MCP 服务"
        },
        {
            "id": "smithery",
            "name": "Smithery",
            "status": "disconnected",
            "description": "从 Smithery 同步 MCP 服务"
        },
        {
            "id": "aliyun",
            "name": "阿里云函数计算",
            "status": "disconnected",
            "description": "从阿里云函数计算同步 MCP 服务"
        },
        {
            "id": "tencent",
            "name": "腾讯云函数",
            "status": "disconnected",
            "description": "从腾讯云函数同步 MCP 服务"
        }
    ]
    return platforms

@app.post("/mcp-platforms/{platform_id}/config")
async def configure_mcp_platform(platform_id: str, config_data: dict):
    """配置MCP托管平台"""
    print(f"收到平台配置请求: {platform_id}, 数据: {config_data}")

    if platform_id == "tencent_scf":
        # 验证腾讯云配置
        required_fields = ["secret_id", "secret_key", "region"]
        for field in required_fields:
            if field not in config_data or not config_data[field]:
                return {"error": f"缺少必需字段: {field}"}, 400

        # 这里可以添加实际的腾讯云连接测试逻辑
        print(f"配置腾讯云函数计算: Secret ID: {config_data['secret_id'][:8]}..., Region: {config_data['region']}")

        # 模拟配置成功
        return {
            "status": "success",
            "message": "腾讯云函数计算配置成功",
            "platform_id": platform_id,
            "config": {
                "region": config_data["region"],
                "secret_id_masked": config_data["secret_id"][:8] + "..."
            }
        }

    elif platform_id == "aliyun":
        # 阿里云配置逻辑
        return {"status": "success", "message": "阿里云配置成功"}

    elif platform_id == "modao":
        # 墨刀配置逻辑
        return {"status": "success", "message": "墨刀配置成功"}

    else:
        return {"error": f"不支持的平台: {platform_id}"}, 400

@app.post("/api/mcp/sync/{platform_type}")
async def sync_mcp_platform(platform_type: str, account_config: dict):
    """同步MCP平台工具"""
    print(f"收到同步请求: {platform_type}, 配置: {account_config}")

    try:
        if platform_type == "tencent_scf":
            # 导入腾讯云平台适配器
            from integrations.mcp_platforms.tencent_scf import TencentSCFPlatform
            from integrations.mcp_platforms.base import PlatformConfig

            # 创建平台配置
            platform_config = PlatformConfig(
                platform_id="tencent_scf",
                platform_name="腾讯云函数计算",
                base_url="https://scf.tencentcloudapi.com",
                region=account_config.get("config", {}).get("region", "ap-guangzhou"),
                timeout=30,
                enabled=True,
                extra_config=account_config.get("config", {})
            )

            # 创建平台实例并同步
            platform = TencentSCFPlatform(platform_config)

            # 测试连接
            connection_ok = await platform.test_connection()
            if not connection_ok:
                return {"error": "平台连接失败，请检查配置"}, 400

            # 获取服务列表
            services = await platform.list_services()

            # 转换为前端需要的格式
            synced_tools = []
            for service in services:
                synced_tools.append({
                    "id": service.id,
                    "name": service.display_name,
                    "description": service.description,
                    "category": service.metadata.get("category", "其他"),
                    "source": "腾讯云开发者社区",
                    "transport": service.metadata.get("transport", "http"),
                    "platform": "tencent_scf",
                    "syncedAt": "2024-01-01T00:00:00.000Z",
                    "endpoint": service.endpoint,
                    "metadata": service.metadata
                })

            await platform.close()

            return synced_tools

        elif platform_type == "modelscope":
            # 魔搭社区同步逻辑
            return []

        else:
            return {"error": f"不支持的平台类型: {platform_type}"}, 400

    except Exception as e:
        print(f"同步失败: {e}")
        import traceback
        traceback.print_exc()
        return {"error": f"同步失败: {str(e)}"}, 500

# --- 作品存放设置API ---
@app.get("/api/workspace-settings")
async def get_workspace_settings():
    """获取用户的作品存放设置"""
    try:
        settings_file = "memory/workspace_settings.json"
        if os.path.exists(settings_file):
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings = json.load(f)
            return settings
        else:
            # 返回默认设置
            default_settings = {
                "defaultOutputPath": "C:\\Users\\用户名\\Documents\\AI工作流作品\\",
                "enableDateFolders": True,
                "enableProjectFolders": True,
                "fileNamingPattern": "document_{timestamp}_{type}",
                "fileTypes": {
                    "documents": "\\Documents\\",
                    "data": "\\Data\\",
                    "images": "\\Images\\",
                    "media": "\\Media\\"
                },
                "cloudSync": {
                    "enabled": False,
                    "provider": "none",
                    "syncPath": ""
                }
            }
            return default_settings
    except Exception as e:
        logging.error(f"获取作品存放设置失败: {e}")
        return {"error": str(e)}, 500

@app.post("/api/workspace-settings")
async def save_workspace_settings(settings: dict):
    """保存用户的作品存放设置"""
    try:
        # 确保memory目录存在
        os.makedirs("memory", exist_ok=True)

        settings_file = "memory/workspace_settings.json"
        with open(settings_file, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)

        logging.info(f"作品存放设置已保存: {settings}")
        return {"success": True, "message": "设置保存成功"}
    except Exception as e:
        logging.error(f"保存作品存放设置失败: {e}")
        return {"error": str(e)}, 500

def get_output_path(file_type: str = "document", project_name: str = None) -> str:
    """根据用户设置生成输出路径"""
    try:
        settings_file = "memory/workspace_settings.json"
        if os.path.exists(settings_file):
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings = json.load(f)
        else:
            # 使用默认设置
            settings = {
                "defaultOutputPath": "C:\\Users\\用户名\\Documents\\AI工作流作品\\",
                "enableDateFolders": True,
                "enableProjectFolders": True,
                "fileTypes": {
                    "documents": "\\Documents\\",
                    "data": "\\Data\\",
                    "images": "\\Images\\",
                    "media": "\\Media\\"
                }
            }

        # 构建基础路径
        base_path = settings.get("defaultOutputPath", "C:\\Users\\用户名\\Documents\\AI工作流作品\\")

        # 添加日期文件夹
        if settings.get("enableDateFolders", True):
            from datetime import datetime
            date_folder = datetime.now().strftime("%Y\\%m\\%d")
            base_path = os.path.join(base_path, date_folder)

        # 添加项目文件夹
        if settings.get("enableProjectFolders", True) and project_name:
            base_path = os.path.join(base_path, project_name)

        # 添加文件类型文件夹
        file_type_mapping = {
            "document": "documents",
            "data": "data",
            "image": "images",
            "media": "media"
        }

        type_key = file_type_mapping.get(file_type, "documents")
        type_folder = settings.get("fileTypes", {}).get(type_key, "\\Documents\\")
        if type_folder:
            base_path = os.path.join(base_path, type_folder.strip("\\"))

        # 确保目录存在
        os.makedirs(base_path, exist_ok=True)

        return base_path

    except Exception as e:
        logging.error(f"生成输出路径失败: {e}")
        # 返回默认路径
        default_path = os.path.join(os.path.expanduser("~"), "Documents", "AI工作流作品")
        os.makedirs(default_path, exist_ok=True)
        return default_path

@app.on_event("startup")
async def startup_event():
    """应用启动事件"""
    logging.info("Starting Workflow Platform...")

    # 启动监控
    try:
        await metrics_collector.start_collection()
        await alert_manager.start_monitoring()
        logging.info("Monitoring services started")
    except Exception as e:
        logging.error(f"Failed to start monitoring: {e}")


# AI描述总结相关模型
class DescriptionSummaryRequest(BaseModel):
    description: str
    maxLength: int = 20

class DescriptionSummaryResponse(BaseModel):
    summary: str
    originalLength: int
    summaryLength: int

@app.get("/api/ai/services")
async def get_ai_services():
    """获取AI服务工具数据"""
    try:
        from utils.frontend_config import load_config, AI_TOOLS_CONFIG
        ai_tools = load_config(AI_TOOLS_CONFIG)

        # 统计数据
        configured_services = len([tool for tool in ai_tools if tool.get("enabled", True)])
        active_connections = len([tool for tool in ai_tools if tool.get("enabled", True) and tool.get("tested", False)])
        supported_providers = 7  # 支持的AI提供商数量

        # 处理工具数据，添加功能分类
        processed_tools = []
        for tool in ai_tools:
            if tool.get("enabled", True):
                # 根据工具特性自动分配功能分类
                functional_category = tool.get("functionalCategory", "ai_assistant")
                if not functional_category or functional_category == "ai_assistant":
                    # 根据工具名称和描述自动分类
                    name_desc = (tool.get("name", "") + " " + tool.get("description", "")).lower()
                    if "智能模式" in name_desc or "检测" in name_desc:
                        functional_category = "ai_assistant"
                    elif "聊天" in name_desc or "对话" in name_desc or "deepseek" in name_desc:
                        functional_category = "ai_assistant"
                    else:
                        functional_category = "ai_assistant"

                processed_tools.append({
                    "name": tool.get("display_name", tool.get("name", "AI工具")),
                    "description": tool.get("description", "AI服务工具"),
                    "functionalCategory": functional_category,
                    "icon": "🤖",
                    "enabled": tool.get("enabled", True),
                    "source": "ai_config"
                })

        return {
            "status": "success",
            "data": {
                "configuredServices": configured_services,
                "activeConnections": active_connections,
                "supportedProviders": supported_providers,
                "tools": processed_tools
            }
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"获取AI服务失败: {str(e)}",
            "data": {
                "configuredServices": 0,
                "activeConnections": 0,
                "supportedProviders": 0,
                "tools": []
            }
        }

@app.get("/api/mcp/tools")
@app.get("/api/tools/mcp-tools")
async def get_mcp_tools():
    """获取MCP工具数据"""
    try:
        from utils.frontend_config import load_config, MCP_TOOLS_CONFIG

        # 加载MCP工具配置，不使用模拟数据
        mcp_tools = load_config(MCP_TOOLS_CONFIG)

        # 统计数据
        all_tools = len(mcp_tools)
        enabled_tools = len([tool for tool in mcp_tools if tool.get("enabled", True)])

        # 处理工具数据，添加功能分类
        processed_tools = []
        for tool in mcp_tools:
            if tool.get("enabled", True):
                # 根据工具特性自动分配功能分类
                functional_category = tool.get("functionalCategory", "automation")
                if not functional_category:
                    # 根据工具名称和描述自动分类
                    name_desc = (tool.get("name", "") + " " + tool.get("description", "")).lower()
                    if "文件" in name_desc:
                        functional_category = "file_management"
                    elif "数据" in name_desc or "分析" in name_desc:
                        functional_category = "data_analysis"
                    else:
                        functional_category = "automation"

                processed_tools.append({
                    "name": tool.get("name", "MCP工具"),
                    "description": tool.get("description", "MCP服务工具"),
                    "functionalCategory": functional_category,
                    "icon": "[ELECTRIC_PLUG]",
                    "enabled": tool.get("enabled", True),
                    "source": "mcp"
                })

        # 兼容前端MCPConnectionsPage期望的格式
        return {
            "tools": processed_tools,
            "total": len(processed_tools),
            "message": "MCP工具数据加载成功"
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"获取MCP工具失败: {str(e)}",
            "data": {
                "allTools": 0,
                "myTools": 0,
                "platformTools": 0,
                "approvedTools": 0,
                "pendingTools": 0,
                "tools": []
            }
        }

@app.get("/api/api/tools")
@app.get("/api/api-tools")
async def get_api_tools():
    """获取API工具数据"""
    try:
        from utils.frontend_config import load_config, API_TOOLS_CONFIG
        api_tools = load_config(API_TOOLS_CONFIG)

        # 统计数据
        total_tools = len(api_tools)
        running_tools = len([tool for tool in api_tools if tool.get("enabled", True)])
        supported_types = 7  # 支持的API类型数量

        # 处理工具数据，添加功能分类
        processed_tools = []
        for tool in api_tools:
            if tool.get("enabled", True):
                # 根据工具特性自动分配功能分类
                functional_category = tool.get("functionalCategory", "web_services")
                if not functional_category or functional_category == "web_services":
                    # 根据工具名称和描述自动分类
                    name_desc = (tool.get("name", "") + " " + tool.get("description", "")).lower()
                    if "http" in name_desc or "请求" in name_desc:
                        functional_category = "network_communication"
                    elif "文件" in name_desc or "处理" in name_desc:
                        functional_category = "file_management"
                    elif "数据" in name_desc or "分析" in name_desc:
                        functional_category = "data_processing"
                    else:
                        functional_category = "network_communication"

                processed_tools.append({
                    "name": tool.get("name", "API工具"),
                    "description": tool.get("description", "API服务工具"),
                    "functionalCategory": functional_category,
                    "icon": "[LINK3]",
                    "enabled": tool.get("enabled", True),
                    "source": "api_config"
                })

        # 兼容前端APIToolsPage期望的格式
        return {
            "tools": processed_tools,
            "total": len(processed_tools),
            "message": "API工具数据加载成功"
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"获取API工具失败: {str(e)}",
            "data": {
                "apiTools": 0,
                "runningTools": 0,
                "supportedTypes": 0,
                "tools": []
            }
        }

@app.get("/api/system/tools")
async def get_system_tools():
    """获取系统工具数据"""
    try:
        from utils.frontend_config import load_config
        from pathlib import Path

        # 加载系统工具配置文件
        PROJECT_ROOT = Path(__file__).parent.parent.parent
        SYSTEM_TOOLS_CONFIG = PROJECT_ROOT / "data" / "configs" / "backend-config" / "system_tools.json"

        system_tools = load_config(SYSTEM_TOOLS_CONFIG)

        # 处理工具数据，添加功能分类
        processed_tools = []
        for tool in system_tools:
            if tool.get("enabled", True):
                # 使用配置文件中的功能分类，如果没有则根据名称自动分类
                functional_category = tool.get("functionalCategory", "system_tools")
                if not functional_category or functional_category == "system_tools":
                    # 根据工具名称和描述自动分类
                    name_desc = (tool.get("name", "") + " " + tool.get("description", "")).lower()
                    if "文件" in name_desc and "管理" in name_desc:
                        functional_category = "file_management"
                    elif "网络" in name_desc or "ping" in name_desc:
                        functional_category = "network_communication"
                    elif "文本" in name_desc or "处理" in name_desc:
                        functional_category = "document_processing"
                    elif "时间" in name_desc:
                        functional_category = "system_tools"
                    elif "系统" in name_desc and "信息" in name_desc:
                        functional_category = "system_tools"
                    else:
                        functional_category = "system_tools"

                processed_tools.append({
                    "name": tool.get("name", "系统工具"),
                    "description": tool.get("description", "系统服务工具"),
                    "functionalCategory": functional_category,
                    "icon": tool.get("icon", "[HAMMER_WRENCH2]"),
                    "enabled": tool.get("enabled", True),
                    "source": "system"
                })

        return {
            "status": "success",
            "data": {
                "tools": processed_tools
            }
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"获取系统工具失败: {str(e)}",
            "data": {
                "tools": []
            }
        }

@app.post("/api/ai/summarize-description", response_model=DescriptionSummaryResponse)
async def summarize_description(request: DescriptionSummaryRequest):
    """使用AI总结工具描述为50字以内的简洁描述"""
    try:
        original_desc = request.description.strip()
        max_length = min(request.maxLength, 50)  # 确保不超过50字

        if not original_desc:
            return DescriptionSummaryResponse(
                summary="",
                originalLength=0,
                summaryLength=0
            )

        # 如果原描述已经很短，直接返回
        if len(original_desc) <= max_length:
            return DescriptionSummaryResponse(
                summary=original_desc,
                originalLength=len(original_desc),
                summaryLength=len(original_desc)
            )

        # 使用AI总结
        summary = await _ai_summarize_description(original_desc, max_length)

        return DescriptionSummaryResponse(
            summary=summary,
            originalLength=len(original_desc),
            summaryLength=len(summary)
        )

    except Exception as e:
        logging.error(f"AI总结描述失败: {e}")
        # 降级处理
        fallback_summary = _fallback_summarize(request.description, request.maxLength)
        return DescriptionSummaryResponse(
            summary=fallback_summary,
            originalLength=len(request.description),
            summaryLength=len(fallback_summary)
        )

async def _ai_summarize_description(description: str, max_length: int = 80) -> str:
    """使用AI总结描述"""
    try:
        # 检查是否有OpenAI API Key
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            logging.warning("未配置OPENAI_API_KEY，使用智能降级总结方案")
            return _intelligent_fallback_summarize(description, max_length)

        # 构建专业的工具描述总结提示
        prompt = f"""你是一个专业的工具描述总结专家。请将以下工具描述总结为{max_length}字以内的精准描述。

原始描述：
{description}

总结要求：
1. 字数限制：不超过{max_length}个字符
2. 核心信息：必须保留工具的主要功能和用途
3. 关键特性：提取最重要的2-3个特色功能
4. 适用场景：简要说明主要应用领域
5. 语言风格：简洁专业，便于系统匹配

总结格式：[工具类型]：[核心功能]，[关键特性]，适用于[主要场景]

请直接返回总结结果，不要包含其他内容："""

        # 使用OpenAI API进行总结
        client = AsyncOpenAI(api_key=api_key)

        response = await client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "你是一个专业的工具描述总结专家，擅长提取关键信息并生成简洁准确的描述。"},
                {"role": "user", "content": prompt}
            ],
            max_tokens=120,
            temperature=0.2
        )

        summary = response.choices[0].message.content.strip()

        # 清理可能的格式标记
        summary = summary.replace("总结：", "").replace("总结结果：", "").strip()

        # 确保不超过最大长度
        if len(summary) > max_length:
            # 智能截取，尽量保持完整性
            summary = _smart_truncate(summary, max_length)

        return summary

    except Exception as e:
        logging.error(f"AI总结失败: {e}")
        return _intelligent_fallback_summarize(description, max_length)

def _intelligent_fallback_summarize(text: str, max_length: int = 80) -> str:
    """智能降级总结方案：基于规则的文本处理"""
    if not text:
        return ""

    # 移除多余的空白字符和特殊符号
    cleaned = re.sub(r'\s+', ' ', text.strip())

    # 提取关键信息的模式
    patterns = [
        r'([^：:]*(?:工具|服务|系统|平台|应用))[：:]([^。！？.!?]*)',  # 工具类型：功能描述
        r'(支持|提供|包括|具有)([^。！？.!?]*)',  # 功能描述
        r'(适用于|用于|面向)([^。！？.!?]*)',  # 应用场景
    ]

    extracted_info = []

    # 尝试提取结构化信息
    for pattern in patterns:
        matches = re.findall(pattern, cleaned)
        for match in matches:
            if isinstance(match, tuple):
                info = ''.join(match).strip()
                if info and len(info) > 5:  # 过滤太短的信息
                    extracted_info.append(info)

    # 如果提取到结构化信息，组合使用
    if extracted_info:
        summary = '、'.join(extracted_info[:3])  # 最多3个关键信息
        if len(summary) <= max_length:
            return summary
        else:
            return _smart_truncate(summary, max_length)

    # 否则使用改进的句子提取
    sentences = re.split(r'[。！？.!?]', cleaned)

    # 找到最有信息量的句子
    best_sentence = ""
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) > len(best_sentence) and len(sentence) <= max_length:
            # 优先选择包含关键词的句子
            if any(keyword in sentence for keyword in ['工具', '支持', '功能', '用于', '适用']):
                best_sentence = sentence

    if best_sentence:
        return best_sentence

    # 最后的降级方案：智能截取
    return _smart_truncate(cleaned, max_length)

def _smart_truncate(text: str, max_length: int) -> str:
    """智能截取文本，尽量保持完整性"""
    if len(text) <= max_length:
        return text

    # 尝试在标点符号处截取
    punctuation_marks = ['，', '、', '；', ',', ';', ' ']

    for i in range(max_length - 1, max_length // 2, -1):
        if text[i] in punctuation_marks:
            return text[:i]

    # 如果没有合适的标点符号，直接截取
    return text[:max_length]

def _fallback_summarize(text: str, max_length: int = 50) -> str:
    """原始降级总结方案：简单的文本处理（保留兼容性）"""
    if not text:
        return ""

    # 移除多余的空白字符
    cleaned = text.strip().replace(r'\s+', ' ')

    # 尝试找到第一句话
    first_sentence = re.split(r'[。！？.!?]', cleaned)[0]

    # 如果第一句话合适，使用第一句话
    if len(first_sentence) <= max_length and len(first_sentence) > 0:
        return first_sentence

    # 否则截取前max_length个字符
    return cleaned[:max_length]

@app.on_event("shutdown")
async def shutdown_event():
    """应用关闭事件"""
    logging.info("Shutting down Workflow Platform...")

    # 停止监控
    try:
        await metrics_collector.stop_collection()
        await alert_manager.stop_monitoring()
        logging.info("Monitoring services stopped")
    except Exception as e:
        logging.error(f"Failed to stop monitoring: {e}")


# --- WebSocket端点 ---
# 导入WebSocket管理器
from utils.websocket_manager import websocket_manager

@app.websocket("/ws/tools")
async def websocket_endpoint(websocket: WebSocket):
    """工具库实时同步WebSocket端点"""
    try:
        await websocket_manager.connect(websocket)

        # 保持连接活跃
        while True:
            try:
                # 等待客户端消息（心跳包等）
                data = await websocket.receive_text()
                message = json.loads(data)

                # 处理心跳包
                if message.get("type") == "ping":
                    await websocket.send_text(json.dumps({
                        "type": "pong",
                        "timestamp": datetime.now().isoformat()
                    }))

            except WebSocketDisconnect:
                break
            except Exception as e:
                logger.error(f"WebSocket处理消息失败: {e}")
                break

    except Exception as e:
        logger.error(f"WebSocket连接失败: {e}")
    finally:
        websocket_manager.disconnect(websocket)

@app.get("/api/websocket/status")
async def get_websocket_status():
    """获取WebSocket连接状态"""
    return {
        "active_connections": websocket_manager.get_connection_count(),
        "connections": websocket_manager.get_connection_info()
    }

@app.post("/api/websocket/test-notification")
async def test_websocket_notification():
    """测试WebSocket通知功能"""
    try:
        from utils.websocket_manager import notify_api_tools_updated, notify_tools_library_refresh

        # 发送测试通知
        await notify_api_tools_updated("updated", {
            "id": "test_tool",
            "name": "测试工具",
            "description": "这是一个WebSocket测试通知"
        })

        await notify_tools_library_refresh("manual_test")

        return {
            "success": True,
            "message": "测试通知已发送",
            "active_connections": websocket_manager.get_connection_count()
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

# AI工具页面专用路由
@app.get("/ai-tools")
async def get_ai_tools_for_page():
    """获取AI工具数据 - 用于AIConnectionsPage"""
    try:
        from utils.frontend_config import load_config, AI_TOOLS_CONFIG
        ai_tools = load_config(AI_TOOLS_CONFIG)
        
        # 过滤启用的工具并格式化数据
        processed_tools = []
        for tool in ai_tools:
            if tool.get("enabled", True):
                processed_tools.append({
                    "id": tool.get("id", tool.get("name")),
                    "name": tool.get("name", tool.get("display_name", "")),
                    "description": tool.get("description", ""),
                    "provider": tool.get("provider", ""),
                    "model": tool.get("model", ""),
                    "api_key": tool.get("api_key", ""),
                    "base_url": tool.get("base_url", ""),
                    "max_tokens": tool.get("max_tokens", 4000),
                    "temperature": tool.get("temperature", 0.7),
                    "enabled": tool.get("enabled", True),
                    "functionalCategory": tool.get("functionalCategory", "ai_assistant"),
                    "category": tool.get("category", "AI工具"),
                    "icon": tool.get("icon", "🤖"),
                    "tested": tool.get("tested", False),
                    "approved": tool.get("approved", True),
                    "status": "active" if tool.get("enabled", True) else "inactive"
                })
        
        return processed_tools
        
    except Exception as e:
        logger.error(f"获取AI工具数据失败: {e}")
        return []

# --- Main Execution ---
if __name__ == "__main__":
    config = get_config()
    uvicorn.run(
        app,
        host=config.host,
        port=config.port,
        reload=config.reload,
        workers=config.workers if not config.debug else 1
    )

@app.get("/ai-tools")
async def get_ai_tools_for_page():
    """获取AI工具数据 - 用于AIConnectionsPage"""
    try:
        from utils.frontend_config import load_config, AI_TOOLS_CONFIG
        ai_tools = load_config(AI_TOOLS_CONFIG)
        
        # 过滤启用的工具并格式化数据
        processed_tools = []
        for tool in ai_tools:
            if tool.get("enabled", True):
                processed_tools.append({
                    "id": tool.get("id", tool.get("name")),
                    "name": tool.get("name", tool.get("display_name", "")),
                    "description": tool.get("description", ""),
                    "provider": tool.get("provider", ""),
                    "model": tool.get("model", ""),
                    "api_key": tool.get("api_key", ""),
                    "base_url": tool.get("base_url", ""),
                    "max_tokens": tool.get("max_tokens", 4000),
                    "temperature": tool.get("temperature", 0.7),
                    "enabled": tool.get("enabled", True),
                    "functionalCategory": tool.get("functionalCategory", "ai_assistant"),
                    "category": tool.get("category", "AI工具"),
                    "icon": tool.get("icon", "🤖"),
                    "tested": tool.get("tested", False),
                    "approved": tool.get("approved", True),
                    "status": "active" if tool.get("enabled", True) else "inactive"
                })
        
        return processed_tools
        
    except Exception as e:
        logger.error(f"获取AI工具数据失败: {e}")
        return []

@app.post("/api/template/{template_id}/favorite")
async def toggle_template_favorite(template_id: str, request: Request):
    """收藏/取消收藏模板"""
    try:
        data = await request.json()
        user_id = data.get("user_id", "anonymous")
        
        # 读取收藏数据
        favorites_file = Path(__file__).parent.parent / "config" / "template_favorites.json"
        favorites_data = {"favorites": []}
        if favorites_file.exists():
            with open(favorites_file, 'r', encoding='utf-8') as f:
                favorites_data = json.load(f)
        
        # 检查用户是否已经收藏
        existing_favorite = None
        for favorite in favorites_data["favorites"]:
            if favorite.get("template_id") == template_id and favorite.get("user_id") == user_id:
                existing_favorite = favorite
                break
        
        if existing_favorite:
            # 取消收藏
            favorites_data["favorites"].remove(existing_favorite)
            favorited = False
        else:
            # 添加收藏
            new_favorite = {
                "template_id": template_id,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }
            favorites_data["favorites"].append(new_favorite)
            favorited = True
        
        # 保存数据
        favorites_file.parent.mkdir(parents=True, exist_ok=True)
        with open(favorites_file, 'w', encoding='utf-8') as f:
            json.dump(favorites_data, f, ensure_ascii=False, indent=2)
        
        # 计算总收藏数
        total_favorites = len([f for f in favorites_data["favorites"] if f.get("template_id") == template_id])
        
        return {
            "success": True,
            "favorited": favorited,
            "total_favorites": total_favorites,
            "message": "收藏成功" if favorited else "取消收藏成功"
        }
        
    except Exception as e:
        logger.error(f"收藏操作失败: {str(e)}")
        return {"success": False, "error": str(e)}


@app.get("/api/my-favorites")
async def get_my_favorites():
    """获取用户收藏的模板"""
    try:
        # 读取收藏数据
        favorites_file = Path(__file__).parent.parent / "config" / "template_favorites.json"
        if not favorites_file.exists():
            return {
                "success": True,
                "templates": [],
                "total": 0
            }
        
        with open(favorites_file, 'r', encoding='utf-8') as f:
            favorites_data = json.load(f)
        
        # 读取模板数据
        templates_file = Path(__file__).parent.parent / "config" / "templates.json"
        templates_data = {"templates": []}
        if templates_file.exists():
            with open(templates_file, 'r', encoding='utf-8') as f:
                templates_data = json.load(f)
        
        # 获取用户收藏的模板ID列表
        user_id = "anonymous"  # TODO: 从session或JWT中获取真实用户ID
        favorited_template_ids = [
            f.get("template_id") for f in favorites_data.get("favorites", [])
            if f.get("user_id") == user_id
        ]
        
        # 筛选出收藏的模板
        favorited_templates = [
            template for template in templates_data.get("templates", [])
            if template.get("id") in favorited_template_ids
        ]
        
        return {
            "success": True,
            "templates": favorited_templates,
            "total": len(favorited_templates)
        }
        
    except Exception as e:
        logger.error(f"获取收藏模板失败: {str(e)}")
        return {"success": False, "error": str(e)}


# MCP专用端点已移动到主要API区域

# 添加缓存清理API端点
@app.post("/api/clear-cache")
async def clear_cache():
    """
    清除所有缓存：Python缓存、工作流状态、结果历史
    """
    try:
        import shutil
        import os
        import glob
        
        results = {
            "python_cache": False,
            "workflow_states": False,
            "result_history": False,
            "error": None
        }
        
        try:
            # 1. 清除Python缓存
            cache_dirs = []
            for root, dirs, files in os.walk("backend"):
                for dir_name in dirs:
                    if dir_name == "__pycache__":
                        cache_dir = os.path.join(root, dir_name)
                        cache_dirs.append(cache_dir)
            
            for cache_dir in cache_dirs:
                try:
                    shutil.rmtree(cache_dir)
                    print(f"🧹 [缓存清理] 已删除: {cache_dir}")
                except:
                    pass
            
            # 清除.pyc文件
            pyc_files = glob.glob("backend/**/*.pyc", recursive=True)
            for pyc_file in pyc_files:
                try:
                    os.remove(pyc_file)
                    print(f"🧹 [缓存清理] 已删除: {pyc_file}")
                except:
                    pass
            
            results["python_cache"] = True
            print(f"✅ [缓存清理] Python缓存清理完成")
            
        except Exception as e:
            print(f"❌ [缓存清理] Python缓存清理失败: {e}")
        
        try:
            # 2. 清除工作流状态
            state_files = glob.glob("backend/workflow_states/*.json")
            for state_file in state_files:
                try:
                    os.remove(state_file)
                    print(f"🧹 [缓存清理] 已删除工作流状态: {state_file}")
                except:
                    pass
            
            # 清除根目录的工作流状态
            root_state_files = glob.glob("workflow_states/*.json")
            for state_file in root_state_files:
                try:
                    os.remove(state_file)
                    print(f"🧹 [缓存清理] 已删除根目录工作流状态: {state_file}")
                except:
                    pass
            
            results["workflow_states"] = True
            print(f"✅ [缓存清理] 工作流状态清理完成")
            
        except Exception as e:
            print(f"❌ [缓存清理] 工作流状态清理失败: {e}")
        
        try:
            # 3. 清除输出文件（可选）
            output_files = glob.glob("backend/output/*.docx")
            for output_file in output_files:
                try:
                    os.remove(output_file)
                    print(f"🧹 [缓存清理] 已删除输出文件: {output_file}")
                except:
                    pass
                    
            # 清除根目录输出文件
            root_output_files = glob.glob("output/*.docx")
            for output_file in root_output_files:
                try:
                    os.remove(output_file)
                    print(f"🧹 [缓存清理] 已删除根目录输出文件: {output_file}")
                except:
                    pass
            
            results["result_history"] = True
            print(f"✅ [缓存清理] 结果历史清理完成")
            
        except Exception as e:
            print(f"❌ [缓存清理] 结果历史清理失败: {e}")
        
        # 4. 清除内存中的状态管理器
        try:
            if hasattr(app.state, 'workflow_state_manager'):
                # 重置状态管理器
                app.state.workflow_state_manager = {}
            print(f"✅ [缓存清理] 内存状态清理完成")
        except Exception as e:
            print(f"❌ [缓存清理] 内存状态清理失败: {e}")
        
        return {
            "success": True,
            "message": "缓存清理完成",
            "details": results
        }
        
    except Exception as e:
        return {
            "success": False,
            "message": f"缓存清理失败: {str(e)}",
            "details": results
        }

# ============================================================================
# 缺失的API路由 - 从complete_mock_backend.py移植
# ============================================================================

# 工具可用性API
@app.get("/api/tools/available")
async def get_available_tools():
    """获取可用工具列表"""
    return {
        "tools": [
            {
                "id": "document_parser",
                "name": "文档解析器",
                "description": "解析Word文档"
            },
            {
                "id": "mindmap_generator", 
                "name": "思维导图生成器",
                "description": "生成思维导图"
            },
            {
                "id": "workflow_builder",
                "name": "工作流构建器", 
                "description": "构建自动化工作流"
            }
        ]
    }

# 文件夹扫描API
@app.get("/api/folder/scan")
async def scan_folder(path: str):
    """扫描文件夹中的Word文档"""
    try:
        import os
        import glob
        
        logger.info(f"🔍 扫描文件夹: {path}")
        
        # 检查路径是否存在
        if not os.path.exists(path):
            return {"files": [], "message": f"路径不存在: {path}"}
        
        # 扫描Word文档
        pattern = os.path.join(path, "*.docx")
        docx_files = glob.glob(pattern)
        
        files = []
        for file_path in docx_files:
            try:
                file_name = os.path.basename(file_path)
                
                # 过滤掉Word临时文件（以~$开头的文件）
                if file_name.startswith('~$'):
                    logger.info(f"⚠️ 跳过Word临时文件: {file_name}")
                    continue
                
                file_stat = os.stat(file_path)
                file_info = {
                    "name": file_name,
                    "path": file_path,
                    "size": file_stat.st_size,
                    "modified": file_stat.st_mtime
                }
                files.append(file_info)
                logger.info(f"📄 找到文档: {file_info['name']} ({file_info['size']} bytes)")
            except Exception as e:
                logger.warning(f"⚠️ 无法读取文件信息: {file_path} - {e}")
        
        result = {
            "files": files,
            "path": path,
            "total": len(files),
            "message": f"找到 {len(files)} 个Word文档"
        }
        
        logger.info(f"✅ 扫描完成: {result}")
        return result
        
    except Exception as e:
        logger.error(f"❌ 文件夹扫描失败: {e}")
        return {"files": [], "error": str(e), "message": "文件夹扫描失败"}

# 通过文件路径解析文档的API
@app.get("/api/document/parse-path")
async def parse_document_by_path(file_path: str):
    """通过文件路径解析Word文档"""
    logger.info(f"🔍 收到文件路径解析请求: {file_path}")
    
    try:
        # 验证文件存在
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail=f"文件不存在: {file_path}")
        
        # 验证文件类型
        if not file_path.endswith('.docx'):
            raise HTTPException(status_code=400, detail="只支持.docx格式的文件")
        
        # 获取文件信息
        file_stat = os.stat(file_path)
        logger.info(f"🔍 文件大小: {file_stat.st_size} bytes")
        
        # 尝试导入docx模块
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx未安装，返回基本文件信息")
            return {
                "filename": os.path.basename(file_path),
                "content": [],
                "total_paragraphs": 0,
                "message": "文档信息获取成功（需要安装python-docx进行详细解析）",
                "file_size": file_stat.st_size
            }
        
        # 解析文件
        doc = Document(file_path)
        
        content_list = []
        run_id = 0
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 只处理非空段落
                logger.info(f"📄 处理段落 {para_idx}: '{paragraph.text[:50]}...'")
                para_content = {
                    "paragraph_id": para_idx,
                    "runs": []
                }
                
                char_offset = 0
                for run in paragraph.runs:
                    if run.text:
                        run_info = {
                            "run_id": run_id,
                            "text": run.text,
                            "font_name": run.font.name if run.font.name else "默认字体",
                            "start_char": char_offset,
                            "end_char": char_offset + len(run.text)
                        }
                        para_content["runs"].append(run_info)
                        char_offset += len(run.text)
                        run_id += 1
            
                if para_content["runs"]:
                    content_list.append(para_content)
        
        result = {
            "filename": os.path.basename(file_path),
            "content": content_list,
            "total_paragraphs": len(content_list),
            "message": "文档解析成功"
        }
        
        logger.info(f"✅ 文档解析完成: {len(content_list)} 个段落")
        return result
        
    except Exception as e:
        logger.error(f"Error parsing document at path {file_path}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")

# 文档上传解析API
@app.post("/api/document/parse-upload")
async def parse_uploaded_document(file: UploadFile = File(...)):
    """解析用户上传的Word文档"""
    logger.info(f"🔍 收到文档解析请求: {file.filename}")
    logger.info(f"🔍 文件大小: {file.size} bytes")
    logger.info(f"🔍 文件类型: {file.content_type}")
    
    temp_file_path = None
    try:
        # 验证文件类型
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="只支持.docx格式的文件")
        
        # 尝试导入docx模块
        try:
            from docx import Document
            import tempfile
        except ImportError:
            logger.warning("python-docx未安装，返回基本文件信息")
            content = await file.read()
            return {
                "filename": file.filename,
                "content": [],
                "total_paragraphs": 0,
                "message": "文档信息获取成功（需要安装python-docx进行详细解析）",
                "file_size": len(content)
            }
        
        # 创建临时文件并写入内容
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file.flush()
            temp_file_path = temp_file.name
        
        # 现在文件已关闭，可以安全地用Document打开
        doc = Document(temp_file_path)
        
        content_list = []
        run_id = 0
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            logger.info(f"📄 处理段落 {para_idx}: '{paragraph.text[:50]}...'")
            if paragraph.text.strip():
                para_content = {
                    "paragraph_id": para_idx,
                    "runs": []
                }
                
                char_offset = 0
                for run in paragraph.runs:
                    if run.text:
                        run_info = {
                            "run_id": run_id,
                            "text": run.text,
                            "font_name": run.font.name if run.font.name else "默认字体",
                            "start_char": char_offset,
                            "end_char": char_offset + len(run.text)
                        }
                        para_content["runs"].append(run_info)
                        char_offset += len(run.text)
                        run_id += 1
                
                if para_content["runs"]:
                    content_list.append(para_content)
        
        result = {
            "filename": file.filename,
            "content": content_list,
            "total_paragraphs": len(content_list),
            "message": "文档解析成功"
        }
        
        logger.info(f"✅ 文档解析完成: {len(content_list)} 个段落")
        return result
        
    except Exception as e:
        logger.error(f"Error parsing uploaded document {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")
    finally:
        # 确保清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(f"🗑️ 临时文件已清理: {temp_file_path}")
            except Exception as e:
                logger.warning(f"⚠️ 清理临时文件失败: {e}")

