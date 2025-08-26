import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

class EnhancedWordTools:
    """增强的Word文档工具类 - 实现完整的Word MCP功能"""
    
    def __init__(self):
        self.current_document = None
        self.current_file_path = None
        
    # ==================== 文档管理 ====================
    
    def create_document(self, filename: str, title: str = None, output_dir: str = None) -> Dict[str, Any]:
        """创建新的Word文档"""
        try:
            self.current_document = Document()
            
            # 设置页面格式
            section = self.current_document.sections[0]
            section.page_height = Inches(11.69)  # A4纸高度
            section.page_width = Inches(8.27)    # A4纸宽度
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # 添加标题（如果提供）
            if title:
                heading = self.current_document.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 确定保存路径
            if not output_dir:
                output_dir = os.path.expanduser("~/Desktop/AI作品")
            
            os.makedirs(output_dir, exist_ok=True)
            
            # 生成文件路径
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            self.current_file_path = os.path.join(output_dir, filename)
            
            # 立即保存文档
            self.current_document.save(self.current_file_path)
            
            return {
                "success": True,
                "message": f"Word文档创建成功: {filename}",
                "document_path": self.current_file_path,
                "filename": filename,
                "output_dir": output_dir
            }
            
        except Exception as e:
            logger.error(f"创建Word文档失败: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def open_document(self, file_path: str) -> Dict[str, Any]:
        """打开现有文档"""
        try:
            if not os.path.exists(file_path):
                return {"success": False, "error": "文档文件不存在"}
            
            self.current_document = Document(file_path)
            self.current_file_path = file_path
            
            return {
                "success": True,
                "message": f"文档打开成功: {file_path}",
                "document_path": file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def save_document(self, file_path: str = None) -> Dict[str, Any]:
        """保存文档"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            save_path = file_path or self.current_file_path
            if not save_path:
                return {"success": False, "error": "没有指定保存路径"}
            
            self.current_document.save(save_path)
            self.current_file_path = save_path
            
            return {
                "success": True,
                "message": f"文档保存成功",
                "file_path": save_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    # ==================== 内容创作 ====================
    
    def add_paragraph(self, text: str, style: str = None, font_name: str = None, 
                     font_size: int = None, bold: bool = False, italic: bool = False,
                     alignment: str = None, color: str = None) -> Dict[str, Any]:
        """添加正文段落"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            paragraph = self.current_document.add_paragraph(text)
            
            # 应用样式
            if style:
                paragraph.style = style
            
            # 设置字体
            if font_name or font_size or bold or italic or color:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                font = run.font
                
                if font_name:
                    font.name = font_name
                if font_size:
                    font.size = Pt(font_size)
                if bold:
                    font.bold = True
                if italic:
                    font.italic = True
                if color:
                    font.color.rgb = self._parse_color(color)
            
            # 设置对齐
            if alignment:
                if alignment.lower() in ['center', '居中']:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment.lower() in ['right', '右对齐']:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment.lower() in ['justify', '两端对齐']:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            return {
                "success": True,
                "message": "段落添加成功",
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def add_heading(self, text: str, level: int = 1) -> Dict[str, Any]:
        """添加标题"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            self.current_document.add_heading(text, level)
            
            return {
                "success": True,
                "message": f"标题添加成功: {text}",
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def add_table(self, rows: int, cols: int, data: List[List[str]] = None) -> Dict[str, Any]:
        """添加表格"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            table = self.current_document.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 填充数据
            if data:
                for i, row_data in enumerate(data):
                    if i < len(table.rows):
                        for j, cell_data in enumerate(row_data):
                            if j < len(table.columns):
                                table.cell(i, j).text = str(cell_data)
            
            return {
                "success": True,
                "message": f"表格添加成功: {rows}x{cols}",
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def add_page_break(self) -> Dict[str, Any]:
        """添加分页符"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            self.current_document.add_page_break()
            
            return {
                "success": True,
                "message": "分页符添加成功",
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    # ==================== 富文本格式 ====================
    
    def format_text(self, text: str, font_name: str = None, font_size: int = None,
                   bold: bool = False, italic: bool = False, underline: bool = False,
                   color: str = None) -> Dict[str, Any]:
        """格式化文本"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            paragraph = self.current_document.add_paragraph()
            run = paragraph.add_run(text)
            font = run.font
            
            if font_name:
                font.name = font_name
            if font_size:
                font.size = Pt(font_size)
            if bold:
                font.bold = True
            if italic:
                font.italic = True
            if underline:
                font.underline = True
            if color:
                font.color.rgb = self._parse_color(color)
            
            return {
                "success": True,
                "message": "文本格式化成功",
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def search_and_replace(self, search_text: str, replace_text: str) -> Dict[str, Any]:
        """搜索和替换文本"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            replacements = 0
            for paragraph in self.current_document.paragraphs:
                if search_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(search_text, replace_text)
                    replacements += 1
            
            return {
                "success": True,
                "message": f"替换完成，共替换 {replacements} 处",
                "replacements": replacements,
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    # ==================== 文档信息 ====================
    
    def get_document_info(self) -> Dict[str, Any]:
        """获取文档信息"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            info = {
                "success": True,
                "document_path": self.current_file_path,
                "paragraphs_count": len(self.current_document.paragraphs),
                "tables_count": len(self.current_document.tables),
                "sections_count": len(self.current_document.sections)
            }
            
            # 获取核心属性
            if hasattr(self.current_document, 'core_properties'):
                core_props = self.current_document.core_properties
                info.update({
                    "title": core_props.title or "未设置",
                    "author": core_props.author or "未设置",
                    "created": str(core_props.created) if core_props.created else "未知",
                    "modified": str(core_props.modified) if core_props.modified else "未知"
                })
            
            return info
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def extract_text(self) -> Dict[str, Any]:
        """提取文档所有文本"""
        try:
            if not self.current_document:
                return {"success": False, "error": "没有活动的文档"}
            
            all_text = []
            for paragraph in self.current_document.paragraphs:
                all_text.append(paragraph.text)
            
            return {
                "success": True,
                "text": "\n".join(all_text),
                "paragraphs": len(all_text),
                "document_path": self.current_file_path
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    # ==================== 辅助方法 ====================
    
    def _parse_color(self, color_str: str) -> RGBColor:
        """解析颜色字符串"""
        color_map = {
            "红色": RGBColor(255, 0, 0),
            "蓝色": RGBColor(0, 0, 255),
            "绿色": RGBColor(0, 128, 0),
            "黄色": RGBColor(255, 255, 0),
            "黑色": RGBColor(0, 0, 0),
            "白色": RGBColor(255, 255, 255),
            "灰色": RGBColor(128, 128, 128)
        }
        
        if color_str in color_map:
            return color_map[color_str]
        
        # 尝试解析十六进制颜色
        if color_str.startswith('#') and len(color_str) == 7:
            try:
                r = int(color_str[1:3], 16)
                g = int(color_str[3:5], 16)
                b = int(color_str[5:7], 16)
                return RGBColor(r, g, b)
            except ValueError:
                pass
        
        # 默认黑色
        return RGBColor(0, 0, 0)

# 创建全局实例
enhanced_word_tools = EnhancedWordTools()

# 全局函数接口
def create_document(filename: str, title: str = None, output_dir: str = None) -> Dict[str, Any]:
    return enhanced_word_tools.create_document(filename, title, output_dir)

def add_paragraph(text: str, **kwargs) -> Dict[str, Any]:
    return enhanced_word_tools.add_paragraph(text, **kwargs)

def add_heading(text: str, level: int = 1) -> Dict[str, Any]:
    return enhanced_word_tools.add_heading(text, level)

def add_table(rows: int, cols: int, data: List[List[str]] = None) -> Dict[str, Any]:
    return enhanced_word_tools.add_table(rows, cols, data)

def format_text(text: str, **kwargs) -> Dict[str, Any]:
    return enhanced_word_tools.format_text(text, **kwargs)

def search_and_replace(search_text: str, replace_text: str) -> Dict[str, Any]:
    return enhanced_word_tools.search_and_replace(search_text, replace_text)

def add_page_break() -> Dict[str, Any]:
    return enhanced_word_tools.add_page_break()

def save_document(file_path: str = None) -> Dict[str, Any]:
    return enhanced_word_tools.save_document(file_path)

def get_document_info() -> Dict[str, Any]:
    return enhanced_word_tools.get_document_info()

def extract_text() -> Dict[str, Any]:
    return enhanced_word_tools.extract_text() 