#!/usr/bin/env python3
"""
完全独立的Word工具模块 - 零依赖MCP服务器
解决所有Windows连接问题的终极方案

Authors: AI Assistant
Date: 2025-08-16
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

# 设置日志
logger = logging.getLogger(__name__)

class SimpleWordTools:
    """完全独立的Word工具类 - 无外部依赖"""
    
    def __init__(self):
        self.output_dir = self._setup_output_dir()
        self.tools_registry = self._register_tools()
        logger.info("✅ SimpleWordTools初始化成功 - 无外部依赖")
    
    def _setup_output_dir(self) -> Path:
        """设置输出目录"""
        # 使用用户期望的AI作品目录
        output_dir = Path("C:/Users/ZhuanZ/Desktop/AI作品")
        output_dir.mkdir(exist_ok=True)
        return output_dir
    
    def _register_tools(self) -> Dict[str, Dict]:
        """注册所有可用工具"""
        return {
            "create_document": {
                "description": "创建新的Word文档",
                "parameters": {
                    "filename": {"type": "string", "description": "文档文件名"},
                    "title": {"type": "string", "description": "文档标题", "required": False},
                    "content": {"type": "string", "description": "文档内容", "required": False}
                }
            },
            "open_document": {
                "description": "打开现有的Word文档",
                "parameters": {
                    "filename": {"type": "string", "description": "要打开的文档文件路径"}
                }
            },
            "add_paragraph": {
                "description": "添加段落到文档",
                "parameters": {
                    "text": {"type": "string", "description": "段落文本"},
                    "filename": {"type": "string", "description": "目标文档文件名", "required": False}
                }
            },
            "add_heading": {
                "description": "添加标题到文档",
                "parameters": {
                    "text": {"type": "string", "description": "标题文本"},
                    "level": {"type": "integer", "description": "标题级别(1-6)", "required": False},
                    "filename": {"type": "string", "description": "目标文档文件名", "required": False}
                }
            },
            "save_document": {
                "description": "保存文档",
                "parameters": {
                    "filename": {"type": "string", "description": "文档文件名"}
                }
            },
            "set_font_name": {
                "description": "设置文档中指定文本的字体",
                "parameters": {
                    "font_name": {"type": "string", "description": "字体名称（如：华文彩云）"},
                    "filename": {"type": "string", "description": "目标文档文件名"},
                    "target_text": {"type": "string", "description": "要修改字体的目标文本", "required": False}
                }
            }
        }
    
    def get_available_tools(self) -> List[Dict]:
        """获取所有可用工具"""
        tools = []
        for tool_name, tool_info in self.tools_registry.items():
            tools.append({
                "name": tool_name,
                "description": tool_info["description"],
                "parameters": tool_info["parameters"]
            })
        return tools
    
    def execute_tool(self, tool_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """执行工具调用"""
        try:
            logger.info(f"🔧 执行工具: {tool_name}, 参数: {parameters}")
            
            if tool_name not in self.tools_registry:
                return {
                    "success": False,
                    "error": f"未知工具: {tool_name}",
                    "available_tools": list(self.tools_registry.keys())
                }
            
            # 根据工具名称调用对应方法
            if tool_name == "create_document":
                return self._create_document(**parameters)
            elif tool_name == "open_document":
                return self._open_document(**parameters)
            elif tool_name == "add_paragraph":
                return self._add_paragraph(**parameters)
            elif tool_name == "add_heading":
                return self._add_heading(**parameters)
            elif tool_name == "save_document":
                return self._save_document(**parameters)
            elif tool_name == "set_font_name":
                return self._set_font_name(**parameters)
            else:
                return {
                    "success": False,
                    "error": f"工具 {tool_name} 未实现"
                }
                
        except Exception as e:
            logger.error(f"❌ 工具执行失败: {tool_name}, 错误: {e}")
            return {
                "success": False,
                "error": f"工具执行失败: {str(e)}"
            }
    
    def _create_document(self, filename: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        """创建Word文档 - 使用原生Python实现"""
        try:
            # 尝试使用python-docx
            try:
                from docx import Document
                from docx.shared import Inches
                
                doc = Document()
                
                # 添加标题
                if title:
                    heading = doc.add_heading(title, 0)
                
                # 添加内容
                if content:
                    doc.add_paragraph(content)
                elif content == "":
                    # 如果明确传递空字符串，创建空文档
                    pass
                else:
                    # 如果没有传递content参数（None），添加默认内容
                    doc.add_paragraph("这是一个新创建的文档。")
                
                # 保存文档
                if not filename.endswith('.docx'):
                    filename += '.docx'
                
                file_path = self.output_dir / filename
                doc.save(str(file_path))
                
                return {
                    "success": True,
                    "message": f"文档 {filename} 创建成功",
                    "file_path": str(file_path),
                    "created_time": datetime.now().isoformat()
                }
                
            except ImportError:
                # 如果没有python-docx，使用RTF格式作为备选
                return self._create_rtf_document(filename, title, content)
                
        except Exception as e:
            logger.error(f"❌ 创建文档失败: {e}")
            return {
                "success": False,
                "error": f"创建文档失败: {str(e)}"
            }
    
    def _open_document(self, filename: str) -> Dict[str, Any]:
        """打开现有的Word文档"""
        try:
            from pathlib import Path
            
            # 规范化文件路径
            if isinstance(filename, str):
                # 处理可能的双反斜杠路径
                filename = filename.replace('\\\\', '\\')
                file_path = Path(filename)
            else:
                file_path = Path(filename)
            
            # 检查文件是否存在
            if not file_path.exists():
                return {
                    "success": False,
                    "error": f"文件不存在: {file_path}",
                    "file_path": str(file_path)
                }
            
            # 检查文件是否可读
            if not file_path.is_file():
                return {
                    "success": False,
                    "error": f"路径不是文件: {file_path}",
                    "file_path": str(file_path)
                }
            
            # 尝试读取文档内容
            content_info = ""
            try:
                # 如果是.docx文件，尝试读取内容
                if file_path.suffix.lower() == '.docx':
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                        content_info = f"文档包含 {len(paragraphs)} 个段落"
                        if paragraphs:
                            preview = paragraphs[0][:100]
                            content_info += f"，内容预览: {preview}..."
                    except ImportError:
                        content_info = "文档格式: .docx (无法读取内容，缺少python-docx库)"
                elif file_path.suffix.lower() == '.rtf':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    content_info = f"RTF文档，大小: {len(content)} 字符"
                else:
                    content_info = f"文档格式: {file_path.suffix}"
            except Exception as read_error:
                content_info = f"无法读取文档内容: {str(read_error)}"
            
            return {
                "success": True,
                "message": f"成功打开文档: {file_path.name}",
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "content_info": content_info,
                "opened_time": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"❌ 打开文档失败: {e}")
            return {
                "success": False,
                "error": f"打开文档失败: {str(e)}"
            }
    
    def _create_rtf_document(self, filename: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        """创建RTF文档作为备选方案"""
        try:
            if not filename.endswith('.rtf'):
                filename = filename.replace('.docx', '.rtf')
                if not filename.endswith('.rtf'):
                    filename += '.rtf'
            
            rtf_content = r"{\rtf1\ansi\deff0"
            
            # 添加字体表
            rtf_content += r"{\fonttbl{\f0 Times New Roman;}{\f1 SimSun;}}"
            
            # 添加标题
            if title:
                rtf_content += r"\f1\fs24\b " + title + r"\b0\par\par"
            
            # 添加内容
            if content:
                rtf_content += r"\f1\fs20 " + content + r"\par"
            elif content == "":
                # 如果明确传递空字符串，创建空文档
                pass
            else:
                # 如果没有传递content参数（None），添加默认内容
                rtf_content += r"\f1\fs20 这是一个新创建的文档。\par"
            
            rtf_content += "}"
            
            file_path = self.output_dir / filename
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(rtf_content)
            
            return {
                "success": True,
                "message": f"RTF文档 {filename} 创建成功",
                "file_path": str(file_path),
                "created_time": datetime.now().isoformat(),
                "note": "使用RTF格式，可在Word中打开"
            }
            
        except Exception as e:
            # 最后的备选方案：纯文本
            return self._create_text_document(filename, title, content)
    
    def _create_text_document(self, filename: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        """创建文本文档作为最终备选方案"""
        try:
            filename = filename.replace('.docx', '.txt').replace('.rtf', '.txt')
            if not filename.endswith('.txt'):
                filename += '.txt'
            
            text_content = ""
            if title:
                text_content += f"{title}\n{'=' * len(title)}\n\n"
            
            if content:
                text_content += content
            elif content == "":
                # 如果明确传递空字符串，创建空文档
                pass
            else:
                # 如果没有传递content参数（None），添加默认内容
                text_content += "这是一个新创建的文档。"
            
            text_content += f"\n\n创建时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
            file_path = self.output_dir / filename
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return {
                "success": True,
                "message": f"文本文档 {filename} 创建成功",
                "file_path": str(file_path),
                "created_time": datetime.now().isoformat(),
                "note": "使用文本格式，兼容性最强"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"创建文档失败: {str(e)}"
            }
    
    def _add_paragraph(self, text: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """添加段落到文档"""
        try:
            if not filename:
                filename = "default_document.txt"
            
            file_path = self.output_dir / filename
            
            # 如果文件不存在，先创建
            if not file_path.exists():
                create_result = self._create_document(filename)
                if not create_result["success"]:
                    return create_result
            
            # 追加内容
            with open(file_path, 'a', encoding='utf-8') as f:
                f.write(f"\n\n{text}")
            
            return {
                "success": True,
                "message": f"段落已添加到 {filename}",
                "added_text": text
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"添加段落失败: {str(e)}"
            }
    
    def _add_heading(self, text: str, level: int = 1, filename: Optional[str] = None) -> Dict[str, Any]:
        """添加标题到文档"""
        try:
            if not filename:
                filename = "default_document.txt"
            
            file_path = self.output_dir / filename
            
            # 如果文件不存在，先创建
            if not file_path.exists():
                create_result = self._create_document(filename)
                if not create_result["success"]:
                    return create_result
            
            # 根据级别生成标题格式
            heading_prefix = "#" * min(level, 6) + " "
            heading_text = f"\n\n{heading_prefix}{text}\n"
            
            # 追加标题
            with open(file_path, 'a', encoding='utf-8') as f:
                f.write(heading_text)
            
            return {
                "success": True,
                "message": f"标题已添加到 {filename}",
                "heading_text": text,
                "level": level
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"添加标题失败: {str(e)}"
            }
    
    def _save_document(self, filename: str) -> Dict[str, Any]:
        """保存文档"""
        try:
            file_path = self.output_dir / filename
            
            if file_path.exists():
                return {
                    "success": True,
                    "message": f"文档 {filename} 已保存",
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size
                }
            else:
                # 如果文件不存在，创建一个空文档
                create_result = self._create_document(filename)
                return create_result
                
        except Exception as e:
            return {
                "success": False,
                "error": f"保存文档失败: {str(e)}"
            }

    def _set_font_name(self, font_name: str, filename: str, target_text: Optional[str] = None) -> Dict[str, Any]:
        """设置文档中指定文本的字体"""
        try:
            from docx import Document
            from docx.shared import RGBColor
            from docx.oxml.ns import qn
            
            file_path = Path(filename)
            if not file_path.is_absolute():
                file_path = self.output_dir / file_path
            
            if not file_path.exists():
                return {
                    "success": False,
                    "error": f"文档文件不存在: {file_path}"
                }
            
            # 打开现有文档
            doc = Document(str(file_path))
            font_changed_count = 0
            
            # 遍历所有段落
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # 如果指定了目标文本，只修改包含该文本的run
                    if target_text:
                        if target_text in run.text:
                            # 设置字体，特别处理中文字体
                            run.font.name = font_name
                            # 确保rPr元素存在
                            if run._element.rPr is None:
                                run._element.get_or_add_rPr()
                            # 确保rFonts元素存在  
                            if run._element.rPr.rFonts is None:
                                from docx.oxml import OxmlElement
                                fonts_element = OxmlElement('w:rFonts')
                                run._element.rPr.append(fonts_element)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            font_changed_count += 1
                    else:
                        # 如果没有指定目标文本，修改所有文本的字体
                        run.font.name = font_name
                        # 确保rPr元素存在
                        if run._element.rPr is None:
                            run._element.get_or_add_rPr()
                        # 确保rFonts元素存在  
                        if run._element.rPr.rFonts is None:
                            from docx.oxml import OxmlElement
                            fonts_element = OxmlElement('w:rFonts')
                            run._element.rPr.append(fonts_element)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        font_changed_count += 1
            
            # 保存文档
            doc.save(str(file_path))
            
            if target_text:
                message = f"成功设置字体 '{font_name}' 到包含 '{target_text}' 的文本，共修改 {font_changed_count} 处"
            else:
                message = f"成功设置字体 '{font_name}' 到整个文档，共修改 {font_changed_count} 处"
            
            logger.info(f"✅ {message}")
            
            return {
                "success": True,
                "message": message,
                "font_name": font_name,
                "target_text": target_text,
                "changes_count": font_changed_count,
                "file_path": str(file_path)
            }
            
        except Exception as e:
            error_msg = f"设置字体失败: {str(e)}"
            logger.error(f"❌ {error_msg}")
            return {
                "success": False,
                "error": error_msg
            }

# 全局实例
_simple_word_tools = None

def get_simple_word_tools() -> SimpleWordTools:
    """获取SimpleWordTools实例（单例模式）"""
    global _simple_word_tools
    if _simple_word_tools is None:
        _simple_word_tools = SimpleWordTools()
    return _simple_word_tools

def execute_word_tool(tool_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
    """执行Word工具（外部接口）"""
    tools = get_simple_word_tools()
    return tools.execute_tool(tool_name, parameters)

def get_word_tools() -> List[Dict]:
    """获取所有Word工具（外部接口）"""
    tools = get_simple_word_tools()
    return tools.get_available_tools()

if __name__ == "__main__":
    # 测试代码
    tools = SimpleWordTools()
    
    print("🧪 测试SimpleWordTools...")
    
    # 测试创建文档
    result = tools.execute_tool("create_document", {
        "filename": "test_simple.docx",
        "title": "测试文档",
        "content": "这是一个测试文档，使用SimpleWordTools创建。"
    })
    print(f"创建文档结果: {result}")
    
    # 测试添加段落
    result = tools.execute_tool("add_paragraph", {
        "text": "这是一个新添加的段落。",
        "filename": "test_simple.docx"
    })
    print(f"添加段落结果: {result}")
    
    print("✅ 测试完成") 