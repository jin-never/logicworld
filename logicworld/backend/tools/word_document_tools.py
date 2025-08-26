import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

class WordDocumentTools:
    """基于 python-docx 的 Word 文档工具类"""
    
    def __init__(self):
        self.current_document = None
        self.current_file_path = None
        
    def create_document(self, filename: str, title: str = None, output_dir: str = None) -> Dict[str, Any]:
        """
        创建新的 Word 文档
        
        Args:
            filename: 文件名（不含扩展名）
            title: 文档标题
            output_dir: 输出目录，默认为用户桌面
        
        Returns:
            包含操作结果的字典
        """
        try:
            # 创建新文档
            self.current_document = Document()
            
            # 设置页面格式
            section = self.current_document.sections[0]
            section.page_height = Inches(11.69)  # A4纸高度
            section.page_width = Inches(8.27)    # A4纸宽度
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # 添加标题（如果提供）
            if title:
                heading = self.current_document.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 确定保存路径
            if not output_dir:
                # 默认保存到用户桌面的AI作品文件夹
                desktop = os.path.join(os.path.expanduser("~"), "Desktop")
                output_dir = os.path.join(desktop, "AI作品")
            
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
            # 完整的文件路径
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            self.current_file_path = os.path.join(output_dir, filename)
            
            # 保存文档
            self.current_document.save(self.current_file_path)
            
            logger.info(f"✅ Word文档创建成功: {self.current_file_path}")
            
            return {
                "success": True,
                "message": f"Word文档创建成功: {filename}",
                "document_path": self.current_file_path,
                "filename": filename,
                "output_dir": output_dir
            }
            
        except Exception as e:
            error_msg = f"创建Word文档失败: {str(e)}"
            logger.error(f"❌ {error_msg}")
            return {
                "success": False,
                "error": error_msg
            }
    
    def add_paragraph(self, text: str, style: str = None, font_name: str = None, 
                     font_size: int = None, bold: bool = False, 
                     alignment: str = None, indent_first_line: float = None,
                     line_spacing: float = None) -> Dict[str, Any]:
        """
        添加正文段落到当前文档
        
        Args:
            text: 正文段落内容
            style: 段落样式
            font_name: 字体名称
            font_size: 字体大小（磅）
            bold: 是否加粗
            alignment: 对齐方式 ('left', 'center', 'right', 'justify')
            indent_first_line: 首行缩进（英寸）
            line_spacing: 行距（倍数）
        
        Returns:
            包含操作结果的字典
        """
        try:
            if not self.current_document:
                return {
                    "success": False,
                    "error": "没有当前文档，请先创建文档"
                }
            
            # 添加段落
            paragraph = self.current_document.add_paragraph(text, style=style)
            
            # 设置字体格式
            if font_name or font_size or bold:
                for run in paragraph.runs:
                    if font_name:
                        run.font.name = font_name
                        # 设置中文字体
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    if font_size:
                        run.font.size = Pt(font_size)
                    if bold:
                        run.font.bold = True
            
            # 设置段落对齐
            if alignment:
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.lower() in alignment_map:
                    paragraph.alignment = alignment_map[alignment.lower()]
            
            # 设置首行缩进
            if indent_first_line:
                paragraph.paragraph_format.first_line_indent = Inches(indent_first_line)
            
            # 设置行距
            if line_spacing:
                paragraph.paragraph_format.line_spacing = line_spacing
            
            # 保存文档
            if self.current_file_path:
                self.current_document.save(self.current_file_path)
            
            logger.info(f"✅ 段落添加成功: {text[:50]}...")
            
            return {
                "success": True,
                "message": f"段落添加成功",
                "document_path": self.current_file_path
            }
            
        except Exception as e:
            error_msg = f"添加正文段落失败: {str(e)}"
            logger.error(f"❌ {error_msg}")
            return {
                "success": False,
                "error": error_msg
            }
    
    def set_page_size(self, width: float = 8.27, height: float = 11.69) -> Dict[str, Any]:
        """
        设置页面尺寸
        
        Args:
            width: 页面宽度（英寸），默认A4
            height: 页面高度（英寸），默认A4
        
        Returns:
            包含操作结果的字典
        """
        try:
            if not self.current_document:
                return {
                    "success": False,
                    "error": "没有当前文档，请先创建文档"
                }
            
            section = self.current_document.sections[0]
            section.page_width = Inches(width)
            section.page_height = Inches(height)
            
            # 保存文档
            if self.current_file_path:
                self.current_document.save(self.current_file_path)
            
            logger.info(f"✅ 页面尺寸设置成功: {width}x{height} 英寸")
            
            return {
                "success": True,
                "message": f"页面尺寸设置为 {width}x{height} 英寸",
                "document_path": self.current_file_path
            }
            
        except Exception as e:
            error_msg = f"设置页面尺寸失败: {str(e)}"
            logger.error(f"❌ {error_msg}")
            return {
                "success": False,
                "error": error_msg
            }
    
    def add_heading(self, text: str, level: int = 1) -> Dict[str, Any]:
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别 (0-9)
        
        Returns:
            包含操作结果的字典
        """
        try:
            if not self.current_document:
                return {
                    "success": False,
                    "error": "没有当前文档，请先创建文档"
                }
            
            heading = self.current_document.add_heading(text, level)
            
            # 保存文档
            if self.current_file_path:
                self.current_document.save(self.current_file_path)
            
            logger.info(f"✅ 标题添加成功: {text}")
            
            return {
                "success": True,
                "message": f"标题添加成功: {text}",
                "document_path": self.current_file_path
            }
            
        except Exception as e:
            error_msg = f"添加标题失败: {str(e)}"
            logger.error(f"❌ {error_msg}")
            return {
                "success": False,
                "error": error_msg
            }
    
    def get_current_document_info(self) -> Dict[str, Any]:
        """
        获取当前文档信息
        
        Returns:
            包含文档信息的字典
        """
        if not self.current_document:
            return {
                "success": False,
                "error": "没有当前文档"
            }
        
        try:
            paragraph_count = len(self.current_document.paragraphs)
            file_size = os.path.getsize(self.current_file_path) if self.current_file_path and os.path.exists(self.current_file_path) else 0
            
            return {
                "success": True,
                "document_path": self.current_file_path,
                "paragraph_count": paragraph_count,
                "file_size": file_size,
                "file_exists": os.path.exists(self.current_file_path) if self.current_file_path else False
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"获取文档信息失败: {str(e)}"
            }

# 全局实例
word_tools = WordDocumentTools()

def create_document(filename: str, title: str = None, output_dir: str = None) -> Dict[str, Any]:
    """创建Word文档的全局函数"""
    return word_tools.create_document(filename, title, output_dir)

def add_paragraph(text: str, **kwargs) -> Dict[str, Any]:
    """添加正文段落的全局函数"""
    return word_tools.add_paragraph(text, **kwargs)

def set_page_size(width: float = 8.27, height: float = 11.69) -> Dict[str, Any]:
    """设置页面尺寸的全局函数"""
    return word_tools.set_page_size(width, height)

def add_heading(text: str, level: int = 1) -> Dict[str, Any]:
    """添加标题的全局函数"""
    return word_tools.add_heading(text, level)

def get_document_info() -> Dict[str, Any]:
    """获取文档信息的全局函数"""
    return word_tools.get_current_document_info() 