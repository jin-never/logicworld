#!/usr/bin/env python3
"""
统一工具系统 - 移除MCP依赖，提供简单可靠的工具调用接口

设计原则：
1. 简单：直接调用，无需外部进程
2. 可靠：内置工具，无网络依赖  
3. 明确：清晰的错误信息和参数验证
4. 扩展：易于添加新工具

Authors: AI Assistant
Date: 2025-08-15
"""

import json
import logging
import asyncio
import os
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple
from datetime import datetime

# Word文档工具
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class UnifiedToolSystem:
    """统一工具系统 - 简单、可靠、无外部依赖"""
    
    def __init__(self):
        self.tools = {}
        self.output_directory = self._setup_output_directory()
        self._register_builtin_tools()
        logger.info(f"✅ [工具系统] 初始化完成，输出目录: {self.output_directory}")
        
    def _setup_output_directory(self) -> Path:
        """设置输出目录"""
        # 尝试多个可能的输出路径
        possible_paths = [
            Path(r"C:\Users\ZhuanZ\Desktop\AI作品"),
            Path("output"),
            Path("../output"),
            Path(".")
        ]
        
        for path in possible_paths:
            try:
                path.mkdir(exist_ok=True)
                if path.exists() and os.access(path, os.W_OK):
                    return path
            except Exception:
                continue
                
        # 默认使用当前目录
        return Path(".")
        
    def _register_builtin_tools(self):
        """注册内置工具"""
        
        # Word文档工具
        if DOCX_AVAILABLE:
            self.tools.update({
                "create_document": {
                    "function": self._create_document,
                    "description": "创建Word文档",
                    "required_params": ["filename"],
                    "optional_params": ["title", "author", "content"]
                },
                "add_paragraph": {
                    "function": self._add_paragraph,
                    "description": "向文档添加段落",
                    "required_params": ["filename", "text"],
                    "optional_params": ["style"]
                },
                "add_heading": {
                    "function": self._add_heading,
                    "description": "向文档添加标题",
                    "required_params": ["filename", "text"],
                    "optional_params": ["level"]
                }
            })
        
        # 文本工具（无依赖）
        self.tools.update({
            "create_text_file": {
                "function": self._create_text_file,
                "description": "创建文本文件",
                "required_params": ["filename", "content"],
                "optional_params": []
            },
            "list_tools": {
                "function": self._list_tools,
                "description": "列出所有可用工具",
                "required_params": [],
                "optional_params": []
            }
        })
        
    async def execute_tool(self, tool_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """执行工具调用"""
        try:
            # 验证工具存在
            if tool_name not in self.tools:
                available_tools = list(self.tools.keys())
                return {
                    "status": "error",
                    "error": f"Unknown tool: {tool_name}",
                    "available_tools": available_tools
                }
                
            tool_info = self.tools[tool_name]
            
            # 验证必需参数
            missing_params = []
            for param in tool_info["required_params"]:
                if param not in parameters:
                    missing_params.append(param)
                    
            if missing_params:
                return {
                    "status": "error", 
                    "error": f"Missing required parameters: {', '.join(missing_params)}",
                    "required_params": tool_info["required_params"]
                }
                
            # 执行工具
            result = await tool_info["function"](**parameters)
            
            return {
                "status": "success",
                "result": result,
                "tool": tool_name,
                "executed_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "tool": tool_name
            }
            
    # ==================== Word文档工具 ====================
    
    async def _create_document(self, filename: str, title: str = "", 
                              author: str = "AI助手", content: str = "") -> str:
        """创建Word文档"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法创建Word文档")
            
        if not filename.endswith('.docx'):
            filename = f"{filename}.docx"
            
        file_path = self.output_directory / filename
        
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, 0)
            
        if content:
            doc.add_paragraph(content)
            
        doc.save(str(file_path))
        
        return f"文档创建成功: {file_path.absolute()}"
        
    async def _add_paragraph(self, filename: str, text: str, style: str = "Normal") -> str:
        """向文档添加段落"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装")
            
        file_path = self.output_directory / filename
        if not file_path.exists():
            raise Exception(f"文档不存在: {file_path}")
            
        doc = Document(str(file_path))
        doc.add_paragraph(text)
        doc.save(str(file_path))
        
        return f"段落添加成功"
        
    async def _add_heading(self, filename: str, text: str, level: int = 1) -> str:
        """向文档添加标题"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装")
            
        file_path = self.output_directory / filename
        if not file_path.exists():
            raise Exception(f"文档不存在: {file_path}")
            
        doc = Document(str(file_path))
        doc.add_heading(text, level)
        doc.save(str(file_path))
        
        return f"标题添加成功: {text}"
        
    # ==================== 文本文件工具 ====================
    
    async def _create_text_file(self, filename: str, content: str) -> str:
        """创建文本文件"""
        if not filename.endswith('.txt'):
            filename = f"{filename}.txt"
            
        file_path = self.output_directory / filename
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        return f"文本文件创建成功: {file_path.absolute()}"
        
    async def _list_tools(self) -> str:
        """列出所有可用工具"""
        tools_info = []
        for tool_name, tool_info in self.tools.items():
            tools_info.append(f"- {tool_name}: {tool_info['description']}")
            
        return f"可用工具 ({len(self.tools)} 个):\n" + "\n".join(tools_info)
        
    def get_available_tools(self) -> List[str]:
        """获取可用工具列表"""
        return list(self.tools.keys())

# 全局工具系统实例
unified_tool_system = UnifiedToolSystem()

# 兼容旧接口的函数
async def execute_tool_call(tool_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
    """执行工具调用 - 兼容旧接口"""
    return await unified_tool_system.execute_tool(tool_name, parameters)

def get_available_tools() -> List[str]:
    """获取可用工具列表 - 兼容旧接口"""  
    return unified_tool_system.get_available_tools()

if __name__ == "__main__":
    print("🧪 测试统一工具系统...")
    print(f"可用工具: {get_available_tools()}")
    
    async def test():
        result = await execute_tool_call("list_tools", {})
        print(f"测试结果: {result}")
        
    asyncio.run(test()) 