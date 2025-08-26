#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from fastapi import APIRouter, HTTPException, File, UploadFile, Query
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
import logging
import os
import json
import tempfile
import mimetypes

logger = logging.getLogger(__name__)
router = APIRouter()

# 新增：解析上传的文档
@router.post("/api/document/parse-upload")
async def parse_uploaded_document(file: UploadFile = File(...)) -> Dict[str, Any]:
    """
    解析用户上传的Word文档
    """
    logger.info(f"🔍 收到文档解析请求: {file.filename}")
    logger.info(f"🔍 文件大小: {file.size} bytes")
    logger.info(f"🔍 文件类型: {file.content_type}")
    
    try:
        # 验证文件类型
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="只支持.docx格式的文件")
        
        # 创建临时文件 - 修复WinError 32问题
        temp_file_path = None
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            # 写入上传的文件内容
            content = await file.read()
            temp_file.write(content)
            temp_file.flush()
            temp_file_path = temp_file.name
            
            try:
            # 使用docx库解析文档 (文件已关闭，现在可以安全访问)
            doc = Document(temp_file_path)
                
                content_list = []
                run_id = 0
                
                for para_idx, paragraph in enumerate(doc.paragraphs):
                logger.info(f"📄 处理段落 {para_idx}: '{paragraph.text[:50]}...'")
                    if paragraph.text.strip():
                        para_content = {
                            "paragraph_id": para_idx,
                            "runs": []
                        }
                        
                        char_offset = 0
                        for run in paragraph.runs:
                            if run.text:
                                run_info = {
                                    "run_id": run_id,
                                    "text": run.text,
                                    "font_name": run.font.name if run.font.name else "默认字体",
                                    "start_char": char_offset,
                                    "end_char": char_offset + len(run.text)
                                }
                                para_content["runs"].append(run_info)
                                char_offset += len(run.text)
                                run_id += 1
                        
                        if para_content["runs"]:
                            content_list.append(para_content)
                
            result = {
                    "filename": file.filename,
                    "content": content_list,
                    "total_paragraphs": len(content_list),
                    "message": "文档解析成功"
                }
            
            logger.info(f"✅ 文档解析完成: {len(content_list)} 个段落")
            logger.info(f"✅ 返回数据: {result}")
            
            return result
                
            finally:
                # 清理临时文件
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                    logger.info(f"🗑️ 临时文件已清理: {temp_file_path}")
                except Exception as e:
                    logger.warning(f"⚠️ 清理临时文件失败: {e}")
                
    except Exception as e:
        logger.error(f"Error parsing uploaded document {file.filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")

@router.get("/api/document/preview/{filename}")
async def get_document_preview(filename: str) -> Dict[str, Any]:
    try:
        possible_paths = [
            Path(filename),
            Path(f"C:\\Users\\ZhuanZ\\Desktop\\AI\\{filename}"),
            Path(f"C:\\Users\\ZhuanZ\\Desktop\\AI作品\\{filename}"),
            Path(f"output/{filename}"),
            Path(f"backend/output/{filename}"),
            Path(f"{filename}")
        ]
        
        document_path = None
        for path in possible_paths:
            if path.exists() and path.suffix.lower() == '.docx':
                document_path = path
                break
        
        if not document_path:
            raise HTTPException(status_code=404, detail=f"Document not found: {filename}")
        
        logger.info(f"Loading document from: {document_path}")
        
        # 读取Word文档
        doc = Document(str(document_path))
        
        content = []
        run_id = 0
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_content = {
                    "paragraph_id": para_idx,
                    "runs": []
                }
                
                char_offset = 0
                for run in paragraph.runs:
                    if run.text:
                        run_info = {
                            "run_id": run_id,
                            "text": run.text,
                            "font_name": run.font.name if run.font.name else "默认字体",
                            "start_char": char_offset,
                            "end_char": char_offset + len(run.text)
                        }
                        para_content["runs"].append(run_info)
                        char_offset += len(run.text)
                        run_id += 1
                
                if para_content["runs"]:
                    content.append(para_content)
        
        return {
            "filename": filename,
            "content": content,
            "total_paragraphs": len(content),
            "path": str(document_path)
        }
        
    except Exception as e:
        logger.error(f"Error loading document {filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error loading document: {str(e)}")

# 新增：设置材料节点文件夹路径
@router.post("/api/material-node/set-folder")
async def set_material_node_folder(folder_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    设置材料节点的文件夹路径
    """
    try:
        folder_path = folder_data.get('folderPath')
        node_id = folder_data.get('nodeId')
        
        if not folder_path:
            return {
                "success": False,
                "message": "文件夹路径不能为空"
            }
        
        # 验证文件夹路径是否存在
        if not Path(folder_path).exists():
            return {
                "success": False,
                "message": f"文件夹路径不存在: {folder_path}"
            }
        
        # TODO: 这里可以将文件夹路径保存到数据库或配置文件
        # 目前先返回成功，实际的路径会通过工作流状态保存
        
        return {
            "success": True,
            "message": "文件夹路径设置成功",
            "folderPath": folder_path
        }
        
    except Exception as e:
        logger.error(f"Error setting material node folder: {str(e)}")
        return {
            "success": False,
            "message": f"设置文件夹路径失败: {str(e)}"
        }

# 新增：用户选择文件夹的API端点
@router.post("/api/material-node/browse-folder")
async def browse_folder() -> Dict[str, Any]:
    """
    打开文件夹选择对话框（这个API主要用于触发前端的文件夹选择）
    """
    return {
        "success": True,
        "message": "请在前端选择文件夹",
        "action": "open_folder_dialog"
    }

# 获取材料节点文档列表
@router.get("/api/material-node/documents")
async def get_material_node_documents(workflow_id: str = None) -> Dict[str, Any]:
    """
    获取当前工作流中材料节点选择的文件夹中的所有Word文档
    """
    try:
        # 查找材料节点的文件夹路径
        material_folder = await find_material_node_folder(workflow_id)
        
        if not material_folder:
            return {
                "success": False,
                "message": "未找到材料节点或未选择文件夹",
                "documents": [],
                "action": "select_folder"  # 提示前端需要选择文件夹
            }
        
        # 扫描文件夹中的Word文档
        documents = []
        folder_path = Path(material_folder)
        
        if folder_path.exists() and folder_path.is_dir():
            for file_path in folder_path.glob("*.docx"):
                if not file_path.name.startswith("~$"):  # 排除临时文件
                    try:
                        # 获取文件基本信息
                        stat = file_path.stat()
                        documents.append({
                            "filename": file_path.name,
                            "path": str(file_path),
                            "size": stat.st_size,
                            "modified": stat.st_mtime,
                            "displayName": file_path.stem  # 不含扩展名的文件名
                        })
                    except Exception as e:
                        logger.warning(f"Error reading file info for {file_path}: {e}")
                        continue
        
        # 按修改时间排序（最新的在前）
        documents.sort(key=lambda x: x["modified"], reverse=True)
        
        return {
            "success": True,
            "folder": material_folder,
            "documents": documents,
            "count": len(documents)
        }
        
    except Exception as e:
        logger.error(f"Error getting material node documents: {str(e)}")
        return {
            "success": False,
            "message": f"获取文档列表失败: {str(e)}",
            "documents": []
        }

async def find_material_node_folder(workflow_id: str = None) -> str:
    """
    查找当前工作流中材料节点设置的文件夹路径
    """
    try:
        # 1. 尝试从工作流状态文件中查找
        workflow_states_dir = Path("workflow_states")
        if workflow_states_dir.exists():
            for state_file in workflow_states_dir.glob("*.json"):
                try:
                    with open(state_file, 'r', encoding='utf-8') as f:
                        workflow_data = json.load(f)
                        
                    # 查找材料节点
                    if 'nodes' in workflow_data:
                        for node in workflow_data['nodes']:
                            if (node.get('type') == 'material-node' or 
                                '材料' in node.get('data', {}).get('label', '') or
                                node.get('data', {}).get('nodeType') == 'material-node'):
                                
                                # 优先使用folderPath（用户选择的文件夹路径）
                                folder_path = node.get('data', {}).get('folderPath')
                                if folder_path:
                                    # 如果是完整路径，直接使用
                                    if Path(folder_path).exists():
                                        logger.info(f"Found material node folder from folderPath: {folder_path}")
                                        return folder_path
                                    
                                    # 如果只是文件夹名称，尝试构建完整路径
                                    possible_base_paths = [
                                        "C:\\Users\\ZhuanZ\\Desktop",
                                        "C:\\Users\\ZhuanZ\\Desktop\\AI作品",
                                        "C:\\Users\\ZhuanZ\\Desktop\\AI",
                                        "output",
                                        "backend/output"
                                    ]
                                    
                                    for base_path in possible_base_paths:
                                        full_path = Path(base_path) / folder_path
                                        if full_path.exists():
                                            logger.info(f"Found material node folder by combining base path: {full_path}")
                                            return str(full_path)
                                
                                # 如果没有folderPath，尝试从selectedFiles推断
                                selected_files = node.get('data', {}).get('selectedFiles', [])
                                if selected_files and len(selected_files) > 0:
                                    # 假设所有文件都在同一个文件夹中
                                    first_file = selected_files[0]
                                    if isinstance(first_file, dict) and 'path' in first_file:
                                        file_path = Path(first_file['path'])
                                        folder_path = str(file_path.parent)
                                        if Path(folder_path).exists():
                                            logger.info(f"Inferred material node folder from files: {folder_path}")
                                            return folder_path
                except Exception as e:
                    logger.warning(f"Error reading workflow state {state_file}: {e}")
                    continue
        
        # 2. 默认文件夹路径
        default_folders = [
            "C:\\Users\\ZhuanZ\\Desktop\\AI HTML",  # 用户指定的路径
            "C:\\Users\\ZhuanZ\\Desktop\\AI作品",
            "C:\\Users\\ZhuanZ\\Desktop\\AI",
            "output",
            "backend/output"
        ]
        
        for folder in default_folders:
            if Path(folder).exists():
                logger.info(f"Using default folder: {folder}")
                return folder
        
        return None
        
    except Exception as e:
        logger.error(f"Error finding material node folder: {str(e)}")
        return None

@router.get("/api/folder/scan")
async def scan_folder(path: str = Query(..., description="要扫描的文件夹路径")) -> Dict[str, Any]:
    """
    扫描指定文件夹中的Word文档
    """
    try:
        folder_path = Path(path)
        
        # 验证路径是否存在
        if not folder_path.exists():
            return {
                "success": False,
                "message": f"路径不存在: {path}",
                "files": []
            }
        
        # 验证是否为文件夹
        if not folder_path.is_dir():
            return {
                "success": False,
                "message": f"指定路径不是文件夹: {path}",
                "files": []
            }
        
        # 扫描Word文档
        word_extensions = ['.docx', '.doc']
        found_files = []
        
        try:
            for file_path in folder_path.iterdir():
                if file_path.is_file() and file_path.suffix.lower() in word_extensions:
                    # 跳过临时文件
                    if file_path.name.startswith("~$"):
                        continue
                        
                    # 获取文件信息
                    stat_info = file_path.stat()
                    file_info = {
                        "name": file_path.name,
                        "fullPath": str(file_path),
                        "size": stat_info.st_size,
                        "lastModified": stat_info.st_mtime * 1000,  # 转换为毫秒
                        "type": mimetypes.guess_type(file_path.name)[0] or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    }
                    found_files.append(file_info)
            
            # 按修改时间排序（最新的在前）
            found_files.sort(key=lambda x: x["lastModified"], reverse=True)
            
            logger.info(f"在路径 {path} 中找到 {len(found_files)} 个Word文档")
            
            return {
                "success": True,
                "message": f"找到 {len(found_files)} 个Word文档",
                "files": found_files,
                "scannedPath": str(folder_path)
            }
            
        except PermissionError:
            return {
                "success": False,
                "message": f"没有权限访问路径: {path}",
                "files": []
            }
            
    except Exception as e:
        logger.error(f"扫描文件夹时出错: {str(e)}")
        return {
            "success": False,
            "message": f"扫描失败: {str(e)}",
            "files": []
        } 